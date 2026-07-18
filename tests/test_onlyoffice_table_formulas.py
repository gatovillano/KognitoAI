import os
import sys
import types
import uuid
import asyncio
from pathlib import Path
from types import SimpleNamespace

import pytest
import docx
import openpyxl

# Fakes para dependencias opcionales
passlib_module = types.ModuleType("passlib")
passlib_context_module = types.ModuleType("passlib.context")

class _FakeCryptContext:
    def __init__(self, *args, **kwargs):
        pass
    def verify(self, *args, **kwargs):
        return True
    def hash(self, value):
        return str(value)

passlib_context_module.CryptContext = _FakeCryptContext
passlib_module.context = passlib_context_module
sys.modules.setdefault("passlib", passlib_module)
sys.modules.setdefault("passlib.context", passlib_context_module)

multipart_module = types.ModuleType("multipart")
multipart_submodule = types.ModuleType("multipart.multipart")
def _fake_parse_options_header(value):
    return value, {}
multipart_submodule.parse_options_header = _fake_parse_options_header
multipart_module.multipart = multipart_submodule
sys.modules.setdefault("multipart", multipart_module)
sys.modules.setdefault("multipart.multipart", submodule_multipart := multipart_submodule)

fitz_module = types.ModuleType("fitz")
sys.modules.setdefault("fitz", fitz_module)

import skills.onlyoffice_skill.scripts.edit_onlyoffice_document_tool as edit_tool_module
from skills.onlyoffice_skill.scripts.edit_onlyoffice_document_tool import EditOnlyOfficeDocumentTool

class _FakeResult:
    def __init__(self, document):
        self._document = document
    def scalar_one_or_none(self):
        return self._document

class _FakeSession:
    def __init__(self, document):
        self._document = document
    async def __aenter__(self):
        return self
    async def __aexit__(self, exc_type, exc, tb):
        return False
    async def execute(self, stmt):
        return _FakeResult(self._document)
    async def commit(self):
        pass

def test_docx_styled_table_with_totals(monkeypatch, tmp_path):
    account_id = str(uuid.uuid4())
    document_id = uuid.uuid4()
    docx_file = tmp_path / account_id / "reporte.docx"
    docx_file.parent.mkdir(parents=True, exist_ok=True)

    # Crear docx base
    doc_obj = docx.Document()
    doc_obj.save(str(docx_file))

    doc_meta = SimpleNamespace(
        id=document_id,
        account_id=uuid.UUID(account_id),
        filename="reporte.docx",
        extension="docx",
        file_path=f"{account_id}/reporte.docx",
        updated_at=None,
    )

    monkeypatch.setattr(edit_tool_module, "SessionLocal", lambda: _FakeSession(doc_meta))
    monkeypatch.setattr(edit_tool_module, "resolve_onlyoffice_file_path", lambda _: docx_file)

    tool = EditOnlyOfficeDocumentTool(account_id=account_id)

    table_data = [
        ["Producto", "Ventas Q1", "Ventas Q2"],
        ["Laptop KAI Pro", "$1,200.00", "$1,500.00"],
        ["Servidor Cloud", "$3,000.00", "$2,800.00"],
        ["TOTAL", "AUTO", "AUTO"]
    ]

    res = asyncio.run(
        tool._arun(
            document_id=str(document_id),
            action="insert_styled_table",
            table_data=table_data,
            theme="emerald_green",
            has_total_row=True,
            auto_calculate_totals=True
        )
    )

    assert "✅" in res
    assert "emerald_green" in res

    # Abrir docx y verificar la tabla
    loaded_doc = docx.Document(str(docx_file))
    assert len(loaded_doc.tables) == 1
    t = loaded_doc.tables[0]
    assert len(t.rows) == 4
    # Verificar total calculado
    total_q1_text = t.cell(3, 1).paragraphs[0].text
    total_q2_text = t.cell(3, 2).paragraphs[0].text
    assert "4,200.00" in total_q1_text
    assert "4,300.00" in total_q2_text


def test_xlsx_create_table_with_formulas(monkeypatch, tmp_path):
    account_id = str(uuid.uuid4())
    document_id = uuid.uuid4()
    xlsx_file = tmp_path / account_id / "finanzas.xlsx"
    xlsx_file.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    wb.save(str(xlsx_file))

    doc_meta = SimpleNamespace(
        id=document_id,
        account_id=uuid.UUID(account_id),
        filename="finanzas.xlsx",
        extension="xlsx",
        file_path=f"{account_id}/finanzas.xlsx",
        updated_at=None,
    )

    monkeypatch.setattr(edit_tool_module, "SessionLocal", lambda: _FakeSession(doc_meta))
    monkeypatch.setattr(edit_tool_module, "resolve_onlyoffice_file_path", lambda _: xlsx_file)

    tool = EditOnlyOfficeDocumentTool(account_id=account_id)

    table_data = [
        ["Categoría", "Monto Q1", "Monto Q2"],
        ["Software", 1500, 2000],
        ["Hardware", 3500, 4200],
        ["Servicios", 800, 1100],
        ["TOTAL", "AUTO", "AUTO"]
    ]

    res = asyncio.run(
        tool._arun(
            document_id=str(document_id),
            action="xlsx_create_table",
            table_data=table_data,
            theme="corporate_blue",
            has_total_row=True,
            column_formats={"B": "$#,##0.00", "C": "$#,##0.00"},
            start_cell="B2"
        )
    )

    assert "✅" in res
    assert "corporate_blue" in res

    # Cargar workbook y verificar fórmulas y formato
    loaded_wb = openpyxl.load_workbook(str(xlsx_file))
    ws = loaded_wb.active

    assert ws["B2"].value == "Categoría"
    assert ws["C2"].value == "Monto Q1"
    assert ws["D2"].value == "Monto Q2"

    # Verificar fórmulas en fila de total
    assert ws["C6"].value == "=SUM(C3:C5)"
    assert ws["D6"].value == "=SUM(D3:D5)"
    assert ws["C3"].number_format == "$#,##0.00"
