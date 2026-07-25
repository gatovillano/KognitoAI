# api/onlyoffice.py

import logging
import mimetypes
import os
import uuid
import secrets
import httpx
import jwt
from datetime import datetime
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query, Request
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, or_
from pydantic import BaseModel

from core.database import (
    SessionLocal,
    Document,
    Account,
    DocumentFolder,
    ChatThread,
    SharedConversationLink,
    OnlyOfficeDocumentShare,
    OnlyOfficeDocumentChat,
)
from core.dependencies import get_db_session
from utils.security import get_current_account_id, check_workspace_permission, decode_access_token
from core.config import settings
from core.onlyoffice_storage import (
    build_onlyoffice_relative_path,
    ensure_onlyoffice_account_dir,
    get_onlyoffice_docs_root,
    resolve_onlyoffice_file_path,
)
from utils.document_parser import extract_text_and_metadata_from_document
from core.memory_manager import delete_document_chunks, process_document_for_rag

def check_office_libs():
    """Verifica dinámicamente qué librerías de oficina están instaladas."""
    docx_lib = None
    openpyxl_lib = None
    pptx_lib = None
    
    try:
        import docx
        docx_lib = docx
    except ImportError:
        pass

    try:
        from openpyxl import Workbook
        openpyxl_lib = Workbook
    except ImportError:
        pass

    try:
        from pptx import Presentation
        pptx_lib = Presentation
    except ImportError:
        pass
        
    success = any([docx_lib, openpyxl_lib, pptx_lib])
    return success, docx_lib, openpyxl_lib, pptx_lib

logger = logging.getLogger(__name__)
router = APIRouter()


class CreateDocumentShareRequest(BaseModel):
    scope: str
    target_account_id: Optional[str] = None
    can_edit: bool = True


class DeleteDocumentShareResponse(BaseModel):
    message: str


async def _can_access_document(
    doc: Document,
    current_account_id: str,
    db: AsyncSession,
    token: Optional[str] = None,
) -> bool:
    if str(doc.account_id) == current_account_id:
        return True

    if doc.workspace_id:
        try:
            if await check_workspace_permission(
                current_account_id,
                str(doc.workspace_id),
                db,
                required_roles=["owner", "editor", "viewer"],
            ):
                return True
        except Exception:
            pass

    share_stmt = select(OnlyOfficeDocumentShare).where(
        OnlyOfficeDocumentShare.document_id == doc.id,
        OnlyOfficeDocumentShare.shared_with_account_id == uuid.UUID(current_account_id),
    )
    share_row = (await db.execute(share_stmt)).scalars().first()
    if share_row:
        return True

    if token:
        public_stmt = select(OnlyOfficeDocumentShare).where(
            OnlyOfficeDocumentShare.document_id == doc.id,
            OnlyOfficeDocumentShare.is_public == True,
            OnlyOfficeDocumentShare.token == token,
        )
        public_share = (await db.execute(public_stmt)).scalars().first()
        if public_share:
            return True

    return False

# Directorio para los documentos de OnlyOffice
# Usar la ruta configurada en settings
DOCUMENTS_ROOT = str(get_onlyoffice_docs_root())
os.makedirs(DOCUMENTS_ROOT, exist_ok=True)

@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    workspace_id: Optional[str] = Form(None),
    folder_id: Optional[str] = Form(None),
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Sube un documento para ser editado con OnlyOffice."""
    # Verificar permisos de workspace si se proporciona
    if workspace_id:
        if not await check_workspace_permission(current_account_id, workspace_id, db, required_roles=["owner", "editor"]):
            raise HTTPException(status_code=403, detail="No tienes permiso para subir documentos a este workspace.")

    account_id_uuid = uuid.UUID(current_account_id)
    
    # Obtener nombre y extensión
    filename = file.filename
    extension = filename.split('.')[-1].lower() if '.' in filename else ""
    
    # Para actuar como una nube, no restringimos las extensiones en la subida.


    # Guardar archivo físicamente
    unique_filename = f"{uuid.uuid4()}.{extension}"
    account_obj = await db.get(Account, uuid.UUID(current_account_id))
    cloud_storage_path = getattr(account_obj, "cloud_storage_path", None)
    user_dir = ensure_onlyoffice_account_dir(current_account_id, cloud_storage_path=cloud_storage_path)
    
    file_path = user_dir / unique_filename
    
    try:
        content = await file.read()
        with open(file_path, "wb") as buffer:
            buffer.write(content)
    except Exception as e:
        logger.error(f"Error al guardar archivo OnlyOffice: {e}")
        raise HTTPException(status_code=500, detail="Error al guardar el archivo en el servidor.")
    
    # Obtener folder_id
    fid = uuid.UUID(folder_id) if folder_id and folder_id != "null" else None
    ws_id = uuid.UUID(workspace_id) if workspace_id and workspace_id != "null" else None
    
    # Inherit workspace_id from folder if not provided
    if fid and not ws_id:
        folder = await db.get(DocumentFolder, fid)
        if folder:
            ws_id = folder.workspace_id

    # Guardar metadatos en la base de datos
    new_doc = Document(
        account_id=account_id_uuid,
        workspace_id=ws_id,
        filename=filename,
        extension=extension,
        file_path=build_onlyoffice_relative_path(current_account_id, unique_filename),
        folder_id=fid
    )
    db.add(new_doc)
    await db.commit()
    await db.refresh(new_doc)
    
    return {"message": "Documento subido correctamente", "id": new_doc.id, "filename": filename}

@router.get("/list")
async def list_documents(
    workspace_id: Optional[str] = Query(None),
    folder_id: Optional[str] = Query(None),
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Lista los documentos editables del usuario o workspace."""
    from core.database import Workspace

    account_uuid = uuid.UUID(current_account_id)
    shared_doc_ids_stmt = select(OnlyOfficeDocumentShare.document_id).where(
        OnlyOfficeDocumentShare.shared_with_account_id == account_uuid
    )
    
    stmt = select(Document, Workspace)\
        .outerjoin(Workspace, Document.workspace_id == Workspace.id)\
        .where(
            or_(
                Document.account_id == account_uuid,
                Document.id.in_(shared_doc_ids_stmt),
            )
        )
    
    if workspace_id and workspace_id != "all":
        stmt = stmt.where(Document.workspace_id == uuid.UUID(workspace_id))
    
    if folder_id:
        if folder_id == "null":
            stmt = stmt.where(Document.folder_id == None)
        else:
            stmt = stmt.where(Document.folder_id == uuid.UUID(folder_id))
        
    result = await db.execute(stmt)
    rows = result.all()
    
    documents = []
    for doc, workspace in rows:
        documents.append({
            "id": str(doc.id),
            "filename": doc.filename,
            "extension": doc.extension,
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
            "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
            "workspace_id": str(doc.workspace_id) if doc.workspace_id else None,
            "workspace_name": workspace.name if workspace else None,
            "workspace_color": workspace.color if workspace else None,
            "folder_id": str(doc.folder_id) if doc.folder_id else None,
            "is_owner": str(doc.account_id) == current_account_id,
        })
    
    return documents

@router.get("/folders")
async def list_folders(
    workspace_id: Optional[str] = Query(None),
    parent_id: Optional[str] = Query(None),
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Lista las carpetas de OnlyOffice."""
    from core.database import Workspace
    
    stmt = select(DocumentFolder, Workspace)\
        .outerjoin(Workspace, DocumentFolder.workspace_id == Workspace.id)\
        .where(DocumentFolder.account_id == uuid.UUID(current_account_id))
    
    if workspace_id and workspace_id != "all":
        stmt = stmt.where(DocumentFolder.workspace_id == uuid.UUID(workspace_id))
    
    if parent_id:
        if parent_id == "null":
            stmt = stmt.where(DocumentFolder.parent_id == None)
        else:
            stmt = stmt.where(DocumentFolder.parent_id == uuid.UUID(parent_id))
            
    result = await db.execute(stmt)
    rows = result.all()
    
    folders = []
    for folder, workspace in rows:
        folders.append({
            "id": str(folder.id),
            "name": folder.name,
            "parent_id": str(folder.parent_id) if folder.parent_id else None,
            "workspace_id": str(folder.workspace_id) if folder.workspace_id else None,
            "workspace_name": workspace.name if workspace else None,
            "workspace_color": workspace.color if workspace else None,
            "created_at": folder.created_at.isoformat() if folder.created_at else None
        })
    
    return folders

async def update_folder_workspace_recursive(folder_id: uuid.UUID, workspace_id: Optional[uuid.UUID], db: AsyncSession):
    # 1. Update all documents in this folder
    docs_stmt = select(Document).where(Document.folder_id == folder_id)
    docs = (await db.execute(docs_stmt)).scalars().all()
    for doc in docs:
        doc.workspace_id = workspace_id
        # Update associated chat thread if it exists
        mapping_stmt = select(OnlyOfficeDocumentChat).where(OnlyOfficeDocumentChat.document_id == doc.id)
        mapping = (await db.execute(mapping_stmt)).scalars().first()
        if mapping:
            thread = await db.get(ChatThread, mapping.thread_id)
            if thread:
                thread.workspace_id = workspace_id

    # 2. Update all subfolders recursively
    subfolders_stmt = select(DocumentFolder).where(DocumentFolder.parent_id == folder_id)
    subfolders = (await db.execute(subfolders_stmt)).scalars().all()
    for subfolder in subfolders:
        subfolder.workspace_id = workspace_id
        await update_folder_workspace_recursive(subfolder.id, workspace_id, db)

@router.post("/folders")
async def create_folder(
    name: str = Form(...),
    workspace_id: Optional[str] = Form(None),
    parent_id: Optional[str] = Form(None),
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Crea una nueva carpeta para OnlyOffice."""
    parent_uuid = uuid.UUID(parent_id) if parent_id and parent_id != "null" else None
    ws_uuid = uuid.UUID(workspace_id) if workspace_id and workspace_id != "null" else None
    
    # Inherit workspace_id from parent folder if not provided
    if parent_uuid and not ws_uuid:
        parent_folder = await db.get(DocumentFolder, parent_uuid)
        if parent_folder:
            ws_uuid = parent_folder.workspace_id

    new_folder = DocumentFolder(
        account_id=uuid.UUID(current_account_id),
        workspace_id=ws_uuid,
        parent_id=parent_uuid,
        name=name
    )
    db.add(new_folder)
    await db.commit()
    await db.refresh(new_folder)
    return {"message": "Carpeta creada", "id": str(new_folder.id)}

@router.post("/create")
async def create_document(
    type: str = Form(...),
    name: str = Form(...),
    workspace_id: Optional[str] = Form(None),
    folder_id: Optional[str] = Form(None),
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Crea un nuevo documento vacío editado con OnlyOffice."""
    libs_ok, docx_lib, Workbook_lib, Presentation_lib = check_office_libs()
    
    if not libs_ok:
        logger.error("Error: Librerías de oficina (python-docx, openpyxl, python-pptx) no encontradas en el entorno Docker.")
        raise HTTPException(
            status_code=500, 
            detail="Las librerías de generación de documentos no se detectan en el contenedor. Si ya las instaló con 'pip install', por favor intente de nuevo. Si persiste, reconstruya el contenedor (docker-compose build)."
        )

    extension = ""
    if type == 'word': extension = 'docx'
    elif type == 'excel': extension = 'xlsx'
    elif type == 'powerpoint': extension = 'pptx'
    else: extension = 'docx'

    if not name.endswith(f".{extension}"):
        full_filename = f"{name}.{extension}"
    else:
        full_filename = name

    # Guardar archivo físicamente
    unique_filename = f"{uuid.uuid4()}.{extension}"
    account_obj = await db.get(Account, uuid.UUID(current_account_id))
    cloud_storage_path = getattr(account_obj, "cloud_storage_path", None)
    user_dir = ensure_onlyoffice_account_dir(current_account_id, cloud_storage_path=cloud_storage_path)
    file_path = os.path.join(str(user_dir), unique_filename)

    try:
        if type == 'word':
            doc = docx_lib.Document()
            doc.save(file_path)
        elif type == 'excel':
            wb = Workbook_lib()
            wb.save(file_path)
        elif type == 'powerpoint':
            prs = Presentation_lib()
            prs.save(file_path)
    except Exception as e:
        logger.error(f"Error al crear archivo OnlyOffice: {e}")
        raise HTTPException(status_code=500, detail="Error al crear el archivo físico.")

    fid = uuid.UUID(folder_id) if folder_id and folder_id != "null" else None
    ws_id = uuid.UUID(workspace_id) if workspace_id and workspace_id != "null" else None
    
    # Inherit workspace_id from folder if not provided
    if fid and not ws_id:
        folder = await db.get(DocumentFolder, fid)
        if folder:
            ws_id = folder.workspace_id

    # Guardar en DB
    new_doc = Document(
        account_id=uuid.UUID(current_account_id),
        workspace_id=ws_id,
        folder_id=fid,
        filename=full_filename,
        extension=extension,
        file_path=os.path.join(current_account_id, unique_filename)
    )
    db.add(new_doc)
    await db.commit()
    await db.refresh(new_doc)
    
    return {"message": "Documento creado correctamente", "id": str(new_doc.id), "filename": full_filename}

@router.post("/{document_id}/meta")
async def update_document_meta(
    document_id: uuid.UUID,
    filename: Optional[str] = Form(None),
    workspace_id: Optional[str] = Form(None),
    folder_id: Optional[str] = Form(None),
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Actualiza los metadatos de un documento (moverlo o renombrarlo)."""
    doc = await db.get(Document, document_id)
    if not doc or str(doc.account_id) != current_account_id:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    
    if filename:
        # Asegurarse de mantener la extensión original si el usuario no la proporciona
        if '.' not in filename:
            filename = f"{filename}.{doc.extension}"
        doc.filename = filename

    if workspace_id:
        doc.workspace_id = uuid.UUID(workspace_id) if workspace_id != "null" else None
    
    if folder_id:
        if folder_id == "null":
            doc.folder_id = None
        else:
            fid = uuid.UUID(folder_id)
            doc.folder_id = fid
            # Inherit workspace_id from folder if not explicitly provided
            if not workspace_id:
                folder = await db.get(DocumentFolder, fid)
                if folder:
                    doc.workspace_id = folder.workspace_id
                    
    # Sync associated chat thread if it exists
    mapping_stmt = select(OnlyOfficeDocumentChat).where(OnlyOfficeDocumentChat.document_id == doc.id)
    mapping = (await db.execute(mapping_stmt)).scalars().first()
    if mapping:
        thread = await db.get(ChatThread, mapping.thread_id)
        if thread:
            thread.workspace_id = doc.workspace_id
        
    await db.commit()
    return {"message": "Metadatos actualizados"}

@router.post("/folders/{folder_id}/meta")
async def update_folder_meta(
    folder_id: uuid.UUID,
    name: Optional[str] = Form(None),
    workspace_id: Optional[str] = Form(None),
    parent_id: Optional[str] = Form(None),
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Actualiza los metadatos de una carpeta (moverla o renombrarla)."""
    folder = await db.get(DocumentFolder, folder_id)
    if not folder or str(folder.account_id) != current_account_id:
        raise HTTPException(status_code=404, detail="Carpeta no encontrada")
    
    if name:
        folder.name = name

    if workspace_id:
        folder.workspace_id = uuid.UUID(workspace_id) if workspace_id != "null" else None
        # Recursively update all subfolders and documents inside this folder
        await update_folder_workspace_recursive(folder.id, folder.workspace_id, db)
    
    if parent_id:
        if parent_id == "null":
            folder.parent_id = None
        else:
            pid = uuid.UUID(parent_id)
            folder.parent_id = pid
            # Inherit workspace from parent folder if not explicitly provided
            if not workspace_id:
                parent_folder = await db.get(DocumentFolder, pid)
                if parent_folder:
                    folder.workspace_id = parent_folder.workspace_id
                    # Recursively update all subfolders and documents inside this folder
                    await update_folder_workspace_recursive(folder.id, folder.workspace_id, db)
        
    await db.commit()
    return {"message": "Metadatos de carpeta actualizados"}

@router.delete("/folders/{folder_id}")
async def delete_folder(
    folder_id: uuid.UUID,
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Elimina una carpeta y su contenido (documentos y subcarpetas)."""
    folder = await db.get(DocumentFolder, folder_id)
    if not folder or str(folder.account_id) != current_account_id:
        raise HTTPException(status_code=404, detail="Carpeta no encontrada")
    
    async def delete_folder_recursive(fid: uuid.UUID):
        # 1. Eliminar documentos hijos
        docs_stmt = select(Document).where(Document.folder_id == fid)
        docs = (await db.execute(docs_stmt)).scalars().all()
        for doc in docs:
            file_full_path = resolve_onlyoffice_file_path(doc.file_path)
            if file_full_path.exists():
                try:
                    os.remove(file_full_path)
                except Exception as e:
                    logger.error(f"Error al eliminar archivo físico: {e}")
            await db.delete(doc)
            
        # 2. Obtener y eliminar subcarpetas
        subfolders_stmt = select(DocumentFolder).where(DocumentFolder.parent_id == fid)
        subfolders = (await db.execute(subfolders_stmt)).scalars().all()
        for subfolder in subfolders:
            await delete_folder_recursive(subfolder.id)
            await db.delete(subfolder)
            
    # Iniciar borrado recursivo de todo el contenido
    await delete_folder_recursive(folder_id)
    
    # Eliminar la carpeta original (y hacer commit final)
    await db.delete(folder)
    await db.commit()
    return {"message": "Carpeta y contenido eliminados con éxito"}

@router.get("/config/{document_id}")
async def get_onlyoffice_config(
    document_id: uuid.UUID,
    share_token: Optional[str] = Query(None),
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Genera la configuración necesaria para el editor OnlyOffice JavaScript."""
    doc = await db.get(Document, document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    if not await _can_access_document(doc, current_account_id, db, token=share_token):
        raise HTTPException(status_code=403, detail="No tienes acceso a este documento")

    can_edit = True
    if share_token:
        share_entry = (await db.execute(
            select(OnlyOfficeDocumentShare).where(OnlyOfficeDocumentShare.token == share_token)
        )).scalars().first()
        if share_entry:
            can_edit = share_entry.can_edit
    elif str(doc.account_id) != current_account_id:
        share_entry = (await db.execute(
            select(OnlyOfficeDocumentShare).where(
                OnlyOfficeDocumentShare.document_id == doc.id,
                OnlyOfficeDocumentShare.shared_with_account_id == uuid.UUID(current_account_id),
            )
        )).scalars().first()
        if share_entry:
            can_edit = share_entry.can_edit

    onlyoffice_api_url = os.getenv("ONLYOFFICE_URL", "/onlyoffice")
    
    onlyoffice_internal_backend = os.getenv(
        "ONLYOFFICE_INTERNAL_BACKEND",
        settings.internal_api_server_url,
    )
    backend_url = onlyoffice_internal_backend.rstrip('/')
    
    file_url = f"{backend_url}/api/onlyoffice/download/{doc.id}"
    callback_url = f"{backend_url}/api/onlyoffice/office-callback/{doc.id}"
    
    from hashlib import md5
    key = md5(f"{doc.id}-{doc.updated_at.isoformat()}".encode()).hexdigest()
    
    editor_mode = "edit" if can_edit else "view"
    
    config = {
        "document": {
            "fileType": doc.extension,
            "key": key,
            "title": doc.filename,
            "url": file_url,
            "permissions": {
                "chat": False,
                "comment": can_edit,
                "edit": can_edit,
            }
        },
        "documentType": get_document_type(doc.extension),
        "editorConfig": {
            "callbackUrl": callback_url,
            "lang": "es",
            "mode": editor_mode,
            "user": {
                "id": current_account_id,
                "name": "Usuario"
            },
            "customization": {
                "forcesave": True,
                "comments": can_edit,
                "help": False,
                "editRights": can_edit,
            }
        }
    }

    jwt_secret = os.getenv("ONLYOFFICE_JWT_SECRET")
    jwt_enabled = os.getenv("ONLYOFFICE_JWT_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
    if jwt_secret and jwt_enabled:
        token = jwt.encode(config, jwt_secret, algorithm="HS256")
        config["token"] = token if isinstance(token, str) else token.decode("utf-8")

    logger.info(
        "OnlyOffice config generated for %s using backend %s (jwt=%s, can_edit=%s)",
        document_id,
        backend_url,
        "enabled" if jwt_secret and jwt_enabled else "disabled",
        can_edit,
    )

    return {"config": config, "onlyoffice_url": onlyoffice_api_url}


@router.get("/share/{token}/config")
async def get_onlyoffice_public_config(
    token: str,
    db: AsyncSession = Depends(get_db_session)
):
    """Obtiene la configuración de OnlyOffice para enlaces publicos de documentos compartidos."""
    share_stmt = select(OnlyOfficeDocumentShare).where(
        OnlyOfficeDocumentShare.token == token,
        OnlyOfficeDocumentShare.is_public == True,
    )
    share_entry = (await db.execute(share_stmt)).scalars().first()
    if not share_entry:
        raise HTTPException(status_code=404, detail="Enlace de comparticion no encontrado")

    doc = await db.get(Document, share_entry.document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    onlyoffice_api_url = os.getenv("ONLYOFFICE_URL", "/onlyoffice")
    onlyoffice_internal_backend = os.getenv(
        "ONLYOFFICE_INTERNAL_BACKEND",
        settings.internal_api_server_url,
    )
    backend_url = onlyoffice_internal_backend.rstrip('/')

    file_url = f"{backend_url}/api/onlyoffice/download/{doc.id}"
    callback_url = f"{backend_url}/api/onlyoffice/office-callback/{doc.id}"

    from hashlib import md5
    key = md5(f"{doc.id}-{doc.updated_at.isoformat()}".encode()).hexdigest()

    editor_mode = "edit" if share_entry.can_edit else "view"

    config = {
        "document": {
            "fileType": doc.extension,
            "key": key,
            "title": doc.filename,
            "url": file_url,
            "permissions": {
                "chat": False,
                "comment": share_entry.can_edit,
                "edit": share_entry.can_edit,
            }
        },
        "documentType": get_document_type(doc.extension),
        "editorConfig": {
            "callbackUrl": callback_url,
            "lang": "es",
            "mode": editor_mode,
            "user": {
                "id": "public-share",
                "name": "Invitado"
            },
            "customization": {
                "forcesave": True,
                "comments": share_entry.can_edit,
                "help": False,
                "editRights": share_entry.can_edit,
            }
        }
    }

    jwt_secret = os.getenv("ONLYOFFICE_JWT_SECRET")
    jwt_enabled = os.getenv("ONLYOFFICE_JWT_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
    if jwt_secret and jwt_enabled:
        token_value = jwt.encode(config, jwt_secret, algorithm="HS256")
        config["token"] = token_value if isinstance(token_value, str) else token_value.decode("utf-8")

    return {"config": config, "onlyoffice_url": onlyoffice_api_url}


@router.post("/{document_id}/share-links")
async def create_document_share_link(
    document_id: uuid.UUID,
    request: CreateDocumentShareRequest,
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session),
):
    """Crea un enlace de comparticion publico o una comparticion privada hacia otra cuenta."""
    doc = await db.get(Document, document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    if str(doc.account_id) != current_account_id:
        raise HTTPException(status_code=403, detail="Solo el propietario puede compartir este documento")

    scope = (request.scope or "").strip().lower()
    if scope not in {"public", "private"}:
        raise HTTPException(status_code=400, detail="scope debe ser 'public' o 'private'")

    if scope == "public":
        token = secrets.token_urlsafe(32)
        share_entry = OnlyOfficeDocumentShare(
            document_id=doc.id,
            owner_account_id=doc.account_id,
            is_public=True,
            can_edit=request.can_edit,
            token=token,
        )
        db.add(share_entry)
        await db.commit()
        await db.refresh(share_entry)
        return {
            "id": str(share_entry.id),
            "scope": "public",
            "token": share_entry.token,
            "share_url": f"/share/document/{share_entry.token}",
            "can_edit": share_entry.can_edit,
            "created_at": share_entry.created_at.isoformat() if share_entry.created_at else None,
        }

    if not request.target_account_id:
        raise HTTPException(status_code=400, detail="target_account_id es obligatorio para comparticion privada")

    try:
        target_uuid = uuid.UUID(request.target_account_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="target_account_id invalido")

    if str(target_uuid) == current_account_id:
        raise HTTPException(status_code=400, detail="No puedes compartir el documento contigo mismo")

    target_account = await db.get(Account, target_uuid)
    if not target_account:
        raise HTTPException(status_code=404, detail="Cuenta destino no encontrada")

    existing_stmt = select(OnlyOfficeDocumentShare).where(
        OnlyOfficeDocumentShare.document_id == doc.id,
        OnlyOfficeDocumentShare.shared_with_account_id == target_uuid,
    )
    existing = (await db.execute(existing_stmt)).scalars().first()
    if existing:
        return {
            "id": str(existing.id),
            "scope": "private",
            "target_account_id": str(target_uuid),
            "can_edit": existing.can_edit,
            "created_at": existing.created_at.isoformat() if existing.created_at else None,
        }

    share_entry = OnlyOfficeDocumentShare(
        document_id=doc.id,
        owner_account_id=doc.account_id,
        shared_with_account_id=target_uuid,
        is_public=False,
        can_edit=request.can_edit,
    )
    db.add(share_entry)
    await db.commit()
    await db.refresh(share_entry)

    return {
        "id": str(share_entry.id),
        "scope": "private",
        "target_account_id": str(target_uuid),
        "can_edit": share_entry.can_edit,
        "created_at": share_entry.created_at.isoformat() if share_entry.created_at else None,
    }


@router.get("/{document_id}/share-links")
async def list_document_share_links(
    document_id: uuid.UUID,
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session),
):
    """Lista los enlaces y permisos de comparticion de un documento."""
    doc = await db.get(Document, document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    if str(doc.account_id) != current_account_id:
        raise HTTPException(status_code=403, detail="Solo el propietario puede ver los enlaces de comparticion")

    stmt = select(OnlyOfficeDocumentShare).where(
        OnlyOfficeDocumentShare.document_id == doc.id
    ).order_by(OnlyOfficeDocumentShare.created_at.desc())
    entries = (await db.execute(stmt)).scalars().all()

    account_ids = [e.shared_with_account_id for e in entries if e.shared_with_account_id]
    account_map = {}
    if account_ids:
        account_rows = (await db.execute(select(Account).where(Account.id.in_(account_ids)))).scalars().all()
        account_map = {acc.id: acc for acc in account_rows}

    response = []
    for entry in entries:
        shared_user = account_map.get(entry.shared_with_account_id)
        response.append({
            "id": str(entry.id),
            "scope": "public" if entry.is_public else "private",
            "can_edit": entry.can_edit,
            "token": entry.token,
            "share_url": f"/share/document/{entry.token}" if entry.is_public and entry.token else None,
            "target_account_id": str(entry.shared_with_account_id) if entry.shared_with_account_id else None,
            "target_email": shared_user.email if shared_user else None,
            "target_username": shared_user.username if shared_user else None,
            "target_name": shared_user.name if shared_user else None,
            "created_at": entry.created_at.isoformat() if entry.created_at else None,
        })

    return response


@router.delete("/share-links/{share_id}", response_model=DeleteDocumentShareResponse)
async def delete_document_share_link(
    share_id: uuid.UUID,
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session),
):
    """Elimina un permiso o enlace de comparticion de documento."""
    share_entry = await db.get(OnlyOfficeDocumentShare, share_id)
    if not share_entry:
        raise HTTPException(status_code=404, detail="Comparticion no encontrada")

    doc = await db.get(Document, share_entry.document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    if str(doc.account_id) != current_account_id:
        raise HTTPException(status_code=403, detail="Solo el propietario puede eliminar comparticiones")

    await db.delete(share_entry)
    await db.commit()
    return {"message": "Comparticion eliminada"}


@router.get("/{document_id}/chat-link")
async def get_or_create_document_chat_link(
    document_id: uuid.UUID,
    force_new: bool = Query(False, description="Si es True, crea un nuevo hilo de chat ignorando el anterior."),
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session),
):
    """Obtiene o crea el chat colaborativo persistente asociado a un documento."""
    doc = await db.get(Document, document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    if not await _can_access_document(doc, current_account_id, db):
        raise HTTPException(status_code=403, detail="No tienes acceso a este documento")

    effective_workspace_id = doc.workspace_id
    if effective_workspace_id is None and doc.folder_id is not None:
        folder = await db.get(DocumentFolder, doc.folder_id)
        if folder:
            effective_workspace_id = folder.workspace_id
            doc.workspace_id = effective_workspace_id
            await db.commit()

    mapping_stmt = select(OnlyOfficeDocumentChat).where(OnlyOfficeDocumentChat.document_id == doc.id)
    mapping = (await db.execute(mapping_stmt)).scalars().first()

    if not mapping or force_new:
        title_base = doc.filename.rsplit('.', 1)[0] if doc.filename else "Documento"
        thread = ChatThread(
            account_id=doc.account_id,
            workspace_id=effective_workspace_id,
            title=f"Chat documento: {title_base} (Nueva sesión)" if force_new else f"Chat documento: {title_base}",
            platform="web",
            created_at=datetime.now(),
            persistent_rag_context=[{
                "id": str(doc.id),
                "type": "document",
                "name": f"[OnlyOffice] {doc.filename}",
                "title": doc.filename,
                "topic": "OnlyOffice",
            }],
        )
        db.add(thread)
        await db.flush()

        if mapping:
            mapping.thread_id = thread.id
            logger.info(f"Chat reiniciado para documento {document_id}. Nuevo thread: {thread.id}")
        else:
            mapping = OnlyOfficeDocumentChat(document_id=doc.id, thread_id=thread.id)
            db.add(mapping)
        
        await db.commit()
        await db.refresh(mapping)
    else:
        # Verify that the existing thread's workspace_id matches the document's workspace_id
        thread = await db.get(ChatThread, mapping.thread_id)
        if thread and thread.workspace_id != effective_workspace_id:
            thread.workspace_id = effective_workspace_id
            await db.commit()

    share_stmt = select(SharedConversationLink).where(
        SharedConversationLink.thread_id == mapping.thread_id,
        SharedConversationLink.allow_reply == True,
        SharedConversationLink.password_hash == None,
        SharedConversationLink.expiry_date == None,
    ).order_by(SharedConversationLink.created_at.desc())
    share_link = (await db.execute(share_stmt)).scalars().first()

    logger.info(f"Chat link requested for document {document_id} (force_new={force_new})")

    if not share_link:
        share_link = SharedConversationLink(
            thread_id=mapping.thread_id,
            token=secrets.token_urlsafe(32),
            allow_reply=True,
        )
        db.add(share_link)
        await db.commit()
        await db.refresh(share_link)
        logger.info(f"New share link created for thread {mapping.thread_id}: {share_link.token}")
    else:
        logger.info(f"Existing share link found for thread {mapping.thread_id}: {share_link.token}")

    return {
        "thread_id": str(mapping.thread_id),
        "token": share_link.token,
        "share_url": f"/share/chat/{share_link.token}",
    }

def get_document_type(ext: str) -> str:
    """Mapea extensiones a tipos de documentos de OnlyOffice."""
    ext = ext.lower()
    if ext in ['doc', 'docx', 'rtf', 'txt', 'odt', 'md']: return 'word'
    if ext in ['xls', 'xlsx', 'csv', 'ods']: return 'cell'
    if ext in ['ppt', 'pptx', 'odp']: return 'slide'
    return 'word'

@router.get("/download/{document_id}")
async def download_document(
    document_id: uuid.UUID,
    request: Request,
    inline: bool = Query(False, description="Servir el archivo inline para previsualización"),
    db = Depends(get_db_session)
):
    """Permite a OnlyOffice descargar el archivo actual."""
    # ── Validación de Acceso (JWT) ──────────────────────────────────────────
    jwt_secret = os.getenv("ONLYOFFICE_JWT_SECRET")
    jwt_enabled = os.getenv("ONLYOFFICE_JWT_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
    
    token = request.query_params.get("token")
    if not token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            token = auth_header.replace("Bearer ", "").strip()
            
    is_authenticated = False
    
    if jwt_enabled and jwt_secret and token:
        try:
            jwt.decode(token, jwt_secret, algorithms=["HS256"])
            is_authenticated = True
            logger.debug(f"Acceso a descarga de documento {document_id} verificado mediante JWT de OnlyOffice.")
        except Exception:
            pass
            
    if not is_authenticated and token:
        try:
            user_id = decode_access_token(token)
            if user_id:
                is_authenticated = True
                logger.debug(f"Acceso a descarga de documento {document_id} verificado mediante JWT de usuario.")
        except Exception:
            pass
            
    if jwt_enabled and not is_authenticated:
        logger.warning(f"Acceso denegado a descarga de documento {document_id}: Autenticación fallida o token ausente.")
        raise HTTPException(
            status_code=403,
            detail="Acceso denegado. Se requiere un token JWT válido de OnlyOffice o de usuario."
        )

    doc = await db.get(Document, document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    
    try:
        file_full_path = resolve_onlyoffice_file_path(doc.file_path)
    except ValueError as exc:
        logger.error("Ruta física inválida para documento %s: %s", document_id, exc)
        raise HTTPException(status_code=500, detail="Ruta física del documento inválida") from exc

    if not file_full_path.exists():
        logger.error(
            "Archivo físico faltante para documento %s en %s",
            document_id,
            file_full_path,
        )
        raise HTTPException(status_code=404, detail="Archivo del documento no encontrado")

    if not file_full_path.is_file():
        logger.error(
            "La ruta del documento %s no es un archivo (es un directorio): %s",
            document_id,
            file_full_path,
        )
        raise HTTPException(status_code=500, detail="La ruta del documento no es un archivo válido")

    media_type, _ = mimetypes.guess_type(doc.filename or "")
    response_headers = None

    if inline:
        safe_filename = (doc.filename or "documento").replace('"', "")
        response_headers = {"Content-Disposition": f'inline; filename="{safe_filename}"'}

    return FileResponse(
        str(file_full_path),
        filename=None if inline else doc.filename,
        media_type=media_type or "application/octet-stream",
        headers=response_headers,
    )



@router.post("/downloadfile/{document_id}")
async def download_file_for_onlyoffice(
    document_id: uuid.UUID,
    request: Request,
    db = Depends(get_db_session),
):
    """Endpoint para OnlyOffice Document Server (POST /downloadfile/{id}).

    El DS llama a este endpoint para obtener el archivo durante la edición.
    Usa la misma lógica de autenticación y resolución de ruta que /download/.
    """
    jwt_secret = os.getenv("ONLYOFFICE_JWT_SECRET")
    jwt_enabled = os.getenv("ONLYOFFICE_JWT_ENABLED", "true").lower() in {"1", "true", "yes", "on"}

    token = request.query_params.get("token")
    if not token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            token = auth_header.replace("Bearer ", "").strip()

    is_authenticated = False

    if jwt_enabled and jwt_secret and token:
        try:
            jwt.decode(token, jwt_secret, algorithms=["HS256"])
            is_authenticated = True
        except Exception:
            pass

    if not is_authenticated and token:
        try:
            user_id = decode_access_token(token)
            if user_id:
                is_authenticated = True
        except Exception:
            pass

    if jwt_enabled and not is_authenticated:
        logger.warning(
            "Acceso denegado a downloadfile de documento %s: Autenticación fallida.",
            document_id,
        )
        raise HTTPException(
            status_code=403,
            detail="Acceso denegado. Se requiere un token JWT válido.",
        )

    doc = await db.get(Document, document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    try:
        file_full_path = resolve_onlyoffice_file_path(doc.file_path)
    except ValueError as exc:
        logger.error("Ruta física inválida para documento %s: %s", document_id, exc)
        raise HTTPException(status_code=500, detail="Ruta física del documento inválida") from exc

    if not file_full_path.exists():
        raise HTTPException(status_code=404, detail="Archivo del documento no encontrado")

    if not file_full_path.is_file():
        logger.error(
            "La ruta del documento %s no es un archivo (es un directorio): %s",
            document_id,
            file_full_path,
        )
        raise HTTPException(status_code=500, detail="La ruta del documento no es un archivo válido")

    media_type, _ = mimetypes.guess_type(doc.filename or "")

    return FileResponse(
        str(file_full_path),
        filename=doc.filename,
        media_type=media_type or "application/octet-stream",
    )


@router.post("/office-callback/{document_id}")
async def onlyoffice_callback(document_id: uuid.UUID, request: Request):
    """Recibe las actualizaciones de OnlyOffice para guardar el archivo."""
    try:
        body = await request.json()
    except Exception as e:
        logger.error(f"Error leyendo JSON del callback: {e}")
        return {"error": 1}
        
    jwt_secret = os.getenv("ONLYOFFICE_JWT_SECRET")
    jwt_enabled = os.getenv("ONLYOFFICE_JWT_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
    
    if jwt_enabled and jwt_secret:
        token = body.get("token")
        if not token:
            auth_header = request.headers.get("Authorization")
            if auth_header and auth_header.startswith("Bearer "):
                token = auth_header.replace("Bearer ", "").strip()
        
        if not token:
            logger.error("Token JWT requerido pero no encontrado en el callback de OnlyOffice.")
            return {"error": 1}
            
        import jwt
        try:
            decoded = jwt.decode(token, jwt_secret, algorithms=["HS256"])
            if isinstance(decoded, dict) and "payload" in decoded:
                body = decoded["payload"]
            else:
                body = decoded
        except Exception as e:
            logger.error(f"Error decodificando JWT en callback: {e}")
            return {"error": 1}

    status = body.get("status")
    logger.info(f"OnlyOffice Callback para {document_id}: Status {status}")
    
    try:
        if status in [2, 6]:
            download_url = body.get("url")
            if not download_url:
                logger.error("No se proporcionó download_url en el callback.")
                return {"error": 1}
                
            logger.info(f"Descargando archivo desde: {download_url}")
            async with httpx.AsyncClient(verify=False, timeout=300.0) as client:
                resp = await client.get(download_url)
                if resp.status_code == 200:
                    async with SessionLocal() as db:
                        doc = await db.get(Document, document_id)
                        if doc:
                            file_full_path = resolve_onlyoffice_file_path(doc.file_path)
                            
                            if file_full_path.exists():
                                backups_dir = file_full_path.parent / '.backups'
                                os.makedirs(backups_dir, exist_ok=True)
                                
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                backup_filename = f"{file_full_path.name}.{timestamp}.bak"
                                backup_full_path = backups_dir / backup_filename
                                
                                import shutil
                                import time
                                shutil.copy2(file_full_path, backup_full_path)
                                
                                try:
                                    for backup_file in os.listdir(backups_dir):
                                        bk_path = backups_dir / backup_file
                                        if bk_path.is_file() and time.time() - bk_path.stat().st_mtime > 2592000:
                                            os.remove(bk_path)
                                except Exception as e:
                                    logger.error(f"Error limpiando backups: {e}")
                            
                            if len(resp.content) > 100:
                                with open(file_full_path, "wb") as f:
                                    f.write(resp.content)
                                
                                doc.updated_at = datetime.now()
                                await db.commit()
                                logger.info(f"✅ Documento guardado (size: {len(resp.content)} bytes). Backup creado.")
                            else:
                                logger.error(f"❌ PELIGRO: OnlyOffice devolvió un archivo vacío. Ignorando.")
                        else:
                            logger.error(f"Documento {document_id} no encontrado en DB durante callback.")
                else:
                    logger.error(f"Error al descargar actualización de OnlyOffice: {resp.status_code} - {resp.text}")
                    return {"error": 1}
    except Exception as e:
        logger.error(f"Excepción en el callback de OnlyOffice: {e}")
        return {"error": 1}
                        
    return {"error": 0}

@router.post("/{document_id}/duplicate")
async def duplicate_document(
    document_id: uuid.UUID,
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Duplica un documento existente creando una copia exacta"""
    original_doc = await db.get(Document, document_id)
    if not original_doc or str(original_doc.account_id) != current_account_id:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    
    # Copiar archivo fisico
    original_path = resolve_onlyoffice_file_path(original_doc.file_path)
    if not original_path.exists():
        raise HTTPException(status_code=404, detail="Archivo fisico no encontrado")
        
    new_uuid = uuid.uuid4()
    new_file_path = build_onlyoffice_relative_path(current_account_id, f"{new_uuid}.{original_doc.extension}")
    new_full_path = resolve_onlyoffice_file_path(new_file_path)
    
    import shutil
    shutil.copy2(original_path, new_full_path)
    
    # Crear nuevo registro en DB
    duplicated_doc = Document(
        account_id=uuid.UUID(current_account_id),
        workspace_id=original_doc.workspace_id,
        folder_id=original_doc.folder_id,
        filename=f"Copia de {original_doc.filename}",
        extension=original_doc.extension,
        file_path=new_file_path
    )
    
    db.add(duplicated_doc)
    await db.commit()
    await db.refresh(duplicated_doc)
    
    return {"id": str(duplicated_doc.id), "filename": duplicated_doc.filename, "message": "Documento duplicado correctamente"}


@router.get("/history/{document_id}")
async def get_document_history(
    document_id: uuid.UUID,
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Obtiene el historial de backups automáticos de un documento."""
    doc = await db.get(Document, document_id)
    if not doc or str(doc.account_id) != current_account_id:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
        
    file_full_path = resolve_onlyoffice_file_path(doc.file_path)
    backups_dir = file_full_path.parent / '.backups'
    
    if not backups_dir.exists():
        return []
        
    backups = []
    base_name = file_full_path.name
    for f in os.listdir(backups_dir):
        if f.startswith(base_name) and f.endswith(".bak"):
            bk_path = backups_dir / f
            # Extraer la fecha del nombre del archivo (ej: nombre.ext.20231015_120000.bak)
            timestamp_str = f.replace(base_name + ".", "").replace(".bak", "")
            try:
                date_obj = datetime.strptime(timestamp_str, "%Y%m%d_%H%M%S")
                backups.append({
                    "filename": f,
                    "date": date_obj.isoformat(),
                    "size": bk_path.stat().st_size
                })
            except:
                pass
                
    # Ordenar por fecha descendente
    backups.sort(key=lambda x: x["date"], reverse=True)
    return backups

@router.post("/{document_id}/restore/{backup_filename}")
async def restore_document_backup(
    document_id: uuid.UUID,
    backup_filename: str,
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Restaura un documento desde un backup específico."""
    doc = await db.get(Document, document_id)
    if not doc or str(doc.account_id) != current_account_id:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
        
    file_full_path = resolve_onlyoffice_file_path(doc.file_path)
    backups_dir = file_full_path.parent / '.backups'
    backup_path = backups_dir / backup_filename
    
    if not backup_path.exists():
        raise HTTPException(status_code=404, detail="Backup no encontrado")
        
    import shutil
    # Crear un último backup del estado actual antes de restaurar
    if file_full_path.exists():
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pre_restore_backup = backups_dir / f"{file_full_path.name}.{timestamp}.bak"
        shutil.copy2(file_full_path, pre_restore_backup)
        
    shutil.copy2(backup_path, file_full_path)
    
    doc.updated_at = datetime.now()
    await db.commit()
    
    return {"message": "Documento restaurado con éxito"}

@router.delete("/{document_id}")
async def delete_document(
    document_id: uuid.UUID,
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Elimina un documento y su archivo físico."""
    doc = await db.get(Document, document_id)
    if not doc or str(doc.account_id) != current_account_id:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    
    file_full_path = resolve_onlyoffice_file_path(doc.file_path)
    if file_full_path.exists():
        try:
            os.remove(file_full_path)
        except Exception as e:
            logger.error(f"Error al eliminar archivo físico: {e}")
        
    # Eliminar chunks de pgvector
    from core.memory_manager import delete_document_chunks
    try:
        await delete_document_chunks(
            account_id=current_account_id,
            file_name=doc.filename,
            workspace_id=str(doc.workspace_id) if doc.workspace_id else None
        )
    except Exception as e:
        logger.error(f"Error al eliminar chunks RAG del documento {doc.filename}: {e}")

    await db.delete(doc)
    await db.commit()
    return {"message": "Documento eliminado con éxito."}


@router.post("/{document_id}/vectorize")
async def vectorize_document(
    document_id: uuid.UUID,
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Vectoriza e indexa un documento de OnlyOffice en el sistema RAG."""
    doc = await db.get(Document, document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
        
    if not await _can_access_document(doc, current_account_id, db):
        raise HTTPException(status_code=403, detail="No tienes acceso a este documento")

    try:
        file_full_path = resolve_onlyoffice_file_path(doc.file_path)
    except ValueError as exc:
        logger.error(f"Ruta física inválida para documento {document_id}: {exc}")
        raise HTTPException(status_code=500, detail="Ruta física del documento inválida")
        
    if not file_full_path.exists():
        raise HTTPException(status_code=404, detail="Archivo físico no encontrado")

    try:
        with open(file_full_path, "rb") as f:
            file_content = f.read()
    except Exception as e:
        logger.error(f"Error al leer el archivo para vectorización: {e}")
        raise HTTPException(status_code=500, detail="Error al leer el archivo en el servidor")

    from utils.document_parser import extract_text_and_metadata_from_document
    from core.memory_manager import delete_document_chunks, process_document_for_rag

    try:
        extracted_text, parser_metadata = await extract_text_and_metadata_from_document(doc.filename, file_content)
    except Exception as e:
        logger.error(f"Error al extraer texto del documento: {e}")
        raise HTTPException(status_code=500, detail="Error al extraer texto del documento")

    if not extracted_text or not extracted_text.strip():
        raise HTTPException(
            status_code=400,
            detail="El documento no contiene texto extraíble. Verifique que no esté vacío o que el formato sea soportado (ej. DOCX)."
        )

    # Determinar el topic (colección)
    topic = "general_documents"
    if doc.folder_id:
        folder = await db.get(DocumentFolder, doc.folder_id)
        if folder and folder.name:
            topic = folder.name

    # Eliminar chunks previos del mismo documento para evitar duplicados
    await delete_document_chunks(
        account_id=current_account_id,
        file_name=doc.filename,
        workspace_id=str(doc.workspace_id) if doc.workspace_id else None
    )

    # Preparar metadatos para el RAG
    metadata = {
        "document_id": str(doc.id),
        "workspace_id": str(doc.workspace_id) if doc.workspace_id else None,
        "folder_id": str(doc.folder_id) if doc.folder_id else None,
        "extension": doc.extension,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
        "file_name": doc.filename,
    }
    if parser_metadata:
        metadata.update(parser_metadata)

    # Llamar a la función de procesamiento RAG
    chunks_processed = await process_document_for_rag(
        file_name=doc.filename,
        extracted_text=extracted_text,
        topic=topic,
        account_id=current_account_id,
        metadata=metadata,
        workspace_id=str(doc.workspace_id) if doc.workspace_id else None
    )

    return {
        "message": "Documento vectorizado e indexado correctamente",
        "chunks_processed": chunks_processed,
        "topic": topic
    }


@router.put("/{document_id}")
async def update_document_content(
    document_id: uuid.UUID,
    file: UploadFile = File(...),
    current_account_id: str = Depends(get_current_account_id),
    db: AsyncSession = Depends(get_db_session)
):
    """Sobrescribe el contenido de un documento existente y actualiza RAG."""
    doc = await db.get(Document, document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
        
    # Verificar acceso usando la misma lógica que otros endpoints
    if not await _can_access_document(doc, current_account_id, db):
        raise HTTPException(status_code=403, detail="No tienes acceso a este documento")

    try:
        file_full_path = resolve_onlyoffice_file_path(doc.file_path)
    except ValueError as exc:
        raise HTTPException(status_code=500, detail="Ruta física del documento inválida")
        
    content = await file.read()
    
    # Crear backup del archivo actual
    if file_full_path.exists():
        backups_dir = file_full_path.parent / '.backups'
        os.makedirs(backups_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_filename = f"{file_full_path.name}.{timestamp}.bak"
        backup_full_path = backups_dir / backup_filename
        import shutil
        try:
            shutil.copy2(file_full_path, backup_full_path)
            logger.info(f"Backup creado exitosamente para {doc.filename} en {backup_full_path}")
        except Exception as e:
            logger.error(f"Error al crear backup para {doc.filename}: {e}")

    # Escribir el nuevo archivo físico
    try:
        with open(file_full_path, "wb") as f:
            f.write(content)
        logger.info(f"Archivo físico {doc.filename} sobrescrito.")
    except Exception as e:
        logger.error(f"Error al escribir archivo {doc.filename}: {e}")
        raise HTTPException(status_code=500, detail="Error al sobrescribir el archivo en el servidor")

    # Actualizar la fecha en DB
    doc.updated_at = datetime.now()
    await db.commit()

    # Re-vectorizar RAG
    try:
        extracted_text, parser_metadata = await extract_text_and_metadata_from_document(doc.filename, content)
    except Exception as e:
        logger.error(f"Error al extraer texto en actualización de {doc.filename}: {e}")
        raise HTTPException(status_code=500, detail="Error al extraer texto del documento")

    if not extracted_text or not extracted_text.strip():
        raise HTTPException(
            status_code=400,
            detail="El documento actualizado no contiene texto extraíble."
        )

    # Eliminar chunks antiguos
    await delete_document_chunks(
        account_id=current_account_id,
        file_name=doc.filename,
        workspace_id=str(doc.workspace_id) if doc.workspace_id else None
    )

    # Re-indexar
    metadata = {
        "document_id": str(doc.id),
        "workspace_id": str(doc.workspace_id) if doc.workspace_id else None,
        "folder_id": str(doc.folder_id) if doc.folder_id else None,
        "extension": doc.extension,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
        "file_name": doc.filename,
    }
    if parser_metadata:
        metadata.update(parser_metadata)

    topic = "general_documents"
    if doc.folder_id:
        folder = await db.get(DocumentFolder, doc.folder_id)
        if folder and folder.name:
            topic = folder.name

    await process_document_for_rag(
        file_name=doc.filename,
        extracted_text=extracted_text,
        topic=topic,
        account_id=current_account_id,
        metadata=metadata,
        workspace_id=str(doc.workspace_id) if doc.workspace_id else None
    )

    return {
        "message": "Documento actualizado con éxito",
        "document_id": str(doc.id),
        "updated_at": doc.updated_at.isoformat()
    }
