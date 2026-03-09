# api/notes.py

import logging
from typing import List, Optional
import uuid
import re
import base64
from datetime import datetime # Importar datetime
import io # Para manejar el PDF en memoria
import markdown # Para convertir Markdown a HTML
from weasyprint import HTML, CSS # Para generar PDF desde HTML

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File
from fastapi.responses import StreamingResponse, JSONResponse # Para devolver el PDF
from pydantic import BaseModel
from sqlalchemy import select, desc
from sqlalchemy.orm import selectinload # Import selectinload
from sqlalchemy.ext.asyncio import AsyncSession

# Ya no necesitamos reportlab
# from reportlab.lib.pagesizes import letter
# from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
# from reportlab.lib.styles import getSampleStyleSheet
# from reportlab.lib.units import inch

from core.database import SessionLocal
from core.notes_manager import NotesManager
from utils.security import get_current_account_id, check_workspace_permission
import os
import shutil
from pathlib import Path
from api.contact_profiles import ContactProfileResponse # Importar ContactProfileResponse
from core.dependencies import get_db_session # Importar dependencia centralizada
import jwt
import time
from core.config import settings
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH



logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

router = APIRouter()

# --- Dependencias de FastAPI ---

# get_db eliminado en favor de core.dependencies.get_db_session

def get_notes_manager(db: AsyncSession = Depends(get_db_session)) -> NotesManager:
    """Inyecta una instancia del gestor de notas."""
    return NotesManager(db)



# --- Modelos Pydantic para la API ---

class ListNotesRequest(BaseModel):
    search_term: Optional[str] = None
    workspace_id: Optional[str] = None
    category: Optional[str] = None
    skip: int = 0
    limit: int = 10

class NoteResponse(BaseModel):
    id: int
    title: Optional[str] # Modificado para ser opcional
    content: str
    category: Optional[str]
    created_at: datetime
    updated_at: datetime
    workspace_id: Optional[str]
    workspace_name: Optional[str] # Added
    workspace_color: Optional[str] # Added
    linked_profiles: List[ContactProfileResponse] = []

class NoteRequest(BaseModel):
    title: Optional[str] = None
    content: str
    category: Optional[str] = None
    workspace_id: Optional[str] = None

class NoteUpdateRequest(BaseModel):
    note_id: int
    title: Optional[str] = None
    content: Optional[str] = None
    category: Optional[str] = None
    workspace_id: Optional[str] = None

class NoteDeleteRequest(BaseModel):
    note_id: int

class PaginatedNotesResponse(BaseModel):
    total: int
    notes: List[NoteResponse]

# Se mantiene ProfileLinkRequest por si se usa en otros endpoints o en el futuro.
# Si no se usa en ningún otro lugar, se podría eliminar.
class ProfileLinkRequest(BaseModel):
    profile_id: uuid.UUID

class GeneratePdfRequest(BaseModel):
    note_id: int
    format: str = "markdown" # Por ahora solo soportamos markdown

class LinkNoteToWorkspaceRequest(BaseModel):
    workspace_id: str # El ID del workspace al que se quiere vincular

# --- Endpoints de la API ---

@router.post("/notes/{note_id}/link-to-workspace", summary="Vincular una nota a un workspace")
async def link_note_to_workspace_endpoint(
    note_id: int,
    request: LinkNoteToWorkspaceRequest,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Vincula una nota a un workspace específico.
    """
    # Obtener la nota para verificar permisos
    note_data = await notes_manager.get_note_by_id(current_account_id, note_id)
    if not note_data:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada.")

    # Verificar permisos de workspace si la nota ya pertenece a uno o si se va a vincular a uno nuevo
    # El usuario debe tener al menos permiso de 'member' en el workspace de destino
    if not await check_workspace_permission(current_account_id, request.workspace_id, notes_manager.db, required_roles=['admin', 'owner', 'member']):
        raise HTTPException(status_code=403, detail="No tienes permiso para vincular notas a este workspace.")

    success = await notes_manager.update_note(
        account_id=current_account_id,
        note_id=note_id,
        new_workspace_id=request.workspace_id
    )
    if not success:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada para vincular.")
    return {"message": f"Nota con ID {note_id} vinculada al workspace {request.workspace_id} correctamente."}

@router.post("/notes/generate-pdf", summary="Generar PDF de una nota")
async def generate_note_pdf_endpoint(
    request: GeneratePdfRequest,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Genera un PDF a partir del contenido de una nota con diseño estilizado.
    """
    note_data = await notes_manager.get_note_by_id(current_account_id, request.note_id)
    if not note_data:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada.")

    # Verificar permisos de workspace si la nota pertenece a uno
    if note_data.get("workspace_id"):
        if not await check_workspace_permission(current_account_id, note_data["workspace_id"], notes_manager.db, required_roles=['admin', 'owner', 'member', 'viewer']):
            raise HTTPException(status_code=403, detail="No tienes permiso para generar PDF de esta nota.")

    buffer = io.BytesIO()

    # Process Mermaid diagrams
    content = note_data["content"]
    def replace_mermaid(match):
        mermaid_code = match.group(1).strip()
        mermaid_base64 = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
        image_url = f"https://mermaid.ink/img/{mermaid_base64}"
        return f'<div class="mermaid-diagram"><img src="{image_url}" alt="Mermaid Diagram" /></div>'

    pattern = r'```mermaid\s+(.*?)```'
    content = re.sub(pattern, replace_mermaid, content, flags=re.DOTALL)

    # Convertir Markdown a HTML, asegurando el soporte para tablas
    html_content = markdown.markdown(content, extensions=['tables'])

    # Convertir created_at a datetime si es una cadena y asegurar que siempre esté definida
    created_at_dt = note_data['created_at']
    if isinstance(created_at_dt, str):
        try:
            created_at_dt = datetime.fromisoformat(created_at_dt.replace('Z', '+00:00'))
        except ValueError:
            pass

    # Leer el logo y convertirlo a base64 para incluirlo en el HTML
    logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'public', 'logo-simple.png')
    try:
        with open(logo_path, 'rb') as logo_file:
            logo_base64 = base64.b64encode(logo_file.read()).decode('utf-8')
            logo_data_uri = f"data:image/png;base64,{logo_base64}"
    except Exception as e:
        logger.warning(f"No se pudo cargar el logo: {e}")
        logo_data_uri = ""

    # Crear un HTML completo con estilos mejorados y profesionales
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>{note_data.get("title", "Nota sin título")}</title>
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');
            
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            
            body {{
                font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
                margin: 0;
                padding: 0;
                font-size: 11pt;
                line-height: 1.6;
                color: #1a1a1a;
                position: relative;
                min-height: 100vh;
            }}
            
            .page-container {{
                margin: 0.5in 0.5in 1in 0.5in;
                padding-bottom: 60px;
            }}
            
            .header {{
                margin-bottom: 2em;
                padding-bottom: 1em;
                border-bottom: 3px solid #6366f1;
            }}
            
            h1 {{
                color: #1a1a1a;
                font-size: 24pt;
                font-weight: 700;
                margin-bottom: 0.5em;
                letter-spacing: -0.02em;
            }}
            
            .note-date {{
                font-size: 9pt;
                color: #6b7280;
                font-weight: 400;
                margin-top: 0.5em;
            }}
            
            .content {{
                margin-top: 2em;
            }}
            
            .content p {{
                margin-bottom: 1em;
                text-align: justify;
                line-height: 1.7;
            }}
            
            .content h2 {{
                color: #374151;
                font-size: 16pt;
                font-weight: 600;
                margin-top: 1.5em;
                margin-bottom: 0.75em;
                border-left: 4px solid #6366f1;
                padding-left: 0.5em;
            }}
            
            .content h3 {{
                color: #4b5563;
                font-size: 13pt;
                font-weight: 600;
                margin-top: 1.2em;
                margin-bottom: 0.6em;
            }}
            
            .content ul, .content ol {{
                margin-left: 1.5em;
                margin-bottom: 1em;
            }}
            
            .content li {{
                margin-bottom: 0.4em;
            }}
            
            pre {{
                background-color: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                padding: 1em;
                overflow-x: auto;
                margin-bottom: 1em;
                font-size: 9pt;
            }}
            
            code {{
                font-family: 'Courier New', monospace;
                background-color: #f3f4f6;
                padding: 0.2em 0.4em;
                border-radius: 3px;
                font-size: 9.5pt;
            }}
            
            pre code {{
                background-color: transparent;
                padding: 0;
            }}
            
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 1.5em;
                font-size: 10pt;
            }}
            
            th, td {{
                border: 1px solid #d1d5db;
                padding: 0.6em 0.8em;
                text-align: left;
            }}
            
            th {{
                background-color: #f3f4f6;
                font-weight: 600;
                color: #374151;
            }}
            
            tr:nth-child(even) {{
                background-color: #f9fafb;
            }}
            
            .footer {{
                position: fixed;
                bottom: 0;
                left: 0;
                right: 0;
                height: 60px;
                display: flex;
                flex-direction: row;
                align-items: center;
                justify-content: flex-start;
                gap: 10px;
                border-top: 1px solid #e5e7eb;
                background-color: #ffffff;
                padding: 15px 0.5in;
            }}
            
            .footer-logo {{
                width: 28px;
                height: 28px;
                opacity: 0.9;
            }}
            
            .footer-text {{
                font-size: 10pt;
                color: #6b7280;
                font-weight: 500;
                letter-spacing: 0.02em;
            }}
            
            blockquote {{
                border-left: 4px solid #e5e7eb;
                padding-left: 1em;
                margin-left: 0;
                margin-bottom: 1em;
                color: #6b7280;
                font-style: italic;
            }}
        </style>
    </head>
    <body>
        <div class="page-container">
            <div class="header">
                <h1>{note_data.get("title", "Nota sin título")}</h1>
                <div class="note-date">
                    Creada el: {created_at_dt.strftime('%d de %B de %Y, %H:%M') if isinstance(created_at_dt, datetime) else str(created_at_dt)}
                </div>
            </div>
            
            <div class="content">
                {html_content}
            </div>
        </div>
        
        <div class="footer">
            {f'<img src="{logo_data_uri}" class="footer-logo" alt="Kognito AI Logo" />' if logo_data_uri else ''}
            <span class="footer-text">Kognito AI</span>
        </div>
    </body>
    </html>
    """

    # Generar PDF usando WeasyPrint
    HTML(string=full_html).write_pdf(buffer)
    buffer.seek(0)

    return StreamingResponse(buffer, media_type="application/pdf", headers={
        "Content-Disposition": f"attachment; filename=\"{note_data.get('title', 'nota')}.pdf\""
    })

@router.get("/notes/{note_id}", response_model=NoteResponse, summary="Obtener una nota por ID")
async def get_note_by_id_endpoint(
    note_id: int,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Obtiene una nota específica por su ID, incluyendo los perfiles de contacto vinculados.
    """
    note = await notes_manager.get_note_by_id(current_account_id, note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada.")
    
    # Si la nota tiene un workspace_id, verificar permisos
    if note.get("workspace_id"):
        if not await check_workspace_permission(current_account_id, note["workspace_id"], notes_manager.db, required_roles=['admin', 'owner', 'member', 'viewer']):
            raise HTTPException(status_code=403, detail="No tienes permiso para acceder a esta nota.")
            
    return NoteResponse(**note)

@router.get("/notes/{note_id}/debug", summary="DEBUG: Obtener contenido crudo de una nota por ID")
async def debug_get_note_content_by_id_endpoint(
    note_id: int,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    DEBUG: Obtiene el contenido crudo de una nota específica por su ID, sin serialización Pydantic.
    Útil para verificar el contenido directamente de la base de datos.
    """
    note_data = await notes_manager.get_note_by_id(current_account_id, note_id)
    if not note_data:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada.")
    return {"id": note_data["id"], "title": note_data["title"], "content": note_data["content"], "category": note_data["category"], "workspace_id": note_data["workspace_id"]}

@router.get("/notes/{note_id}/linked-profiles", response_model=List[ContactProfileResponse], summary="Obtener perfiles vinculados a una nota")
async def get_linked_profiles_for_note_endpoint(
    note_id: int,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Obtiene la lista de perfiles de contacto vinculados a una nota específica.
    """
    note = await notes_manager.get_note_by_id(current_account_id, note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada.")
    
    # Si la nota tiene un workspace_id, verificar permisos
    if note.get("workspace_id"):
        if not await check_workspace_permission(current_account_id, note["workspace_id"], notes_manager.db, required_roles=['admin', 'owner', 'member', 'viewer']):
            raise HTTPException(status_code=403, detail="No tienes permiso para acceder a esta nota.")
            
    return note["linked_profiles"]

@router.post("/notes/list-notes", response_model=PaginatedNotesResponse, summary="Listar notas del usuario con paginación")
async def list_notes_endpoint(
    request: ListNotesRequest,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Devuelve todas las notas de un usuario, incluyendo personales y de equipos, o filtradas por workspace, con paginación.
    """
    total, notes = await notes_manager.get_notes_as_dicts(
        current_account_id, 
        request.search_term, 
        workspace_id=request.workspace_id, 
        category=request.category,
        skip=request.skip, 
        limit=request.limit
    )
    
    return PaginatedNotesResponse(total=total, notes=[NoteResponse(**note) for note in notes])

@router.post("/notes/{note_id}/link-profile", summary="Vincular perfil a una nota") # CAMBIO EN LA RUTA
async def link_profile_to_note_endpoint(
    note_id: int,
    profile_link_request: ProfileLinkRequest, # CAMBIO: Ahora profile_id es un path parameter
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Vincula un perfil de contacto a una nota.
    """
    # Obtener la nota para verificar el workspace_id
    note_data = await notes_manager.get_note_by_id(current_account_id, note_id)
    if not note_data:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada.")

    # Verificar permisos de workspace si la nota pertenece a uno
    if note_data.get("workspace_id"):
                    if not await check_workspace_permission(current_account_id, note_data["workspace_id"], notes_manager.db, required_roles=['admin', 'owner', 'member']):
                        raise HTTPException(status_code=403, detail="No tienes permiso para vincular perfiles a esta nota.")
    success = await notes_manager.link_profile_to_note(
        account_id=current_account_id,
        note_id=note_id,
        profile_id=profile_link_request.profile_id # CAMBIO: Se pasa directamente el profile_id
    )
    if not success:
        raise HTTPException(status_code=404, detail="Nota o perfil no encontrado, o no autorizado.")
    return {"message": f"Perfil {profile_link_request.profile_id} vinculado a la nota {note_id} correctamente."}

@router.post("/notes/{note_id}/unlink-profile", summary="Desvincular perfil de una nota")
async def unlink_profile_from_note_endpoint(
    note_id: int,
    profile_link_request: ProfileLinkRequest,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Desvincula un perfil de contacto de una nota.
    """
    # Obtener la nota para verificar el workspace_id
    note_data = await notes_manager.get_note_by_id(current_account_id, note_id)
    if not note_data:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada.")

    # Verificar permisos de workspace si la nota pertenece a uno
    if note_data.get("workspace_id"):
        if not await check_workspace_permission(current_account_id, note_data["workspace_id"], notes_manager.db, required_roles=['admin', 'owner', 'member']):
            raise HTTPException(status_code=403, detail="No tienes permiso para desvincular perfiles de esta nota.")

    success = await notes_manager.unlink_profile_from_note(
        account_id=current_account_id,\
        note_id=note_id,\
        profile_id=profile_link_request.profile_id\
    )
    if not success:
        raise HTTPException(status_code=404, detail="Vínculo no encontrado, o nota/perfil no autorizado.")
    return {"message": f"Perfil {profile_link_request.profile_id} desvinculado de la nota {note_id} correctamente."}

@router.post("/notes/{note_id}/unshare", summary="Descompartir una nota de su workspace")
async def unshare_note_endpoint(
    note_id: int,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Desvincula una nota de su workspace, estableciendo su workspace_id a None.
    """
    success = await notes_manager.unshare_note_from_workspace(
        account_id=current_account_id,
        note_id=note_id
    )
    if not success:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada para descompartir.")
    return {"message": f"Nota con ID {note_id} descompartida de su workspace correctamente."}

@router.post("/add-note")
async def add_note_endpoint(
    note: NoteRequest,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """Añade una nueva nota para el usuario."""
    new_note = await notes_manager.add_note(
        account_id=current_account_id,\
        title=note.title or "",\
        content=note.content,\
        category=note.category or "",\
        workspace_id=note.workspace_id\
    )
    return new_note

@router.post("/update-note")
async def update_note_endpoint(
    note: NoteUpdateRequest,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """Actualiza una nota existente del usuario."""
    # Obtener la nota para verificar el workspace_id
    note_data = await notes_manager.get_note_by_id(current_account_id, note.note_id)
    if not note_data:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada.")

    # Verificar permisos de workspace si la nota pertenece a uno
    if note_data.get("workspace_id"):
        if not await check_workspace_permission(current_account_id, note_data["workspace_id"], notes_manager.db, required_roles=['admin', 'owner', 'member']):
            raise HTTPException(status_code=403, detail="No tienes permiso para actualizar esta nota.")

    success = await notes_manager.update_note(
        account_id=current_account_id,
        note_id=note.note_id,
        new_title=note.title,
        new_content=note.content,
        new_category=note.category,
        new_workspace_id=note.workspace_id
    )
    if not success:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no pertenece al usuario.")
    return {"message": f"Nota con ID {note.note_id} actualizada correctamente."}

@router.post("/delete-note")
async def delete_note_endpoint(
    note: NoteDeleteRequest,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """Elimina una nota del usuario."""
    # Obtener la nota para verificar el workspace_id
    note_data = await notes_manager.get_note_by_id(current_account_id, note.note_id)
    if not note_data:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no autorizada.")

    # Verificar permisos de workspace si la nota pertenece a uno
    if note_data.get("workspace_id"):
        if not await check_workspace_permission(current_account_id, note_data["workspace_id"], notes_manager.db, required_roles=['admin', 'owner', 'member']):
            raise HTTPException(status_code=403, detail="No tienes permiso para eliminar esta nota.")

    success = await notes_manager.delete_note(current_account_id, note.note_id)
    if not success:
        raise HTTPException(status_code=404, detail="Nota no encontrada o no pertenece al usuario.")
    return {"message": f"Nota con ID {note.note_id} eliminada."}


@router.post("/notes/delete-all-embeddings", summary="Delete all note embeddings")
async def delete_all_note_embeddings_endpoint(
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Deletes all embeddings for all notes of the current user.
    """
    deleted_count = await notes_manager.delete_all_note_embeddings(current_account_id)
    return {"message": f"Deleted {deleted_count} note embeddings."}


@router.post("/notes/revectorize-all", summary="Re-vectorize all notes")
async def revectorize_all_notes_endpoint(
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Re-vectorizes all notes for the current user.
    """
    revectorized_count = await notes_manager.revectorize_all_notes(current_account_id)
    return {"message": f"Re-vectorized {revectorized_count} notes."}


@router.post("/upload-image", summary="Subir imagen para notas")
async def upload_image(
    file: UploadFile = File(...),
    current_account_id: str = Depends(get_current_account_id)
):
    """
    Sube una imagen y devuelve la URL para usarla en las notas.
    """
    # Validar tipo de archivo
    allowed_types = ["image/jpeg", "image/png", "image/gif", "image/webp"]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Tipo de archivo no permitido. Solo imágenes JPEG, PNG, GIF y WebP.")

    # Validar tamaño del archivo (máximo 5MB)
    max_size = 5 * 1024 * 1024  # 5MB
    file_content = await file.read()
    if len(file_content) > max_size:
        raise HTTPException(status_code=400, detail="El archivo es demasiado grande. Máximo 5MB.")

    # Crear directorio si no existe
    upload_dir = Path("media/uploads/notes")
    upload_dir.mkdir(parents=True, exist_ok=True)

    # Generar nombre único para el archivo
    file_extension = Path(file.filename).suffix.lower()
    unique_filename = f"{current_account_id}_{uuid.uuid4()}{file_extension}"
    file_path = upload_dir / unique_filename

    # Guardar el archivo
    with open(file_path, "wb") as buffer:
        buffer.write(file_content)

    # Devolver la URL relativa
    image_url = f"/media/uploads/notes/{unique_filename}"
    return {"url": image_url}

@router.post("/notes/{note_id}/convert-to-word", summary="Convertir nota a DOCX")
async def convert_note_to_word(
    note_id: int,
    current_account_id: str = Depends(get_current_account_id),
    notes_manager: NotesManager = Depends(get_notes_manager)
):
    """
    Convierte el contenido de una nota (Markdown/HTML) a un archivo .docx real.
    """
    note = await notes_manager.get_note_by_id(current_account_id, note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Nota no encontrada.")

    # Convertir Markdown a HTML
    html_content = markdown.markdown(note.get("content", ""))
    soup = BeautifulSoup(html_content, "html.parser")

    doc = Document()
    doc.add_heading(note.get("title") or "Nota sin título", level=0)

    # Simple HTML to Docx mapping
    for element in soup.recursiveChildGenerator():
        if element.name:
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level=level)
            elif element.name == 'p':
                p = doc.add_paragraph(element.get_text())
                # Basic formatting check (bold, italic could be recursive but keeping it simple for now)
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(), style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(), style='List Number')
            elif element.name == 'pre':
                 doc.add_paragraph(element.get_text(), style='No Spacing') # Code block approximation

    # Save logic
    filename = f"note_{note_id}.docx"
    file_path = os.path.join("media/notes", filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)

    return {"message": "Nota convertida a Word correctamente", "file_path": file_path}


