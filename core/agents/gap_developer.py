# core/agents/gap_developer.py

import asyncio
import uuid
import json
import logging
from typing import Any, Dict, List, Optional, cast

from langchain_core.messages import (
    AIMessage,
    BaseMessage,
    HumanMessage,
    SystemMessage,
    ToolMessage,
)
from langchain_core.runnables import RunnableConfig
from langgraph.graph import END, START, StateGraph

from core.llm_manager import get_main_llm, get_fast_llm, get_llm_for_user # Importar get_llm_for_user
from core.utils.llm_utils import safe_bind_tools
from core.utils.date_utils import get_today_str
from core.skill_manager import get_skill_manager

logger = logging.getLogger(__name__)

# --- Imports and helpers for DOCX generation and OnlyOffice ---
import re
from core.database import SessionLocal, Document
from core.onlyoffice_storage import build_onlyoffice_relative_path, ensure_onlyoffice_account_dir

try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

inline_regex = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|__.*?__|__.*?__|\*.*?\*|_.*?_|`.*?`)')

def set_cell_background(cell, hex_color: str):
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)
    except Exception as e:
        logger.warning(f"No se pudo establecer el fondo de celda: {e}")

def set_table_borders(table, hex_color: str = "CBD5E1"):
    try:
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'  <w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)
    except Exception as e:
        logger.warning(f"No se pudieron aplicar bordes de tabla: {e}")

def _add_inline_formatted_text(paragraph, text: str, default_font: str = "Segoe UI", default_size = None, default_color = None):
    if default_size is None:
        default_size = Pt(10.5)
    if default_color is None:
        default_color = RGBColor(51, 65, 85)
    if not text:
        return
    tokens = inline_regex.split(text)
    for token in tokens:
        if not token:
            continue
        is_bold = False
        is_italic = False
        is_code = False
        clean_text = token
        if token.startswith("***") and token.endswith("***") and len(token) > 6:
            is_bold = True
            is_italic = True
            clean_text = token[3:-3]
        elif token.startswith("**") and token.endswith("**") and len(token) > 4:
            is_bold = True
            clean_text = token[2:-2]
        elif token.startswith("__") and token.endswith("__") and len(token) > 4:
            is_bold = True
            clean_text = token[2:-2]
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            is_italic = True
            clean_text = token[1:-1]
        elif token.startswith("_") and token.endswith("_") and len(token) > 2:
            is_italic = True
            clean_text = token[1:-1]
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            is_code = True
            clean_text = token[1:-1]
            
        run = paragraph.add_run(clean_text)
        if is_code:
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(239, 68, 68)
            try:
                rPr = run._r.get_or_add_rPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                rPr.append(shd)
            except Exception:
                pass
        else:
            run.font.name = default_font
            run.font.size = default_size
            run.font.color.rgb = default_color
            run.bold = is_bold
            run.italic = is_italic

def _parse_and_render_markdown(doc_obj, md_text: str):
    if not md_text:
        return
    lines = md_text.split('\n')
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < n:
                i += 1
            code_text = "\n".join(code_lines)
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.left_indent = Pt(12)
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(15, 23, 42)
            try:
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
                pPr.append(shd)
                pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="CBD5E1"/></w:pBdr>')
                pPr.append(pbdr)
            except Exception:
                pass
            continue
        h_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if h_match:
            level = len(h_match.group(1))
            content = h_match.group(2).strip()
            p = doc_obj.add_paragraph()
            p.paragraph_format.keep_with_next = True
            if level == 1:
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
                _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(18), default_color=RGBColor(30, 58, 138))
                for r in p.runs:
                    r.bold = True
            elif level == 2:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(14), default_color=RGBColor(37, 99, 235))
                for r in p.runs:
                    r.bold = True
            else:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(12), default_color=RGBColor(30, 41, 59))
                for r in p.runs:
                    r.bold = True
            i += 1
            continue
        if stripped in ("---", "***", "___") or re.match(r'^[-*_]{3,}$', stripped):
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            try:
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E2E8F0"/></w:pBdr>')
                pPr.append(pBdr)
            except Exception:
                pass
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for tl in table_lines:
                cols = [col.strip() for col in tl.split("|")[1:-1]]
                rows.append(cols)
            if len(rows) > 1 and all(re.match(r'^:?-+:?$', c) for c in rows[1]):
                rows.pop(1)
            if rows:
                num_rows = len(rows)
                num_cols = max(len(r) for r in rows)
                table = doc_obj.add_table(rows=num_rows, cols=num_cols)
                table.autofit = True
                set_table_borders(table)
                for row_idx, r_data in enumerate(rows):
                    for col_idx, cell_val in enumerate(r_data):
                        if col_idx < num_cols:
                            cell = table.cell(row_idx, col_idx)
                            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                            p.paragraph_format.space_before = Pt(4)
                            p.paragraph_format.space_after = Pt(4)
                            p.paragraph_format.line_spacing = 1.0
                            if row_idx == 0:
                                _add_inline_formatted_text(p, cell_val, default_font="Segoe UI", default_size=Pt(10), default_color=RGBColor(255, 255, 255))
                                for run in p.runs:
                                    run.bold = True
                                set_cell_background(cell, "1E3A8A")
                            else:
                                _add_inline_formatted_text(p, cell_val, default_font="Segoe UI", default_size=Pt(9.5), default_color=RGBColor(51, 65, 85))
                                if row_idx % 2 == 0:
                                    set_cell_background(cell, "F8FAFC")
                                else:
                                    set_cell_background(cell, "FFFFFF")
                p_after = doc_obj.add_paragraph()
                p_after.paragraph_format.space_before = Pt(0)
                p_after.paragraph_format.space_after = Pt(6)
            continue
        bullet_match = re.match(r'^([\s]*)([-*+])\s+(.*)$', line)
        num_match = re.match(r'^([\s]*)(\d+)\.\s+(.*)$', line)
        if bullet_match or num_match:
            list_lines = []
            while i < n:
                curr_line = lines[i]
                if not curr_line.strip():
                    break
                b_m = re.match(r'^([\s]*)([-*+])\s+(.*)$', curr_line)
                n_m = re.match(r'^([\s]*)(\d+)\.\s+(.*)$', curr_line)
                if b_m or n_m:
                    list_lines.append((curr_line, b_m, n_m))
                    i += 1
                else:
                    if list_lines and curr_line.startswith("   "):
                        prev_text, prev_bm, prev_nm = list_lines[-1]
                        updated_text = prev_text + " " + curr_line.strip()
                        b_m_up = re.match(r'^([\s]*)([-*+])\s+(.*)$', updated_text)
                        n_m_up = re.match(r'^([\s]*)(\d+)\.\s+(.*)$', updated_text)
                        list_lines[-1] = (updated_text, b_m_up, n_m_up)
                        i += 1
                    else:
                        break
            for item_line, bm, nm in list_lines:
                p = doc_obj.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                if bm:
                    indent_level = len(bm.group(1)) // 2
                    content = bm.group(3).strip()
                    p.style = 'List Bullet'
                    p.paragraph_format.left_indent = Pt(18 * (indent_level + 1))
                    _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(10.5), default_color=RGBColor(51, 65, 85))
                else:
                    indent_level = len(nm.group(1)) // 2
                    content = nm.group(3).strip()
                    p.style = 'List Number'
                    p.paragraph_format.left_indent = Pt(18 * (indent_level + 1))
                    _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(10.5), default_color=RGBColor(51, 65, 85))
            continue
        para_lines = []
        while i < n:
            curr_line = lines[i]
            curr_stripped = curr_line.strip()
            if not curr_stripped:
                break
            if (curr_stripped.startswith("```") or 
                re.match(r'^(#{1,6})\s+(.*)$', curr_stripped) or 
                curr_stripped.startswith("|") or 
                re.match(r'^([\s]*)([-*+])\s+(.*)$', curr_line) or 
                re.match(r'^([\s]*)(\d+)\.\s+(.*)$', curr_line) or
                curr_stripped in ("---", "***", "___") or 
                re.match(r'^[-*_]{3,}$', curr_stripped)):
                break
            para_lines.append(curr_stripped)
            i += 1
        if para_lines:
            para_text = " ".join(para_lines)
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            _add_inline_formatted_text(p, para_text, default_font="Segoe UI", default_size=Pt(10.5), default_color=RGBColor(51, 65, 85))


# --- State Definition ---
from typing import TypedDict, Annotated
import operator

class GapDeveloperState(TypedDict):
    messages: Annotated[List[BaseMessage], operator.add]
    account_id: str
    gap_id: str
    context: str
    full_analysis_context: Optional[str] # NUEVO
    research_results: str
    document_id: Optional[str]
    progress: int
    workspace_id: Optional[str]
    sources: Annotated[List[Dict[str, Any]], operator.add] # NUEVO
    visual_schema: Optional[str] # NUEVO

# --- Prompts ---

GAP_ANALYSIS_PROMPT = """Eres un Agente Especialista en Desarrollo de Brechas de Conocimiento. 
Tu objetivo es analizar una brecha de información detectada y proponer una solución estructurada.

Brecha a resolver: {gap_id}
Contexto proporcionado: {context}
Análisis detallado de origen: {full_analysis_context}

Instrucciones:
1. Analiza cómo la brecha se relaciona con el contexto actual y el análisis de origen.
2. Identifica qué información específica falta para cerrar esta brecha.
3. Si es necesario, utiliza tus herramientas de búsqueda para encontrar soluciones o marcos de trabajo externos.
4. Genera una propuesta de "Documento Borrador" que articule la solución.

Fecha actual: {date}
"""

DRAFT_WRITER_PROMPT = """Eres un Redactor Experto de Propuestas Técnicas. 
Basándote en la investigación realizada, redacta un documento borrador completo en formato Markdown.

El documento debe incluir:
1. **Título Sugerido** (Debe ser la primera línea del documento, comenzando con '#')
2. **Introducción y Contexto** (Relación de la brecha con el conocimiento actual)
3. **Análisis del Problema** (Por qué esta brecha es crítica)
4. **Propuesta de Solución Detallada** (Desarrollo técnico o conceptual basado en la investigación)
5. **Conclusiones y Recomendaciones**
6. **Bibliografía** (Lista de fuentes consultadas con sus enlaces si están disponibles)
7. **Esquema Visual** (Obligatorio. Utiliza HTML inline y Tailwind para crear una representación visual de la solución, dentro de etiquetas <visual_schema> y </visual_schema>)

INSTRUCCIONES DE CITACIÓN (ESTÁNDAR KOGNITO):
Cuando uses información de las fuentes proporcionadas, SIEMPRE cita la fuente usando el formato [número] al final de la oración o párrafo que use esa información.

Reglas Críticas:
1. Usa SOLO el número entre corchetes (ej. [1]). NUNCA incluyas palabras como "Fuente", "Ref" o "Cita" dentro de los corchetes.
2. Si usas múltiples fuentes, sepáralas así: [1] [2]. NO uses [1, 2].
3. Coloca las citas al final de las oraciones.
4. NO inventes números de citas.
5. Al final del documento, incluye una sección de "Bibliografía" o "Fuentes" que liste las fuentes utilizadas.

Investigación previa y fuentes:
{research_results}

Escribe el documento de forma profesional, clara y accionable. No uses introducciones como "Aquí tienes el borrador", simplemente escribe el contenido del documento en Markdown.
"""

# --- Nodes ---

async def research_node(state: GapDeveloperState, config: RunnableConfig) -> dict:
    logger.info("--- [GapDeveloper] Node: research_node ---")
    progress_callback = config.get("configurable", {}).get("progress_callback")
    
    if progress_callback:
        await progress_callback(20, "Investigando contexto y buscando soluciones externas...", "research")

    account_id = state.get("account_id")
    llm = await get_llm_for_user(account_id, purpose="main")
    
    if not llm:
        logger.warning(f"No se encontró LLM para el usuario {account_id}, usando fallback.")
        llm = get_main_llm()
    
    # Obtener herramientas de búsqueda y notas
    skill_manager = get_skill_manager()
    tools = await skill_manager.load_skills(
        account_id=account_id,
        relevant_categories=["rag_skill", "notes_skill", "document_management_skill", "analysis_and_insights_skill"]
    )
    
    # Filtrar herramientas de búsqueda web específicamente (Tavily o similar)
    search_tools = [t for t in tools if t.name in ["web_search", "deep_research", "web_scraper_tool"]]
    llm_with_tools = safe_bind_tools(llm, search_tools)

    prompt = GAP_ANALYSIS_PROMPT.format(
        gap_id=state["gap_id"],
        context=state["context"],
        full_analysis_context=state.get("full_analysis_context", "No se proporcionó contexto de análisis adicional."),
        date=get_today_str()
    )

    # Realizar investigación inicial
    response = await llm_with_tools.ainvoke([HumanMessage(content=prompt)])
    
    # Por simplicidad en esta versión, si el LLM decide usar herramientas, las ejecutamos linealmente
    # (En una versión más compleja usaríamos un bucle de ReAct)
    research_summary = response.content
    extracted_sources = []
    
    if hasattr(response, "tool_calls") and response.tool_calls:
        # Ejecutar la primera herramienta de búsqueda sugerida
        tool_call = response.tool_calls[0]
        tool_to_use = next((t for t in search_tools if t.name == tool_call["name"]), None)
        if tool_to_use:
            logger.info(f"Ejecutando herramienta de búsqueda: {tool_call['name']}")
            tool_result = await tool_to_use.ainvoke(tool_call["args"])
            
            # Formatear resultados para el LLM y extraer fuentes estructuradas
            if isinstance(tool_result, list):
                formatted_results = "\n\nFUENTES ENCONTRADAS:\n"
                for idx, item in enumerate(tool_result):
                    source_id = idx + 1
                    title = item.get("title") or item.get("name") or "Fuente externa"
                    url = item.get("url") or item.get("link") or ""
                    snippet = item.get("snippet") or item.get("content") or ""
                    
                    formatted_results += f"[{source_id}] TÍTULO: {title}\nURL: {url}\nCONTENIDO: {snippet}\n\n"
                    
                    if url:
                        extracted_sources.append({
                            "id": source_id,
                            "title": title,
                            "url": url,
                            "snippet": snippet,
                            "type": "web"
                        })
                research_summary += formatted_results
            elif isinstance(tool_result, dict):
                # Caso: DeepResearchTool devuelve un dict con 'context_for_llm' y 'sources'
                context = tool_result.get("context_for_llm") or str(tool_result)
                research_summary += f"\n\nResultados de investigación externa:\n{context}"
                
                tool_sources = tool_result.get("sources", [])
                if isinstance(tool_sources, list):
                    for idx, s in enumerate(tool_sources):
                        # Evitar duplicados de IDs si ya había fuentes (aunque en este nodo es el primer tool_call)
                        source_id = len(extracted_sources) + 1
                        
                        source_data = s if isinstance(s, dict) else (s.model_dump() if hasattr(s, 'model_dump') else {})
                        
                        title = source_data.get("title") or "Fuente externa"
                        url = source_data.get("url") or ""
                        snippet = source_data.get("snippet") or ""
                        
                        extracted_sources.append({
                            "id": source_id,
                            "title": title,
                            "url": url,
                            "snippet": snippet,
                            "type": source_data.get("type", "web")
                        })
                        
                        # Añadir al research_summary para que el redactor vea los IDs y el contenido
                        research_summary += f"\n[{source_id}] TÍTULO: {title}\nURL: {url}\nCONTENIDO: {snippet}\n"
            elif isinstance(tool_result, str):
                research_summary += f"\n\nResultados de investigación externa:\n{tool_result}"
            else:
                research_summary += f"\n\nResultados de investigación externa:\n{str(tool_result)}"

    return {"research_results": research_summary, "progress": 50, "sources": extracted_sources}

async def draft_writer_node(state: GapDeveloperState, config: RunnableConfig) -> dict:
    logger.info("--- [GapDeveloper] Node: draft_writer_node ---")
    progress_callback = config.get("configurable", {}).get("progress_callback")
    
    if progress_callback:
        await progress_callback(60, "Redactando el documento borrador detallado...", "writing")

    account_id = state.get("account_id")
    llm = await get_llm_for_user(account_id, purpose="main")
    
    if not llm:
        llm = get_main_llm()
    
    prompt = DRAFT_WRITER_PROMPT.format(
        research_results=state["research_results"]
    )
    
    response = await llm.ainvoke([HumanMessage(content=prompt)])
    draft_content = response.content
    
    import re
    visual_schema = ""
    schema_match = re.search(r"<visual_schema>(.*?)</visual_schema>", draft_content, re.DOTALL | re.IGNORECASE)
    if schema_match:
        visual_schema = schema_match.group(1).strip()
        draft_content = re.sub(r"<visual_schema>.*?</visual_schema>", "", draft_content, flags=re.DOTALL | re.IGNORECASE).strip()
    else:
        fallback_match = re.search(r"(<div style=.*?>.*?</div>)", draft_content, re.DOTALL | re.IGNORECASE)
        if fallback_match:
            visual_schema = fallback_match.group(1).strip()
            
    response.content = draft_content
    
    return {"messages": [response], "progress": 80, "visual_schema": visual_schema if visual_schema else None}

async def persistence_node(state: GapDeveloperState, config: RunnableConfig) -> dict:
    logger.info("--- [GapDeveloper] Node: persistence_node ---")
    progress_callback = config.get("configurable", {}).get("progress_callback")
    
    if progress_callback:
        await progress_callback(90, "Guardando documento en OnlyOffice...", "persistence")

    account_id = state.get("account_id")
    draft_content = state["messages"][-1].content if state["messages"] else "Sin contenido"
    
    # 1. Extraer título
    extracted_title = None
    for line in str(draft_content).split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('#'):
            extracted_title = line.lstrip('#').strip()
            break
        if "título" in line.lower() and ":" in line:
            extracted_title = line.split(":", 1)[1].strip()
            break
            
    if extracted_title:
        extracted_title = extracted_title.replace("**", "").replace("*", "").strip()
        title = extracted_title
        if len(title) > 100:
            title = title[:97] + "..."
    else:
        title = f"Borrador: Solución a Brecha - {state['gap_id'][:50]}"

    clean_content = str(draft_content).replace("```markdown", "").replace("```", "").strip()
    
    doc_id = None
    
    if DOCX_AVAILABLE:
        try:
            # Crear archivo físicamente
            doc_obj = docx.Document()
            # Agregar el título principal al documento
            doc_obj.add_heading(title, level=1)
            # Renderizar Markdown
            _parse_and_render_markdown(doc_obj, clean_content)
            
            # Guardar físicamente en OnlyOffice
            user_dir = ensure_onlyoffice_account_dir(account_id)
            unique_filename = f"{uuid.uuid4()}.docx"
            file_path = user_dir / unique_filename
            doc_obj.save(file_path)
            
            # Registrar en DB
            display_filename = f"Borrador - {title}.docx"
            # Limpiar caracteres inválidos para el nombre de archivo
            display_filename = re.sub(r'[\\/*?:"<>|]', "", display_filename)
            
            async with SessionLocal() as db:
                workspace_uuid = None
                ws_id = state.get("workspace_id")
                if ws_id and ws_id != "null":
                    try:
                        workspace_uuid = uuid.UUID(str(ws_id))
                    except ValueError:
                        pass
                
                new_doc = Document(
                    account_id=uuid.UUID(account_id),
                    workspace_id=workspace_uuid,
                    filename=display_filename,
                    extension="docx",
                    file_path=build_onlyoffice_relative_path(account_id, unique_filename)
                )
                db.add(new_doc)
                await db.commit()
                await db.refresh(new_doc)
                
                doc_id = str(new_doc.id)
                logger.info(f"Documento de OnlyOffice creado exitosamente con ID: {doc_id}")
        except Exception as e:
            logger.error(f"Error al generar/guardar el documento docx en OnlyOffice: {e}")
            doc_id = f"error_docx: {str(e)}"
    
    # Fallback si docx no está disponible o falló
    if not doc_id or doc_id.startswith("error"):
        logger.warning("Intentando fallback al sistema de notas...")
        skill_manager = get_skill_manager()
        tools = await skill_manager.load_skills(
            account_id=account_id,
            relevant_categories=["notes_skill"]
        )
        add_note_tool = next((t for t in tools if t.name == "add_note"), None)
        if add_note_tool:
            try:
                result = await add_note_tool.ainvoke({
                    "title": title,
                    "content": clean_content,
                    "workspace_id": state.get("workspace_id")
                })
                if isinstance(result, str):
                    try:
                        res_data = json.loads(result)
                        doc_id = str(res_data.get("id"))
                    except:
                        doc_id = "created_successfully"
                elif isinstance(result, dict):
                    doc_id = str(result.get("id"))
                logger.info(f"Nota de borrador creada (fallback) con ID: {doc_id}")
            except Exception as e:
                logger.error(f"Error al guardar nota de fallback: {e}")
                
    if progress_callback:
        await progress_callback(100, "Documento borrador desarrollado con éxito.", "complete")

    return {"document_id": doc_id, "progress": 100}


# --- Graph Assembly ---

def compile_gap_developer_graph():
    workflow = StateGraph(GapDeveloperState)
    
    workflow.add_node("research", research_node)
    workflow.add_node("writer", draft_writer_node)
    workflow.add_node("persistence", persistence_node)
    
    workflow.add_edge(START, "research")
    workflow.add_edge("research", "writer")
    workflow.add_edge("writer", "persistence")
    workflow.add_edge("persistence", END)
    
    return workflow.compile()
