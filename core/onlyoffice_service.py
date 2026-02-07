# core/onlyoffice_service.py

import os
import logging
import uuid
import io
import httpx
import jwt
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any, List
from sqlalchemy import select, update, text
from sqlalchemy.ext.asyncio import AsyncSession
from docx import Document

from core.database import Nota
from core.config import settings
from utils.onlyoffice_client import onlyoffice_client
from core.notes_manager import NotesManager

logger = logging.getLogger(__name__)

class OnlyOfficeService:
    """
    Servicio centralizado para la gestión de OnlyOffice.
    Maneja la lógica de negocio, seguridad y persistencia.
    """
    def __init__(self, db: AsyncSession):
        self.db = db
        self.notes_manager = NotesManager(db)

    async def get_editor_config(self, note_id: int, account_id: str) -> Optional[Dict[str, Any]]:
        """
        Genera la configuración segura para el editor de OnlyOffice.
        """
        note_data = await self.notes_manager.get_note_by_id(account_id, note_id)
        if not note_data:
            logger.warning(f"🚫 [OnlyOffice] Nota {note_id} no encontrada o sin acceso para {account_id}")
            return None

        # URLs base. OnlyOffice Server necesita acceder a estas.
        # INTERNAL_DOC_API_URL es la URL desde la perspectiva del servidor OnlyOffice (usualmente red docker)
        api_base_url = os.getenv("INTERNAL_DOC_API_URL") or settings.internal_api_server_url
        
        # URL de descarga del documento raw
        file_url = f"{api_base_url}/api/notes/{note_id}/download-raw"
        
        # Generamos una clave única basada en la última actualización
        updated_at = note_data.get('updated_at')
        if isinstance(updated_at, str):
            updated_at_dt = datetime.fromisoformat(updated_at.replace('Z', '+00:00'))
        else:
            updated_at_dt = datetime.now()
        
        document_key = f"note_{note_id}_{int(updated_at_dt.timestamp())}"
        
        # Callback URL robusta con token de seguridad
        callback_url = f"{api_base_url}/api/notes/onlyoffice-callback?note_id={note_id}&account_id={account_id}"

        config = {
            "document": {
                "fileType": "docx",
                "key": document_key,
                "title": note_data.get("title") or "Nota sin título",
                "url": file_url,
                "permissions": {
                    "edit": True,
                    "download": True,
                    "comment": True,
                    "print": True
                }
            },
            "documentType": "word",
            "editorConfig": {
                "callbackUrl": callback_url,
                "lang": "es",
                "mode": "edit",
                "user": {
                    "id": account_id,
                    "name": note_data.get("created_by_email", "Usuario Kognito")
                },
                "customization": {
                    "forcesave": True,
                    "chat": False,
                    "help": False,
                    "plugins": False,
                    "compactToolbar": False
                }
            },
            "width": "100%",
            "height": "100%",
            "type": "desktop"
        }

        # Aplicar seguridad JWT si está configurada
        if onlyoffice_client.secret:
            token = jwt.encode(config, onlyoffice_client.secret, algorithm="HS256")
            return {
                "token": token,
                "onlyoffice_url": settings.onlyoffice_url,
                "config": config # Opcional, pero útil para depuración
            }
        
        return {
            "token": "",
            "onlyoffice_url": settings.onlyoffice_url,
            "config": config
        }

    async def handle_callback(self, note_id: int, account_id: str, data: Dict[str, Any], remote_ip: Optional[str] = None) -> Dict[str, Any]:
        """
        Procesa el callback de OnlyOffice de forma segura con logging de auditoría.
        """
        status = data.get("status")
        logger.info(f"📩 [OnlyOffice-Audit] Callback recibido: Nota {note_id}, Status {status}, IP: {remote_ip}")

        # Registro de intentos de callback (Detección de intrusiones/Auditoría)
        if status is None:
            logger.warning(f"🚨 [OnlyOffice-Security] Intento de callback malformado detectado desde IP {remote_ip} para nota {note_id}")
            return {"error": 1, "message": "Invalid status"}

        # Status 2: El documento está listo para ser guardado (después de cerrar el editor o forzar guardado)
        # Status 6: El documento se está editando pero se fuerza el guardado
        if status in [2, 6]:
            download_url = data.get("url")
            if not download_url:
                logger.error(f"❌ [OnlyOffice] No se recibió URL de descarga en callback de nota {note_id}")
                return {"error": 1}

            content = await onlyoffice_client.download_file(download_url)
            if not content:
                logger.error(f"❌ [OnlyOffice] Error al descargar el documento actualizado para nota {note_id}")
                return {"error": 1}

            # Guardar archivo físico
            upload_dir = Path("media/notes")
            upload_dir.mkdir(parents=True, exist_ok=True)
            file_path = upload_dir / f"note_{note_id}.docx"
            
            try:
                with open(file_path, "wb") as f:
                    f.write(content)
                
                # Extraer texto del DOCX para la IA
                extracted_text = self._extract_text_from_docx(content)
                
                # Actualizar la nota en la DB (usamos SQL crudo para visual_content si no está en el ORM base, 
                # aunque parece que sí está por los logs previos)
                await self.db.execute(
                    text("UPDATE notas SET visual_content = :vc, content = :content, updated_at = :now WHERE id = :id"),
                    {
                        "vc": str(file_path), 
                        "content": extracted_text, 
                        "id": note_id,
                        "now": datetime.utcnow()
                    }
                )
                await self.db.commit()
                
                logger.info(f"✅ [OnlyOffice] Nota {note_id} actualizada correctamente desde OnlyOffice.")
                return {"error": 0}
            except Exception as e:
                logger.error(f"⚠️ [OnlyOffice] Error guardando nota {note_id}: {e}")
                return {"error": 1}

        return {"error": 0}

    def _extract_text_from_docx(self, docx_content: bytes) -> str:
        """Extrae el texto de un archivo DOCX para mantener la sincronía con la IA."""
        try:
            doc = Document(io.BytesIO(docx_content))
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"❌ [OnlyOffice] Error extrayendo texto de DOCX: {e}")
            return ""

    async def get_raw_document(self, note_id: int) -> Optional[io.BytesIO]:
        """Genera o recupera el documento DOCX para descarga de OnlyOffice."""
        # Buscamos si existe la nota
        result = await self.db.execute(select(Nota).where(Nota.id == note_id))
        note = result.scalar_one_or_none()
        if not note:
            return None

        file_path = os.path.join("media/notes", f"note_{note_id}.docx")
        
        if os.path.exists(file_path):
            with open(file_path, "rb") as f:
                return io.BytesIO(f.read())

        # Si no existe, generamos uno a partir del contenido Markdown/Texto
        doc = Document()
        doc.add_heading(note.title or "Nota sin título", level=1)
        for paragraph in note.content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
