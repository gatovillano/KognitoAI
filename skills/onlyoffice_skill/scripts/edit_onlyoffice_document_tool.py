# skills/onlyoffice_skill/scripts/edit_onlyoffice_document_tool.py

"""
Herramienta de Edición Avanzada para Documentos OnlyOffice.
Soporta Word (.docx) y Excel (.xlsx) con múltiples acciones de formato.
"""

import os
import re
import logging
import uuid
import shutil
import json
from typing import Any, Type, Optional, List
from pydantic import BaseModel, Field, field_validator
from langchain_core.tools import BaseTool
from sqlalchemy import select
from core.database import SessionLocal, Document
from core.onlyoffice_storage import resolve_onlyoffice_file_path
from datetime import datetime

# Librerías de procesamiento de documentos (Office)
DOCX_IMPORT_ERROR = None
try:
    import docx
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
except Exception as e:
    DOCX_AVAILABLE = False
    DOCX_IMPORT_ERROR = str(e)


try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    import openpyxl.styles.numbers as numbers
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

logger = logging.getLogger(__name__)

# Intentar importar settings, con fallback si no está disponible
try:
    from core.config import settings
    DEFAULT_DOCS_ROOT = os.path.join(settings.media_root, "documents")
except ImportError:
    DEFAULT_DOCS_ROOT = os.path.join(os.getenv("MEDIA_ROOT", "/media/documents"), "documents")
DOCUMENTS_ROOT = os.environ.get("ONLYOFFICE_DOCS_ROOT", DEFAULT_DOCS_ROOT)

# Regex para tokenizar estilos inline en markdown: bold-italic, bold, italic, code
inline_regex = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|__.*?__|__.*?__|\*.*?\*|_.*?_|`.*?`)')

def set_cell_background(cell, hex_color: str):
    """Establece el color de fondo (sombreado) de una celda en Word."""
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)
    except Exception as e:
        logger.warning(f"No se pudo establecer el fondo de celda: {e}")

def set_table_borders(table, hex_color: str = "CBD5E1"):
    """Aplica bordes sutiles horizontales de estilo minimalista a una tabla."""
    try:
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'  <w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)
    except Exception as e:
        logger.warning(f"No se pudieron aplicar bordes de tabla: {e}")

PREMIUM_THEMES = {
    "corporate_blue": {
        "header_bg": "1E3A8A",
        "header_fg": "FFFFFF",
        "zebra_bg": "F8FAFC",
        "accent_bg": "E2E8F0",
        "border_color": "CBD5E1",
        "font_name": "Segoe UI",
    },
    "emerald_green": {
        "header_bg": "065F46",
        "header_fg": "FFFFFF",
        "zebra_bg": "F0FDF4",
        "accent_bg": "DCFCE7",
        "border_color": "A7F3D0",
        "font_name": "Segoe UI",
    },
    "dark_slate": {
        "header_bg": "0F172A",
        "header_fg": "FFFFFF",
        "zebra_bg": "F1F5F9",
        "accent_bg": "E2E8F0",
        "border_color": "94A3B8",
        "font_name": "Segoe UI",
    },
    "sunset_amber": {
        "header_bg": "9A3412",
        "header_fg": "FFFFFF",
        "zebra_bg": "FFFBEB",
        "accent_bg": "FEF3C7",
        "border_color": "FDE68A",
        "font_name": "Segoe UI",
    },
    "violet_gold": {
        "header_bg": "581C87",
        "header_fg": "FFFFFF",
        "zebra_bg": "FAF5FF",
        "accent_bg": "F3E8FF",
        "border_color": "D8B4FE",
        "font_name": "Segoe UI",
    },
    "minimalist_gray": {
        "header_bg": "374151",
        "header_fg": "FFFFFF",
        "zebra_bg": "F9FAFB",
        "accent_bg": "E5E7EB",
        "border_color": "D1D5DB",
        "font_name": "Segoe UI",
    },
}

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    """Establece márgenes/padding internos (en dxa: 20 dxa = 1 pt) de una celda en Word."""
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="{top}" w:type="dxa"/>'
            f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'  <w:left w:w="{left}" w:type="dxa"/>'
            f'  <w:right w:w="{right}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)
    except Exception as e:
        logger.warning(f"No se pudieron aplicar márgenes a celda: {e}")

def set_row_header(row):
    """Marca una fila como encabezado repetible en Word (<w:tblHeader/>)."""
    try:
        trPr = row._tr.get_or_add_trPr()
        tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
        trPr.append(tblHeader)
    except Exception as e:
        logger.warning(f"No se pudo establecer tblHeader: {e}")

def set_row_cant_split(row):
    """Evita que una fila se corte entre páginas en Word (<w:cantSplit/>)."""
    try:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
        trPr.append(cantSplit)
    except Exception as e:
        logger.warning(f"No se pudo establecer cantSplit: {e}")

def set_total_row_borders(row, hex_color="CBD5E1"):
    """Borde superior sencillo y borde inferior doble en fila de totales en Word."""
    for cell in row.cells:
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{hex_color}"/>'
                f'  <w:bottom w:val="double" w:sz="12" w:space="0" w:color="{hex_color}"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)
        except Exception:
            pass

def _add_inline_formatted_text(paragraph, text: str, default_font: str = "Segoe UI", default_size = Pt(10.5), default_color = RGBColor(51, 65, 85)):
    """Divide un texto por marcadores inline de Markdown y añade Runs formateados."""
    if not text:
        return
    
    tokens = inline_regex.split(text)
    for token in tokens:
        if not token:
            continue
        
        is_bold = False
        is_italic = False
        is_code = False
        clean_text = token
        
        if token.startswith("***") and token.endswith("***") and len(token) > 6:
            is_bold = True
            is_italic = True
            clean_text = token[3:-3]
        elif token.startswith("**") and token.endswith("**") and len(token) > 4:
            is_bold = True
            clean_text = token[2:-2]
        elif token.startswith("__") and token.endswith("__") and len(token) > 4:
            is_bold = True
            clean_text = token[2:-2]
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            is_italic = True
            clean_text = token[1:-1]
        elif token.startswith("_") and token.endswith("_") and len(token) > 2:
            is_italic = True
            clean_text = token[1:-1]
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            is_code = True
            clean_text = token[1:-1]
            
        run = paragraph.add_run(clean_text)
        if is_code:
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(239, 68, 68) # soft red-500
            try:
                rPr = run._r.get_or_add_rPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>') # slate-100
                rPr.append(shd)
            except Exception:
                pass
        else:
            run.font.name = default_font
            run.font.size = default_size
            run.font.color.rgb = default_color
            run.bold = is_bold
            run.italic = is_italic

def _parse_and_render_markdown(doc_obj, md_text: str):
    """Parsea texto Markdown estructurado y lo añade como elementos nativos y estilizados a docx."""
    if not md_text:
        return
        
    lines = md_text.split('\n')
    i = 0
    n = len(lines)
    
    while i < n:
        line = lines[i]
        stripped = line.strip()
        
        # 1. Líneas vacías
        if not stripped:
            i += 1
            continue
            
        # 2. Bloques de código
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < n:
                i += 1 # Consumir cierre
                
            code_text = "\n".join(code_lines)
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.left_indent = Pt(12)
            
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(15, 23, 42) # dark slate
            
            try:
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>') # slate-50 background
                pPr.append(shd)
                pbdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="CBD5E1"/>' # 2.25pt border left
                    f'</w:pBdr>'
                )
                pPr.append(pbdr)
            except Exception:
                pass
            continue
            
        # 3. Títulos (H1, H2, H3, H4)
        h_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if h_match:
            level = len(h_match.group(1))
            content = h_match.group(2).strip()
            
            p = doc_obj.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            if level == 1:
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
                _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(18), default_color=RGBColor(30, 58, 138))
                for r in p.runs:
                    r.bold = True
            elif level == 2:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(14), default_color=RGBColor(37, 99, 235))
                for r in p.runs:
                    r.bold = True
            else:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(12), default_color=RGBColor(30, 41, 59))
                for r in p.runs:
                    r.bold = True
            i += 1
            continue
            
        # 4. Regla horizontal
        if stripped in ("---", "***", "___") or re.match(r'^[-*_]{3,}$', stripped):
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            try:
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E2E8F0"/></w:pBdr>')
                pPr.append(pBdr)
            except Exception:
                pass
            i += 1
            continue
            
        # 5. Tablas
        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
                
            rows = []
            for tl in table_lines:
                cols = [col.strip() for col in tl.split("|")[1:-1]]
                rows.append(cols)
                
            if len(rows) > 1 and all(re.match(r'^:?-+:?$', c) for c in rows[1]):
                rows.pop(1) # Remover fila de guiones divisores
                
            if rows:
                num_rows = len(rows)
                num_cols = max(len(r) for r in rows)
                
                table = doc_obj.add_table(rows=num_rows, cols=num_cols)
                table.autofit = True
                set_table_borders(table)
                
                for row_idx, r_data in enumerate(rows):
                    for col_idx, cell_val in enumerate(r_data):
                        if col_idx < num_cols:
                            cell = table.cell(row_idx, col_idx)
                            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                            p.paragraph_format.space_before = Pt(4)
                            p.paragraph_format.space_after = Pt(4)
                            p.paragraph_format.line_spacing = 1.0
                            
                            if row_idx == 0:
                                _add_inline_formatted_text(p, cell_val, default_font="Segoe UI", default_size=Pt(10), default_color=RGBColor(255, 255, 255))
                                for run in p.runs:
                                    run.bold = True
                                set_cell_background(cell, "1E3A8A") # azul rey
                            else:
                                _add_inline_formatted_text(p, cell_val, default_font="Segoe UI", default_size=Pt(9.5), default_color=RGBColor(51, 65, 85))
                                if row_idx % 2 == 0:
                                    set_cell_background(cell, "F8FAFC") # zebra
                                else:
                                    set_cell_background(cell, "FFFFFF")
                
                # Párrafo vacío para espaciado post-tabla
                p_after = doc_obj.add_paragraph()
                p_after.paragraph_format.space_before = Pt(0)
                p_after.paragraph_format.space_after = Pt(6)
            continue
            
        # 6. Listas (viñetas o numeradas)
        bullet_match = re.match(r'^([\s]*)([-*+])\s+(.*)$', line)
        num_match = re.match(r'^([\s]*)(\d+)\.\s+(.*)$', line)
        if bullet_match or num_match:
            list_lines = []
            while i < n:
                curr_line = lines[i]
                if not curr_line.strip():
                    break
                b_m = re.match(r'^([\s]*)([-*+])\s+(.*)$', curr_line)
                n_m = re.match(r'^([\s]*)(\d+)\.\s+(.*)$', curr_line)
                if b_m or n_m:
                    list_lines.append((curr_line, b_m, n_m))
                    i += 1
                else:
                    # Continuación identada
                    if list_lines and curr_line.startswith("   "):
                        prev_text, prev_bm, prev_nm = list_lines[-1]
                        updated_text = prev_text + " " + curr_line.strip()
                        b_m_up = re.match(r'^([\s]*)([-*+])\s+(.*)$', updated_text)
                        n_m_up = re.match(r'^([\s]*)(\d+)\.\s+(.*)$', updated_text)
                        list_lines[-1] = (updated_text, b_m_up, n_m_up)
                        i += 1
                    else:
                        break
                        
            for item_line, bm, nm in list_lines:
                p = doc_obj.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                
                if bm:
                    indent_level = len(bm.group(1)) // 2
                    content = bm.group(3).strip()
                    p.style = 'List Bullet'
                    p.paragraph_format.left_indent = Pt(18 * (indent_level + 1))
                    _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(10.5), default_color=RGBColor(51, 65, 85))
                else:
                    indent_level = len(nm.group(1)) // 2
                    content = nm.group(3).strip()
                    p.style = 'List Number'
                    p.paragraph_format.left_indent = Pt(18 * (indent_level + 1))
                    _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(10.5), default_color=RGBColor(51, 65, 85))
            continue
            
        # 7. Párrafo estándar (agrupamiento de líneas contiguas)
        para_lines = []
        while i < n:
            curr_line = lines[i]
            curr_stripped = curr_line.strip()
            if not curr_stripped:
                break
            if (curr_stripped.startswith("```") or 
                re.match(r'^(#{1,6})\s+(.*)$', curr_stripped) or 
                curr_stripped.startswith("|") or 
                re.match(r'^([\s]*)([-*+])\s+(.*)$', curr_line) or 
                re.match(r'^([\s]*)(\d+)\.\s+(.*)$', curr_line) or
                curr_stripped in ("---", "***", "___") or 
                re.match(r'^[-*_]{3,}$', curr_stripped)):
                break
            para_lines.append(curr_stripped)
            i += 1
            
        if para_lines:
            para_text = " ".join(para_lines)
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            _add_inline_formatted_text(p, para_text, default_font="Segoe UI", default_size=Pt(10.5), default_color=RGBColor(51, 65, 85))


# --- Acciones disponibles ---
SUPPORTED_ACTIONS = [
    "append",            # Añadir párrafo de texto al final
    "append_heading",    # Añadir un título (H1, H2, H3)
    "append_list",       # Añadir una lista de viñetas
    "replace",           # Buscar y reemplazar texto en todo el documento
    "replace_section",   # Reemplazar la sección que empieza con cierto texto
    "insert_table",      # Insertar una tabla al final
    "insert_styled_table", # Insertar una tabla con tema premium y totales en Word
    "apply_bold",        # Poner en negrita las ocurrencias de un texto
    "clear_and_write",   # Borrar todo el contenido y escribir nuevo
    # DOCX new actions
    "edit_paragraph",    # Editar párrafo por índice o búsqueda
    "insert_image",      # Insertar imagen en el documento
    "set_page_layout",   # Configurar márgenes y orientación de página
    "add_header_footer", # Añadir encabezados y pies de página
    "apply_cell_style",  # Aplicar estilo a celdas de tabla
    # XLSX - Gestión de pestañas/hojas
    "xlsx_create_sheet",      # Crear una nueva pestaña/hoja en Excel
    "xlsx_rename_sheet",      # Renombrar una pestaña/hoja en Excel
    "xlsx_delete_sheet",      # Eliminar una pestaña/hoja en Excel
    "xlsx_list_sheets",       # Listar las pestañas/hojas disponibles en Excel
    "xlsx_copy_sheet",        # Copiar una pestaña/hoja con nuevo nombre
    "xlsx_move_sheet",        # Mover/reordenar una pestaña a una posición
    "xlsx_clear_sheet",       # Borrar todo el contenido de una hoja sin eliminarla
    "xlsx_get_sheet_info",    # Obtener info de una hoja (dimensiones, rango, datos)
    "xlsx_freeze_panes",      # Congelar filas/columnas superiores/izquierdas
    "xlsx_protect_sheet",     # Proteger hoja con contraseña (solo lectura)
    "xlsx_set_sheet_tab_color", # Cambiar color de pestaña de hoja existente
    # XLSX - Edición de contenido
    "xlsx_create_table",      # Crear tabla con diseño premium y fórmulas en Excel
    "xlsx_write_cell",        # Escribir en una celda específica de una hoja Excel
    "xlsx_append_row",        # Añadir una fila al final de una hoja Excel
    "xlsx_write_range",       # Escribir en un rango de celdas
    "xlsx_format_cells",      # Formatear celdas con estilos
    "xlsx_insert_row",        # Insertar fila en posición específica
    "xlsx_insert_column",     # Insertar columna en posición específica
    "xlsx_delete_row",        # Eliminar fila
    "xlsx_delete_column",     # Eliminar columna
    "xlsx_merge_cells",       # Fusionar celdas
    "xlsx_set_column_width",  # Establecer ancho de columna
    "xlsx_set_row_height",    # Establecer altura de fila
]


class EditOnlyOfficeInput(BaseModel):
    document_id: str = Field(..., description="ID UUID del documento a editar.")
    action: str = Field(
        ...,
        description=(
            "Acción a realizar. Opciones disponibles: "
            # --- DOCX ---
            "'append' (añadir texto/Markdown formateado al final), "
            "'append_heading' (añadir título - requiere heading_level 1/2/3), "
            "'append_list' (añadir lista de viñetas - requiere list_items), "
            "'replace' (buscar y reemplazar - requiere search_text), "
            "'replace_section' (reemplaza la sección que empieza con search_text), "
            "'insert_table' (insertar tabla - requiere table_data), "
            "'insert_styled_table' (insertar tabla premium con temas y totales en Word - requiere table_data), "
            "'apply_bold' (poner en negrita - requiere search_text), "
            "'clear_and_write' (borrar todo el contenido y escribir nuevo texto/Markdown formateado), "
            "'edit_paragraph' (editar párrafo por índice o búsqueda - requiere paragraph_index o search_text), "
            "'insert_image' (insertar imagen - requiere image_path y opcionalmente image_width/image_height), "
            "'set_page_layout' (configurar márgenes/orientación - requiere page_margins y/o page_orientation), "
            "'add_header_footer' (añadir encabezado/pie - requiere header_footer_type y content), "
            "'apply_cell_style' (aplicar estilo a celda - requiere cell_style y cell_coords), "
            # --- XLSX - Gestión de pestañas ---
            "'xlsx_create_sheet' (crear nueva pestaña en Excel - requiere sheet_name, opcional tab_color, sheet_index), "
            "'xlsx_rename_sheet' (renombrar pestaña en Excel - requiere sheet_name (nuevo nombre) y opcionalmente search_text como nombre anterior), "
            "'xlsx_delete_sheet' (eliminar pestaña en Excel - requiere sheet_name), "
            "'xlsx_list_sheets' (listar todas las pestañas del archivo Excel - devuelve nombre, color y estadísticas), "
            "'xlsx_copy_sheet' (copiar una pestaña a un nuevo nombre - requiere sheet_name como origen y text como nuevo nombre, opcional sheet_index), "
            "'xlsx_move_sheet' (mover/reordenar pestaña a posición 0-based - requiere sheet_name y sheet_index), "
            "'xlsx_clear_sheet' (borrar TODO el contenido de una hoja sin eliminarla - requiere sheet_name), "
            "'xlsx_get_sheet_info' (obtener estadísticas de una hoja: dimensiones, rango usado, muestra de datos - requiere sheet_name), "
            "'xlsx_freeze_panes' (congelar filas/columnas: cell='B2' congela fila 1 y columna A - requiere cell y sheet_name), "
            "'xlsx_protect_sheet' (proteger hoja con contraseña para evitar edición manual - requiere sheet_name, opcional text como contraseña), "
            "'xlsx_set_sheet_tab_color' (cambiar color de pestaña de hoja existente - requiere sheet_name y tab_color hex), "
            # --- XLSX - Edición de contenido ---
            "'xlsx_create_table' (crear tabla premium con fórmulas en Excel - requiere table_data), "
            "'xlsx_write_cell' (escribir valor o fórmula en celda Excel - requiere cell y text; si text empieza con '=' se guarda como fórmula), "
            "'xlsx_append_row' (añadir fila a Excel - requiere row_data), "
            "'xlsx_write_range' (escribir en rango - requiere range_address y range_data), "
            "'xlsx_format_cells' (formatear celdas - requiere range_address y format_options), "
            "'xlsx_insert_row' (insertar fila - requiere row_index), "
            "'xlsx_insert_column' (insertar columna - requiere column_index), "
            "'xlsx_delete_row' (eliminar fila - requiere row_index), "
            "'xlsx_delete_column' (eliminar columna - requiere column_index), "
            "'xlsx_merge_cells' (fusionar celdas - requiere range_address), "
            "'xlsx_set_column_width' (ancho columna - requiere column_index y width), "
            "'xlsx_set_row_height' (altura fila - requiere row_index y height)."
        ),
    )
    text: Optional[str] = Field(None, description="Texto principal a insertar (soporta Markdown enriquecido) o nuevo texto en 'replace'. Para Excel, si empieza con '=' se tratará como fórmula (ej. '=SUM(A1:A10)').")
    search_text: Optional[str] = Field(None, description="Texto a buscar (para 'replace', 'replace_section', 'apply_bold', o nombre de pestaña anterior en 'xlsx_rename_sheet').")
    heading_level: Optional[int] = Field(None, description="Nivel del título: 1 (H1), 2 (H2), 3 (H3). Usar con 'append_heading'.")
    list_items: Optional[List[str]] = Field(None, description="Lista de cadenas para crear una lista de viñetas. Usar con 'append_list'.")
    table_data: Optional[List[List[str]]] = Field(None, description="Matriz de filas/columnas para crear una tabla. La primera fila es el encabezado. Usar con 'insert_table', 'insert_styled_table', 'xlsx_create_table'.")
    theme: Optional[str] = Field(None, description="Tema de diseño premium para tablas ('corporate_blue', 'emerald_green', 'dark_slate', 'sunset_amber', 'violet_gold', 'minimalist_gray').")
    has_total_row: Optional[bool] = Field(None, description="Si es True, aplica formato de totales (borde doble inferior, texto en negrita y fondo de acento) a la última fila.")
    auto_calculate_totals: Optional[bool] = Field(None, description="Si es True, calcula automáticamente sumas de columnas numéricas en la fila de totales.")
    col_alignments: Optional[List[str]] = Field(None, description="Lista de alineaciones por columna ('left', 'center', 'right').")
    col_widths: Optional[List[float]] = Field(None, description="Lista de anchos de columna en cm o caracteres.")
    cell_padding: Optional[dict] = Field(None, description="Margen/padding interno de celda en pt {'top': 6, 'bottom': 6, 'left': 8, 'right': 8}.")
    total_formulas: Optional[Any] = Field(None, description="Fórmulas para la fila de totales en Excel (ej: {'B': 'SUM', 'C': 'AVERAGE'} o lista de fórmulas).")
    column_formats: Optional[dict] = Field(None, description="Formatos numéricos por columna en Excel (ej: {'B': '$#,##0.00', 'C': '0.0%'}).")
    auto_fit_columns: Optional[bool] = Field(True, description="Ajustar automáticamente el ancho de columnas en Excel.")
    start_cell: Optional[str] = Field("A1", description="Celda de inicio en Excel para crear tabla (ej: 'A1'). Usar con 'xlsx_create_table'.")
    # Para Excel
    sheet_name: Optional[str] = Field(None, description="Nombre de la hoja/pestaña de Excel. Si la pestaña indicada no existe, se creará automáticamente.")
    tab_color: Optional[str] = Field(None, description="Color hexadecimal para la pestaña de Excel (ej: '1E3A8A'). Usar con 'xlsx_create_sheet' o 'xlsx_rename_sheet'.")
    sheet_index: Optional[int] = Field(None, description="Posición 0-based de la pestaña en el libro Excel. Usar con 'xlsx_create_sheet'.")
    cell: Optional[str] = Field(None, description="Coordenada de celda Excel (ej: 'A1', 'B3'). Usar con 'xlsx_write_cell'.")
    row_data: Optional[List[str]] = Field(None, description="Lista de valores para añadir como fila. Para fórmulas, iniciar con '='. Usar con 'xlsx_append_row'.")
    # Nuevos campos DOCX
    paragraph_index: Optional[int] = Field(None, description="Índice del párrafo a editar (0-based). Usar con 'edit_paragraph'.")
    image_path: Optional[str] = Field(None, description="Ruta del archivo de imagen a insertar. Usar con 'insert_image'.")
    image_width: Optional[float] = Field(None, description="Ancho de la imagen en cm. Por defecto 10cm. Usar con 'insert_image'.")
    image_height: Optional[float] = Field(None, description="Alto de la imagen en cm. Por defecto autoajustado. Usar con 'insert_image'.")
    page_margins: Optional[dict] = Field(None, description="Diccionario con márgenes en cm: {'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5}. Usar con 'set_page_layout'.")
    page_orientation: Optional[str] = Field(None, description="Orientación de página: 'portrait' o 'landscape'. Usar con 'set_page_layout'.")
    header_footer_type: Optional[str] = Field(None, description="Tipo: 'header' o 'footer'. Usar con 'add_header_footer'.")
    header_footer_content: Optional[str] = Field(None, description="Contenido del encabezado/pie de página. Usar con 'add_header_footer'.")
    cell_style: Optional[dict] = Field(None, description="Estilo a aplicar: {'bold': True, 'italic': False, 'font_size': 12, 'font_color': 'FF0000'}. Usar con 'apply_cell_style'.")
    cell_coords: Optional[str] = Field(None, description="Coordenadas de celda: 'A1'. Usar con 'apply_cell_style'.")
    # Nuevos campos XLSX
    range_address: Optional[str] = Field(None, description="Rango de celdas (ej: 'A1:C3'). Usar con 'xlsx_write_range', 'xlsx_format_cells', 'xlsx_merge_cells'.")
    range_data: Optional[List[List[str]]] = Field(None, description="Datos para escribir en el rango. Fórmulas deben empezar con '='. Usar con 'xlsx_write_range'.")
    format_options: Optional[dict] = Field(None, description="Opciones de formato: {'bold': True, 'italic': True, 'font_size': 12, 'font_color': 'FF0000', 'font_name': 'Arial', 'fill_color': 'FFFF00', 'align': 'center', 'valign': 'center', 'wrap_text': True, 'border': 'thin', 'number_format': '0.00%'}. Usar con 'xlsx_format_cells'.")
    row_index: Optional[int] = Field(None, description="Índice de fila (0-based). Usar con 'xlsx_insert_row', 'xlsx_delete_row'.")
    column_index: Optional[int] = Field(None, description="Índice de columna (0-based). Usar con 'xlsx_insert_column', 'xlsx_delete_column'.")
    width: Optional[float] = Field(None, description="Ancho en cm. Usar con 'xlsx_set_column_width'.")
    height: Optional[float] = Field(None, description="Alto en filas. Usar con 'xlsx_set_row_height'.")

    @field_validator(
        "list_items",
        "table_data",
        "row_data",
        "page_margins",
        "cell_style",
        "range_data",
        "format_options",
        "col_alignments",
        "col_widths",
        "cell_padding",
        "total_formulas",
        "column_formats",
        mode="before"
    )
    @classmethod
    def parse_json_fields(cls, v: Any) -> Any:
        if isinstance(v, str):
            v_stripped = v.strip()
            if (v_stripped.startswith('[') and v_stripped.endswith(']')) or \
               (v_stripped.startswith('{') and v_stripped.endswith('}')):
                try:
                    return json.loads(v)
                except Exception:
                    pass
        return v


class EditOnlyOfficeDocumentTool(BaseTool):
    name: str = "edit_onlyoffice_document"
    description: str = (
        "Edita documentos de OnlyOffice (Word .docx y Excel .xlsx) con operaciones avanzadas. "
        "Soporta añadir texto/Markdown completo con formatos complejos (negrita, cursiva, "
        "bloques de código, listas numeradas/viñetas, tablas estructuradas con diseño premium y fórmulas, títulos) "
        "que se traducen automáticamente a estilos nativos premium sin dejar caracteres crudos de Markdown. "
        "Permite gestionar múltiples pestañas en Excel (crear, renombrar, eliminar, cambiar color de pestaña), "
        "buscar/reemplazar texto, poner texto en negrita, reescribir secciones o modificar celdas de Excel. "
        "Los cambios se guardan directamente en el servidor con respaldo automático. "
        "El usuario verá los cambios al recargar el editor. "
        "Usa esta herramienta cuando el usuario pida redactar, corregir o dar formato a su documento."
    )
    args_schema: Type[BaseModel] = EditOnlyOfficeInput

    account_id: str = Field(..., description="ID de cuenta del usuario, inyectado automáticamente.")

    async def _arun(
        self,
        document_id: str,
        action: str,
        text: Optional[str] = None,
        search_text: Optional[str] = None,
        heading_level: Optional[int] = None,
        list_items: Optional[List[str]] = None,
        table_data: Optional[List[List[str]]] = None,
        theme: Optional[str] = None,
        has_total_row: Optional[bool] = None,
        auto_calculate_totals: Optional[bool] = None,
        col_alignments: Optional[List[str]] = None,
        col_widths: Optional[List[float]] = None,
        cell_padding: Optional[dict] = None,
        total_formulas: Optional[Any] = None,
        column_formats: Optional[dict] = None,
        auto_fit_columns: Optional[bool] = True,
        start_cell: Optional[str] = "A1",
        sheet_name: Optional[str] = None,
        tab_color: Optional[str] = None,
        sheet_index: Optional[int] = None,
        cell: Optional[str] = None,
        row_data: Optional[List[str]] = None,
        # Nuevos parámetros DOCX
        paragraph_index: Optional[int] = None,
        image_path: Optional[str] = None,
        image_width: Optional[float] = None,
        image_height: Optional[float] = None,
        page_margins: Optional[dict] = None,
        page_orientation: Optional[str] = None,
        header_footer_type: Optional[str] = None,
        header_footer_content: Optional[str] = None,
        cell_style: Optional[dict] = None,
        cell_coords: Optional[str] = None,
        # Nuevos parámetros XLSX
        range_address: Optional[str] = None,
        range_data: Optional[List[List[str]]] = None,
        format_options: Optional[dict] = None,
        row_index: Optional[int] = None,
        column_index: Optional[int] = None,
        width: Optional[float] = None,
        height: Optional[float] = None,
        **kwargs: Any,
    ) -> str:
        try:
            acc_id = uuid.UUID(self.account_id)
            doc_id = uuid.UUID(document_id)

            async with SessionLocal() as db:
                stmt = select(Document).where(
                    Document.id == doc_id,
                    Document.account_id == acc_id,
                )
                result = await db.execute(stmt)
                doc = result.scalar_one_or_none()

                if not doc:
                    return f"❌ Documento con ID '{document_id}' no encontrado o no tienes permiso para editarlo."

                ext = doc.extension.lower().lstrip(".")

                # Construir ruta física
                try:
                    file_path = str(resolve_onlyoffice_file_path(doc.file_path))
                except Exception as exc:
                    return f"❌ Ruta física inválida para documento: {exc}"

                if not os.path.exists(file_path):
                    return f"❌ El archivo físico '{doc.filename}' no se encuentra en el servidor ({file_path})."

                # Dispatch por tipo de archivo
                if ext == "docx":
                    msg = await self._edit_docx(
                        file_path, action, text, search_text,
                        heading_level, list_items, table_data,
                        theme, has_total_row, auto_calculate_totals,
                        col_alignments, col_widths, cell_padding,
                        paragraph_index, image_path, image_width, image_height,
                        page_margins, page_orientation, header_footer_type, header_footer_content,
                        cell_style, cell_coords, doc.filename
                    )
                elif ext in ("xlsx", "xls"):
                    msg = await self._edit_xlsx(
                        file_path, action, text, cell,
                        sheet_name, tab_color, sheet_index, search_text,
                        row_data, table_data, theme,
                        has_total_row, total_formulas, column_formats,
                        auto_fit_columns, start_cell, range_address, range_data,
                        format_options, row_index, column_index, width, height,
                        doc.filename
                    )
                else:
                    return (
                        f"❌ El tipo de archivo '.{ext}' no está soportado para edición. "
                        "Se soportan: .docx (Word) y .xlsx (Excel)."
                    )

                if msg.startswith("❌"):
                    return msg

                # Actualizar timestamp en DB para invalidar caché de OnlyOffice
                doc.updated_at = datetime.now()
                await db.commit()
                return msg

        except Exception as e:
            logger.error(f"Error en EditOnlyOfficeDocumentTool: {e}", exc_info=True)
            return f"❌ Error inesperado al editar el documento: {str(e)}"

    # ------------------------------------------------------------------ #
    #  DOCX Editing
    # ------------------------------------------------------------------ #
    async def _edit_docx(
        self,
        file_path: str,
        action: str,
        text: Optional[str],
        search_text: Optional[str],
        heading_level: Optional[int],
        list_items: Optional[List[str]],
        table_data: Optional[List[List[str]]],
        theme: Optional[str],
        has_total_row: Optional[bool],
        auto_calculate_totals: Optional[bool],
        col_alignments: Optional[List[str]],
        col_widths: Optional[List[float]],
        cell_padding: Optional[dict],
        paragraph_index: Optional[int],
        image_path: Optional[str],
        image_width: Optional[float],
        image_height: Optional[float],
        page_margins: Optional[dict],
        page_orientation: Optional[str],
        header_footer_type: Optional[str],
        header_footer_content: Optional[str],
        cell_style: Optional[dict],
        cell_coords: Optional[str],
        filename: str,
    ) -> str:
        if not DOCX_AVAILABLE:
            return f"❌ La librería 'python-docx' no está instalada. No se puede editar documentos Word. Error de importación: {DOCX_IMPORT_ERROR}"

        try:
            doc_obj = docx.Document(file_path)

            if action == "append":
                if not text:
                    return "❌ La acción 'append' requiere el campo 'text'."
                _parse_and_render_markdown(doc_obj, text)
                msg = f"✅ Contenido (Markdown parsed) añadido al final de '{filename}'."

            elif action == "append_heading":
                if not text:
                    return "❌ La acción 'append_heading' requiere el campo 'text'."
                level = heading_level if heading_level in (1, 2, 3) else 2
                doc_obj.add_heading(text, level=level)
                msg = f"✅ Título H{level} '{text}' añadido a '{filename}'."

            elif action == "append_list":
                if not list_items:
                    return "❌ La acción 'append_list' requiere el campo 'list_items'."
                for item in list_items:
                    doc_obj.add_paragraph(item, style="List Bullet")
                msg = f"✅ Lista de {len(list_items)} ítems añadida a '{filename}'."

            elif action == "replace":
                if not search_text:
                    return "❌ La acción 'replace' requiere 'search_text'."
                if text is None:
                    return "❌ La acción 'replace' requiere 'text' con el nuevo contenido."
                count = self._replace_in_docx(doc_obj, search_text, text)
                if count == 0:
                    return f"⚠️ No se encontró '{search_text}' en el documento."
                msg = f"✅ Reemplazadas {count} ocurrencia(s) de '{search_text}' en '{filename}'."

            elif action == "replace_section":
                if not search_text:
                    return "❌ La acción 'replace_section' requiere 'search_text' (inicio de la sección)."
                if text is None:
                    return "❌ La acción 'replace_section' requiere 'text' con el nuevo contenido."
                replaced = self._replace_section_in_docx(doc_obj, search_text, text)
                if not replaced:
                    return f"⚠️ No se encontró ningún párrafo que empiece con '{search_text}'."
                msg = f"✅ Sección que comenzaba con '{search_text}' reemplazada en '{filename}'."

            elif action in ("insert_table", "insert_styled_table"):
                msg = self._handle_insert_styled_table(
                    doc_obj, table_data, theme, has_total_row,
                    auto_calculate_totals, col_alignments, col_widths, cell_padding, filename
                )

            elif action == "apply_bold":
                if not search_text:
                    return "❌ La acción 'apply_bold' requiere 'search_text'."
                count = self._apply_bold_in_docx(doc_obj, search_text)
                if count == 0:
                    return f"⚠️ No se encontró '{search_text}' para poner en negrita."
                msg = f"✅ '{search_text}' puesto en negrita ({count} ocurrencia(s)) en '{filename}'."

            elif action == "clear_and_write":
                if not text:
                    return "❌ La acción 'clear_and_write' requiere 'text'."
                # Eliminar todos los elementos existentes en el cuerpo para una limpieza completa
                for element in list(doc_obj.element.body):
                    if element.tag.endswith('sectPr'):
                        continue
                    doc_obj.element.body.remove(element)
                # Escribir el nuevo contenido parsed a partir de Markdown
                _parse_and_render_markdown(doc_obj, text)
                msg = f"✅ Documento '{filename}' reescrito con nuevo contenido estructurado."

            elif action == "edit_paragraph":
                if paragraph_index is None and not search_text:
                    return "❌ La acción 'edit_paragraph' requiere 'paragraph_index' o 'search_text'."
                msg = self._handle_edit_paragraph(doc_obj, paragraph_index, search_text, text)

            elif action == "insert_image":
                if not image_path:
                    return "❌ La acción 'insert_image' requiere 'image_path'."
                msg = self._handle_insert_image(doc_obj, image_path, image_width, image_height)

            elif action == "set_page_layout":
                msg = self._handle_set_page_layout(doc_obj, page_margins, page_orientation)

            elif action == "add_header_footer":
                if not header_footer_type:
                    return "❌ La acción 'add_header_footer' requiere 'header_footer_type' ('header' o 'footer')."
                if not header_footer_content:
                    return "❌ La acción 'add_header_footer' requiere 'header_footer_content'."
                msg = self._handle_header_footer(doc_obj, header_footer_type, header_footer_content)

            elif action == "apply_cell_style":
                if not cell_coords:
                    return "❌ La acción 'apply_cell_style' requiere 'cell_coords' (ej: 'A1')."
                if not cell_style:
                    return "❌ La acción 'apply_cell_style' requiere 'cell_style' (diccionario con estilos)."
                msg = self._handle_cell_style(doc_obj, cell_coords, cell_style)

            else:
                return (
                    f"❌ Acción '{action}' no reconocida para .docx. "
                    f"Acciones disponibles: {', '.join(SUPPORTED_ACTIONS)}"
                )

            # Backup y guardado
            self._backup_file(file_path)
            doc_obj.save(file_path)
            return msg

        except Exception as e:
            logger.error(f"Error procesando DOCX: {e}")
            return f"❌ Error al procesar el archivo Word: {str(e)}"

    def _handle_insert_styled_table(
        self,
        doc_obj: Any,
        table_data: Optional[List[List[Any]]],
        theme: Optional[str],
        has_total_row: Optional[bool],
        auto_calculate_totals: Optional[bool],
        col_alignments: Optional[List[str]],
        col_widths: Optional[List[float]],
        cell_padding: Optional[dict],
        filename: str,
    ) -> str:
        if not table_data or len(table_data) < 1:
            return "❌ La acción de tabla requiere 'table_data' con al menos una fila."

        theme_name = (theme or "corporate_blue").lower()
        t_info = PREMIUM_THEMES.get(theme_name, PREMIUM_THEMES["corporate_blue"])

        num_rows = len(table_data)
        num_cols = max(len(r) for r in table_data)

        # Copiar datos para no modificar la lista original
        rows_data = [list(r) + [""] * (num_cols - len(r)) for r in table_data]

        # Verificar si la última fila es de totales
        is_total = bool(has_total_row)
        if not is_total and len(rows_data) > 1:
            first_val = str(rows_data[-1][0]).strip().upper()
            if first_val in ("TOTAL", "TOTALES", "SUMA", "SUM"):
                is_total = True

        if (is_total or auto_calculate_totals) and len(rows_data) > 1:
            last_row = rows_data[-1]
            if not str(last_row[0]).strip():
                last_row[0] = "TOTAL"
            for c_idx in range(1, num_cols):
                val_str = str(last_row[c_idx]).strip()
                if not val_str or val_str.upper() in ("AUTO", "=SUM", "=SUM()"):
                    col_vals = []
                    has_currency = False
                    curr_symbol = "$"
                    for r_idx in range(1, len(rows_data) - 1):
                        cell_raw = str(rows_data[r_idx][c_idx]).strip()
                        if "$" in cell_raw:
                            has_currency = True
                        clean_num = re.sub(r'[^\d.-]', '', cell_raw)
                        if clean_num:
                            try:
                                col_vals.append(float(clean_num))
                            except ValueError:
                                pass
                    if col_vals:
                        total_val = sum(col_vals)
                        if has_currency:
                            last_row[c_idx] = f"{curr_symbol}{total_val:,.2f}"
                        elif any('.' in str(v) for v in col_vals):
                            last_row[c_idx] = f"{total_val:,.2f}"
                        else:
                            last_row[c_idx] = f"{int(total_val):,}"

        table = doc_obj.add_table(rows=num_rows, cols=num_cols)
        table.alignment = 1  # Centrado
        set_table_borders(table, t_info["border_color"])

        # Padding de celda en dxa (1pt = 20 dxa)
        pad_top = int((cell_padding.get('top', 6) if cell_padding else 6) * 20)
        pad_bot = int((cell_padding.get('bottom', 6) if cell_padding else 6) * 20)
        pad_left = int((cell_padding.get('left', 8) if cell_padding else 8) * 20)
        pad_right = int((cell_padding.get('right', 8) if cell_padding else 8) * 20)

        for r_idx, r_data in enumerate(rows_data):
            row = table.rows[r_idx]
            set_row_cant_split(row)

            is_header_row = (r_idx == 0)
            is_last_row_total = (r_idx == num_rows - 1 and is_total)

            if is_header_row:
                set_row_header(row)

            if is_last_row_total:
                set_total_row_borders(row, t_info["border_color"])

            for c_idx, val in enumerate(r_data):
                cell = row.cells[c_idx]
                set_cell_margins(cell, pad_top, pad_bot, pad_left, pad_right)

                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0

                # Alineación
                if col_alignments and c_idx < len(col_alignments):
                    align_str = col_alignments[c_idx].lower()
                    if align_str == 'center':
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align_str == 'right':
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    str_val = str(val).strip()
                    if is_header_row:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif re.match(r'^\$?[\d,]+(\.\d+)?%?$', str_val):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Fondos y texto
                if is_header_row:
                    set_cell_background(cell, t_info["header_bg"])
                    _add_inline_formatted_text(p, str(val), default_font=t_info["font_name"], default_size=Pt(10), default_color=RGBColor(255, 255, 255))
                    for r in p.runs:
                        r.bold = True
                elif is_last_row_total:
                    set_cell_background(cell, t_info["accent_bg"])
                    _add_inline_formatted_text(p, str(val), default_font=t_info["font_name"], default_size=Pt(9.5), default_color=RGBColor(15, 23, 42))
                    for r in p.runs:
                        r.bold = True
                else:
                    bg = t_info["zebra_bg"] if r_idx % 2 == 1 else "FFFFFF"
                    set_cell_background(cell, bg)
                    _add_inline_formatted_text(p, str(val), default_font=t_info["font_name"], default_size=Pt(9.5), default_color=RGBColor(51, 65, 85))

                if col_widths and c_idx < len(col_widths):
                    cell.width = Cm(col_widths[c_idx])

        # Párrafo posterior para espaciado
        p_after = doc_obj.add_paragraph()
        p_after.paragraph_format.space_before = Pt(0)
        p_after.paragraph_format.space_after = Pt(6)

        return f"✅ Tabla premium de {num_rows} fila(s) × {num_cols} columna(s) (tema '{theme_name}') insertada en '{filename}'."

    def _replace_in_docx(self, doc_obj, search: str, replacement: str) -> int:
        """Reemplaza texto en párrafos y tablas, formateando con Markdown en el fallback."""
        count = 0
        all_paragraphs = list(doc_obj.paragraphs)
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        for paragraph in all_paragraphs:
            if search in paragraph.text:
                for run in paragraph.runs:
                    if search in run.text:
                        # Si el reemplazo tiene marcadores markdown, usamos el fallback de formateado
                        if any(marker in replacement for marker in ('**', '*', '__', '_', '`')):
                            break
                        run.text = run.text.replace(search, replacement)
                        count += 1
                        
                # Fallback: si el texto está dividido entre runs o contiene Markdown para formatear
                if search in paragraph.text:
                    full_text = paragraph.text.replace(search, replacement)
                    p_element = paragraph._p
                    runs_to_remove = list(paragraph.runs)
                    for r in runs_to_remove:
                        p_element.remove(r._r)
                    _add_inline_formatted_text(paragraph, full_text)
                    count += 1
        return count

    def _replace_section_in_docx(self, doc_obj, section_start: str, new_content: str) -> bool:
        """Reemplaza el primer párrafo cuyo texto empiece con section_start aplicando inline Markdown."""
        for paragraph in doc_obj.paragraphs:
            if paragraph.text.strip().startswith(section_start):
                p_element = paragraph._p
                runs_to_remove = list(paragraph.runs)
                for r in runs_to_remove:
                    p_element.remove(r._r)
                _add_inline_formatted_text(paragraph, new_content)
                return True
        return False

    def _apply_bold_in_docx(self, doc_obj, search: str) -> int:
        """Divide los runs para poner en negrita solo las ocurrencias del texto."""
        count = 0
        for paragraph in doc_obj.paragraphs:
            if search not in paragraph.text:
                continue
            # Reconstruir el párrafo split por el término buscado
            full_text = paragraph.text
            parts = full_text.split(search)
            if len(parts) < 2:
                continue
            # Guardar el estilo del primer run
            existing_style = {}
            if paragraph.runs:
                r = paragraph.runs[0]
                existing_style = {"font_size": r.font.size, "color": r.font.color.rgb if r.font.color.type else None}
            # Limpiar párrafo
            for run in paragraph.runs:
                run.text = ""
            # Reconstruir con el texto en negrita
            for i, part in enumerate(parts):
                if part:
                    run = paragraph.add_run(part)
                    run.bold = False
                if i < len(parts) - 1:
                    bold_run = paragraph.add_run(search)
                    bold_run.bold = True
                    count += 1
        return count

    def _handle_edit_paragraph(self, doc_obj, paragraph_index: Optional[int], search_text: Optional[str], new_text: Optional[str]) -> str:
        """Edita un párrafo por índice o búsqueda."""
        if new_text is None:
            return "❌ La acción 'edit_paragraph' requiere 'text' con el nuevo contenido."
        
        if paragraph_index is not None:
            if paragraph_index < 0:
                return "❌ El índice del párrafo debe ser >= 0."
            paragraphs = list(doc_obj.paragraphs)
            # También incluir párrafos de tablas
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)
            
            if paragraph_index >= len(paragraphs):
                return f"❌ Índice de párrafo {paragraph_index} fuera de rango (máximo: {len(paragraphs)-1})."
            
            paragraph = paragraphs[paragraph_index]
            p_element = paragraph._p
            runs_to_remove = list(paragraph.runs)
            for r in runs_to_remove:
                p_element.remove(r._r)
            _add_inline_formatted_text(paragraph, new_text)
            return f"✅ Párrafo en posición {paragraph_index} actualizado."
        
        elif search_text:
            paragraphs = list(doc_obj.paragraphs)
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)
            
            for idx, paragraph in enumerate(paragraphs):
                if search_text in paragraph.text:
                    p_element = paragraph._p
                    runs_to_remove = list(paragraph.runs)
                    for r in runs_to_remove:
                        p_element.remove(r._r)
                    _add_inline_formatted_text(paragraph, new_text)
                    return f"✅ Párrafo encontrado con '{search_text}' actualizado."
            return f"⚠️ No se encontró ningún párrafo con '{search_text}'."
        else:
            return "❌ La acción 'edit_paragraph' requiere 'paragraph_index' o 'search_text'."

    def _handle_insert_image(self, doc_obj, image_path: str, image_width: Optional[float], image_height: Optional[float]) -> str:
        """Inserta una imagen en el documento."""
        if not os.path.exists(image_path):
            return f"❌ La imagen '{image_path}' no existe."
        
        try:
            width_cm = image_width if image_width else 10.0
            height_cm = image_height
            
            if height_cm:
                # Tamaño fijo especificado
                run = doc_obj.add_paragraph().add_run()
                run.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))
            else:
                # Ancho fijo, alto automático
                run = doc_obj.add_paragraph().add_run()
                run.add_picture(image_path, width=Cm(width_cm))
            
            return f"✅ Imagen '{os.path.basename(image_path)}' insertada (ancho: {width_cm}cm)."
        except Exception as e:
            logger.error(f"Error insertando imagen: {e}")
            return f"❌ Error al insertar imagen: {str(e)}"

    def _handle_set_page_layout(self, doc_obj, page_margins: Optional[dict], page_orientation: Optional[str]) -> str:
        """Configura márgenes y orientación de página."""
        try:
            sections = doc_obj.sections
            if not sections:
                return "❌ No hay secciones en el documento."
            
            section = sections[0]  # Modificar la primera sección
            
            if page_margins:
                top = Cm(page_margins.get('top', 2.5))
                bottom = Cm(page_margins.get('bottom', 2.5))
                left = Cm(page_margins.get('left', 2.5))
                right = Cm(page_margins.get('right', 2.5))
                section.top_margin = top
                section.bottom_margin = bottom
                section.left_margin = left
                section.right_margin = right
            
            if page_orientation:
                if page_orientation.lower() == 'landscape':
                    section.orientation = WD_ORIENT.LANDSCAPE
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
            
            return "✅ Configuración de página actualizada."
        except Exception as e:
            logger.error(f"Error configurando layout: {e}")
            return f"❌ Error al configurar la página: {str(e)}"

    def _handle_header_footer(self, doc_obj, header_footer_type: str, content: str) -> str:
        """Añade encabezado o pie de página."""
        try:
            if header_footer_type not in ('header', 'footer'):
                return "❌ Tipo debe ser 'header' o 'footer'."
            
            section = doc_obj.sections[0]
            
            if header_footer_type == 'header':
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.text = content
            else:
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.text = content
            
            return f"✅ {header_footer_type.capitalize()} añadido con contenido: '{content}'."
        except Exception as e:
            logger.error(f"Error añadiendo header/footer: {e}")
            return f"❌ Error al añadir {header_footer_type}: {str(e)}"

    def _handle_cell_style(self, doc_obj, cell_coords: str, cell_style: dict) -> str:
        """Aplica estilo a una celda de tabla."""
        try:
            # Buscar la primera tabla del documento
            if not doc_obj.tables:
                return "❌ No hay tablas en el documento."
            
            table = doc_obj.tables[0]
            
            # Parsear coordenadas (ej: "A1", "B2")
            col_letter = ''.join([c for c in cell_coords if c.isalpha()]).upper()
            row_num = int(''.join([c for c in cell_coords if c.isdigit()]))
            
            # Convertir letra a índice (A=0, B=1, etc.)
            col_idx = 0
            for c in col_letter:
                col_idx = col_idx * 26 + (ord(c.upper()) - ord('A') + 1)
            col_idx -= 1
            
            row_idx = row_num - 1  # 1-based a 0-based
            
            if row_idx >= len(table.rows) or col_idx >= len(table.columns):
                return f"❌ Coordenadas {cell_coords} fuera de rango."
            
            cell = table.cell(row_idx, col_idx)
            
            # Aplicar estilos
            if cell_style.get('bold'):
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            
            if cell_style.get('italic'):
                for run in cell.paragraphs[0].runs:
                    run.italic = True
            
            if cell_style.get('font_size'):
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(cell_style['font_size'])
            
            if cell_style.get('font_color'):
                color_hex = cell_style['font_color']
                if color_hex.startswith('#'):
                    color_hex = color_hex[1:]
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
            
            return f"✅ Estilo aplicado a celda {cell_coords}."
        except Exception as e:
            logger.error(f"Error aplicando estilo a celda: {e}")
            return f"❌ Error al aplicar estilo: {str(e)}"

    # ------------------------------------------------------------------ #
    #  XLSX Editing
    # ------------------------------------------------------------------ #
    # ------------------------------------------------------------------ #
    async def _edit_xlsx(
        self,
        file_path: str,
        action: str,
        text: Optional[str],
        cell: Optional[str],
        sheet_name: Optional[str],
        tab_color: Optional[str],
        sheet_index: Optional[int],
        search_text: Optional[str],
        row_data: Optional[List[str]],
        table_data: Optional[List[List[str]]],
        theme: Optional[str],
        has_total_row: Optional[bool],
        total_formulas: Optional[Any],
        column_formats: Optional[dict],
        auto_fit_columns: Optional[bool],
        start_cell: Optional[str],
        range_address: Optional[str],
        range_data: Optional[List[List[str]]],
        format_options: Optional[dict],
        row_index: Optional[int],
        column_index: Optional[int],
        width: Optional[float],
        height: Optional[float],
        filename: str,
    ) -> str:
        if not XLSX_AVAILABLE:
            return "❌ La librería 'openpyxl' no está instalada. No se pueden editar archivos Excel."

        try:
            wb = openpyxl.load_workbook(file_path)

            if action == "xlsx_list_sheets":
                sheets_str = ", ".join(f"'{s}'" for s in wb.sheetnames)
                return f"📋 Pestañas (hojas) disponibles en '{filename}': {sheets_str}"

            if action == "xlsx_create_sheet":
                if not sheet_name:
                    return "❌ La acción 'xlsx_create_sheet' requiere 'sheet_name'."
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    msg = f"⚠️ La pestaña '{sheet_name}' ya existe en '{filename}'."
                else:
                    idx = sheet_index if (sheet_index is not None and 0 <= sheet_index <= len(wb.sheetnames)) else len(wb.sheetnames)
                    ws = wb.create_sheet(title=sheet_name, index=idx)
                    msg = f"✅ Pestaña '{sheet_name}' creada con éxito en '{filename}'."
                if tab_color:
                    clean_color = tab_color.lstrip('#').upper()
                    ws.sheet_properties.tabColor = clean_color
                    msg += f" (Color: #{clean_color})"
                self._backup_file(file_path)
                wb.save(file_path)
                return msg

            elif action == "xlsx_rename_sheet":
                if not sheet_name:
                    return "❌ La acción 'xlsx_rename_sheet' requiere 'sheet_name' (nuevo nombre)."
                target_sheet = search_text if (search_text and search_text in wb.sheetnames) else (sheet_name if sheet_name in wb.sheetnames else wb.active.title)
                if target_sheet not in wb.sheetnames:
                    return f"❌ La pestaña a renombrar '{target_sheet}' no se encuentra en el documento."
                wb[target_sheet].title = sheet_name
                if tab_color:
                    wb[sheet_name].sheet_properties.tabColor = tab_color.lstrip('#').upper()
                msg = f"✅ Pestaña '{target_sheet}' renombrada a '{sheet_name}' en '{filename}'."
                self._backup_file(file_path)
                wb.save(file_path)
                return msg

            elif action == "xlsx_delete_sheet":
                target_del = sheet_name or search_text
                if not target_del or target_del not in wb.sheetnames:
                    return f"❌ La pestaña '{target_del}' no existe en el documento."
                if len(wb.sheetnames) <= 1:
                    return "❌ No se puede eliminar la única pestaña del documento Excel."
                wb.remove(wb[target_del])
                msg = f"✅ Pestaña '{target_del}' eliminada de '{filename}'."
                self._backup_file(file_path)
                wb.save(file_path)
                return msg

            # Seleccionar o auto-crear hoja para el resto de acciones
            if sheet_name:
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                else:
                    ws = wb.create_sheet(title=sheet_name)
                    if tab_color:
                        ws.sheet_properties.tabColor = tab_color.lstrip('#').upper()
            else:
                ws = wb.active
                sheet_name = ws.title

            if action == "xlsx_create_table":
                msg = self._handle_xlsx_create_table(
                    ws, table_data, theme, has_total_row, total_formulas,
                    column_formats, auto_fit_columns, start_cell, filename, sheet_name
                )

            elif action == "xlsx_write_cell":
                if not cell:
                    return "❌ La acción 'xlsx_write_cell' requiere el campo 'cell' (ej: 'A1')."
                if text is None:
                    return "❌ La acción 'xlsx_write_cell' requiere 'text' con el valor a escribir."
                cell_ref = ws[cell.upper()]
                str_text = str(text).strip()
                if str_text.startswith("="):
                    # Guardar como fórmula evaluable
                    cell_ref.value = str_text
                else:
                    # Intentar parsear como número si aplica
                    try:
                        if "." in str_text:
                            cell_ref.value = float(str_text)
                        else:
                            cell_ref.value = int(str_text)
                    except ValueError:
                        cell_ref.value = text
                msg = f"✅ Celda '{cell.upper()}' de la hoja '{sheet_name}' en '{filename}' actualizada con: {text}"

            elif action == "xlsx_append_row":
                if not row_data:
                    return "❌ La acción 'xlsx_append_row' requiere 'row_data' (lista de valores)."
                parsed_row = []
                for val in row_data:
                    s_val = str(val).strip()
                    if s_val.startswith("="):
                        parsed_row.append(s_val)
                    else:
                        try:
                            parsed_row.append(float(s_val) if "." in s_val else int(s_val))
                        except ValueError:
                            parsed_row.append(val)
                ws.append(parsed_row)
                msg = f"✅ Fila añadida a la hoja '{sheet_name}' de '{filename}': {row_data}"

            elif action == "xlsx_write_range":
                if not range_address:
                    return "❌ La acción 'xlsx_write_range' requiere 'range_address' (ej: 'A1:C3')."
                if not range_data:
                    return "❌ La acción 'xlsx_write_range' requiere 'range_data' (matriz 2D)."
                msg = self._xlsx_write_range(ws, range_address, range_data, filename, sheet_name)

            elif action == "xlsx_format_cells":
                if not range_address:
                    return "❌ La acción 'xlsx_format_cells' requiere 'range_address' (ej: 'A1:C3')."
                if not format_options:
                    return "❌ La acción 'xlsx_format_cells' requiere 'format_options'."
                msg = self._xlsx_format_cells(ws, range_address, format_options, filename, sheet_name)

            elif action == "xlsx_insert_row":
                if row_index is None:
                    return "❌ La acción 'xlsx_insert_row' requiere 'row_index'."
                msg = self._xlsx_insert_row(ws, row_index, filename, sheet_name)

            elif action == "xlsx_insert_column":
                if column_index is None:
                    return "❌ La acción 'xlsx_insert_column' requiere 'column_index'."
                msg = self._xlsx_insert_column(ws, column_index, filename, sheet_name)

            elif action == "xlsx_delete_row":
                if row_index is None:
                    return "❌ La acción 'xlsx_delete_row' requiere 'row_index'."
                msg = self._xlsx_delete_row(ws, row_index, filename, sheet_name)

            elif action == "xlsx_delete_column":
                if column_index is None:
                    return "❌ La acción 'xlsx_delete_column' requiere 'column_index'."
                msg = self._xlsx_delete_column(ws, column_index, filename, sheet_name)

            elif action == "xlsx_merge_cells":
                if not range_address:
                    return "❌ La acción 'xlsx_merge_cells' requiere 'range_address' (ej: 'A1:C3')."
                msg = self._xlsx_merge_cells(ws, range_address, filename, sheet_name)

            elif action == "xlsx_set_column_width":
                if column_index is None or width is None:
                    return "❌ La acción 'xlsx_set_column_width' requiere 'column_index' y 'width'."
                msg = self._xlsx_set_column_width(ws, column_index, width, filename, sheet_name)

            elif action == "xlsx_set_row_height":
                if row_index is None or height is None:
                    return "❌ La acción 'xlsx_set_row_height' requiere 'row_index' y 'height'."
                msg = self._xlsx_set_row_height(ws, row_index, height, filename, sheet_name)

            elif action == "xlsx_copy_sheet":
                if not sheet_name:
                    return "❌ La acción 'xlsx_copy_sheet' requiere 'sheet_name' (nombre de la hoja origen)."
                if not text:
                    return "❌ La acción 'xlsx_copy_sheet' requiere 'text' con el nuevo nombre para la copia."
                msg = self._xlsx_copy_sheet(wb, sheet_name, text, sheet_index, filename)

            elif action == "xlsx_move_sheet":
                if not sheet_name:
                    return "❌ La acción 'xlsx_move_sheet' requiere 'sheet_name'."
                if sheet_index is None:
                    return "❌ La acción 'xlsx_move_sheet' requiere 'sheet_index' (posición 0-based destino)."
                msg = self._xlsx_move_sheet(wb, sheet_name, sheet_index, filename)

            elif action == "xlsx_clear_sheet":
                target_clear = sheet_name or (wb.active.title if wb.active else None)
                if not target_clear or target_clear not in wb.sheetnames:
                    return f"❌ La hoja '{target_clear}' no existe en el documento."
                ws_clear = wb[target_clear]
                msg = self._xlsx_clear_sheet(ws_clear, target_clear, filename)

            elif action == "xlsx_get_sheet_info":
                target_info = sheet_name or (wb.active.title if wb.active else None)
                if not target_info or target_info not in wb.sheetnames:
                    return f"❌ La hoja '{target_info}' no existe en el documento."
                ws_info = wb[target_info]
                return self._xlsx_get_sheet_info(ws_info, target_info, filename)

            elif action == "xlsx_freeze_panes":
                if not cell:
                    return "❌ La acción 'xlsx_freeze_panes' requiere 'cell' (ej: 'B2' para congelar fila 1 y columna A)."
                msg = self._xlsx_freeze_panes(ws, cell, filename, sheet_name)

            elif action == "xlsx_protect_sheet":
                if not sheet_name:
                    return "❌ La acción 'xlsx_protect_sheet' requiere 'sheet_name'."
                password = text  # text se usa como contraseña (opcional)
                msg = self._xlsx_protect_sheet(ws, password, filename, sheet_name)

            elif action == "xlsx_set_sheet_tab_color":
                if not sheet_name:
                    return "❌ La acción 'xlsx_set_sheet_tab_color' requiere 'sheet_name'."
                if not tab_color:
                    return "❌ La acción 'xlsx_set_sheet_tab_color' requiere 'tab_color' (código hex, ej: 'FF0000')."
                msg = self._xlsx_set_sheet_tab_color(ws, tab_color, filename, sheet_name)

            else:
                return (
                    f"❌ Acción '{action}' no válida para .xlsx. "
                    f"Acciones disponibles: {', '.join(SUPPORTED_ACTIONS)}"
                )

            self._backup_file(file_path)
            wb.save(file_path)
            return msg

        except Exception as e:
            logger.error(f"Error procesando XLSX: {e}")
            return f"❌ Error al procesar el archivo Excel: {str(e)}"

    # ------------------------------------------------------------------ #
    #  XLSX Helper Methods
    # ------------------------------------------------------------------ #
    def _handle_xlsx_create_table(
        self,
        ws: Any,
        table_data: Optional[List[List[Any]]],
        theme: Optional[str],
        has_total_row: Optional[bool],
        total_formulas: Optional[Any],
        column_formats: Optional[dict],
        auto_fit_columns: Optional[bool],
        start_cell: Optional[str],
        filename: str,
        sheet_name: str,
    ) -> str:
        if not table_data or len(table_data) < 1:
            return "❌ La acción 'xlsx_create_table' requiere 'table_data' con al menos una fila."

        theme_name = (theme or "corporate_blue").lower()
        t_info = PREMIUM_THEMES.get(theme_name, PREMIUM_THEMES["corporate_blue"])

        from openpyxl.utils import coordinate_to_tuple, get_column_letter

        start_address = (start_cell or "A1").upper()
        start_row, start_col = coordinate_to_tuple(start_address)

        num_rows = len(table_data)
        num_cols = max(len(r) for r in table_data)

        rows_data = [list(r) + [""] * (num_cols - len(r)) for r in table_data]

        is_total = bool(has_total_row)
        if not is_total and len(rows_data) > 1:
            first_val = str(rows_data[-1][0]).strip().upper()
            if first_val in ("TOTAL", "TOTALES", "SUMA", "SUM"):
                is_total = True

        data_start_row = start_row + 1
        data_end_row = start_row + num_rows - (1 if is_total else 0) - 1

        header_fill = PatternFill(start_color=t_info["header_bg"], end_color=t_info["header_bg"], fill_type="solid")
        header_font = Font(name=t_info["font_name"], size=11, bold=True, color=t_info["header_fg"])

        zebra_fill = PatternFill(start_color=t_info["zebra_bg"], end_color=t_info["zebra_bg"], fill_type="solid")
        accent_fill = PatternFill(start_color=t_info["accent_bg"], end_color=t_info["accent_bg"], fill_type="solid")

        thin_side = Side(border_style="thin", color=t_info["border_color"])
        double_side = Side(border_style="double", color=t_info["border_color"])

        normal_border = Border(bottom=thin_side)
        header_border = Border(top=thin_side, bottom=thin_side, left=thin_side, right=thin_side)
        total_border = Border(top=thin_side, bottom=double_side)

        for r_offset, r_data in enumerate(rows_data):
            curr_row = start_row + r_offset
            is_header = (r_offset == 0)
            is_last_total = (r_offset == num_rows - 1 and is_total)

            for c_offset, val in enumerate(r_data):
                curr_col = start_col + c_offset
                col_letter = get_column_letter(curr_col)
                cell = ws.cell(row=curr_row, column=curr_col)

                str_val = str(val).strip()

                if is_last_total and c_offset > 0 and (not str_val or str_val.upper() in ("AUTO", "=SUM", "=SUM()")):
                    formula_op = "SUM"
                    if isinstance(total_formulas, dict):
                        formula_op = total_formulas.get(col_letter, total_formulas.get(str(c_offset), "SUM"))
                    elif isinstance(total_formulas, list) and c_offset < len(total_formulas):
                        f_item = str(total_formulas[c_offset]).strip()
                        if f_item:
                            formula_op = f_item

                    if str(formula_op).startswith("="):
                        cell.value = formula_op
                    else:
                        cell.value = f"={formula_op}({col_letter}{data_start_row}:{col_letter}{data_end_row})"
                elif str_val.startswith("="):
                    cell.value = str_val
                else:
                    try:
                        if "." in str_val:
                            cell.value = float(str_val)
                        else:
                            cell.value = int(str_val)
                    except ValueError:
                        cell.value = val

                if column_formats and not is_header:
                    fmt = column_formats.get(col_letter) or column_formats.get(str(c_offset))
                    if fmt:
                        cell.number_format = fmt

                if is_header:
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                    cell.border = header_border
                elif is_last_total:
                    cell.fill = accent_fill
                    cell.font = Font(name=t_info["font_name"], size=10.5, bold=True, color="0F172A")
                    cell.border = total_border
                    if isinstance(cell.value, (int, float)) or str(cell.value or "").startswith("="):
                        cell.alignment = Alignment(horizontal="right", vertical="center")
                    else:
                        cell.alignment = Alignment(horizontal="left", vertical="center")
                else:
                    if r_offset % 2 == 1:
                        cell.fill = zebra_fill
                    cell.font = Font(name=t_info["font_name"], size=10)
                    cell.border = normal_border
                    if isinstance(cell.value, (int, float)) or str(cell.value or "").startswith("="):
                        cell.alignment = Alignment(horizontal="right", vertical="center")
                    else:
                        cell.alignment = Alignment(horizontal="left", vertical="center")

        if auto_fit_columns is not False:
            for c_offset in range(num_cols):
                curr_col = start_col + c_offset
                col_letter = get_column_letter(curr_col)
                max_len = 0
                for r_offset in range(num_rows):
                    cell_val = ws.cell(row=start_row + r_offset, column=curr_col).value
                    v_str = str(cell_val or '')
                    if len(v_str) > max_len:
                        max_len = len(v_str)
                ws.column_dimensions[col_letter].width = max(max_len + 4, 12)

        return f"✅ Tabla Excel premium de {num_rows} fila(s) × {num_cols} columna(s) (tema '{theme_name}') creada en '{start_address}' de hoja '{sheet_name}'."

    def _xlsx_write_range(self, ws, range_address: str, range_data: List[List[str]], filename: str, sheet_name: str) -> str:
        """Escribe datos en un rango de celdas."""
        try:
            from openpyxl.utils import coordinate_to_tuple
            start_row, start_col = coordinate_to_tuple(range_address.split(':')[0])
            
            for i, row in enumerate(range_data):
                for j, value in enumerate(row):
                    val_str = str(value).strip()
                    cell = ws.cell(row=start_row + i, column=start_col + j)
                    if val_str.startswith("="):
                        cell.value = val_str
                    else:
                        try:
                            cell.value = float(val_str) if "." in val_str else int(val_str)
                        except ValueError:
                            cell.value = value
            
            return f"✅ Rango '{range_address}' en hoja '{sheet_name}' actualizado."
        except Exception as e:
            logger.error(f"Error escribiendo rango: {e}")
            return f"❌ Error al escribir en rango: {str(e)}"

    def _xlsx_format_cells(self, ws, range_address: str, format_options: dict, filename: str, sheet_name: str) -> str:
        """Formatea celdas con estilos avanzados."""
        try:
            if ':' in range_address:
                range_ref = [cell for row in ws[range_address] for cell in row]
            else:
                range_ref = [ws[range_address]]
            
            for cell in range_ref:
                # Font
                font_kwargs = {}
                if 'bold' in format_options:
                    font_kwargs['bold'] = format_options.get('bold')
                if 'italic' in format_options:
                    font_kwargs['italic'] = format_options.get('italic')
                if format_options.get('font_size'):
                    font_kwargs['size'] = format_options['font_size']
                if format_options.get('font_color'):
                    font_kwargs['color'] = str(format_options['font_color']).lstrip('#')
                if format_options.get('font_name'):
                    font_kwargs['name'] = format_options['font_name']
                
                if font_kwargs or getattr(cell, 'font', None):
                    if cell.font:
                        for attr in ['name', 'size', 'bold', 'italic', 'color', 'underline']:
                            if attr not in font_kwargs and getattr(cell.font, attr) is not None:
                                font_kwargs[attr] = getattr(cell.font, attr)
                    cell.font = Font(**font_kwargs)
                
                # Fill
                if format_options.get('fill_color'):
                    color = str(format_options['fill_color']).lstrip('#')
                    fill = PatternFill(start_color=color, end_color=color, fill_type='solid')
                    cell.fill = fill
                
                # Alignment
                align_kwargs = {}
                if format_options.get('align'):
                    align_kwargs['horizontal'] = format_options['align'].lower()
                if format_options.get('valign'):
                    align_kwargs['vertical'] = format_options['valign'].lower()
                if 'wrap_text' in format_options:
                    align_kwargs['wrap_text'] = format_options.get('wrap_text')
                
                if align_kwargs or getattr(cell, 'alignment', None):
                    if cell.alignment:
                        for attr in ['horizontal', 'vertical', 'wrap_text']:
                            if attr not in align_kwargs and getattr(cell.alignment, attr) is not None:
                                align_kwargs[attr] = getattr(cell.alignment, attr)
                    cell.alignment = Alignment(**align_kwargs)
                
                # Borders
                if format_options.get('border'):
                    border_style = format_options['border']
                    side = Side(border_style=border_style, color='000000')
                    cell.border = Border(left=side, right=side, top=side, bottom=side)
                
                # Number Format
                if format_options.get('number_format'):
                    cell.number_format = format_options['number_format']
            
            return f"✅ Formato aplicado al rango '{range_address}' en hoja '{sheet_name}'."
        except Exception as e:
            logger.error(f"Error formateando celdas: {e}")
            return f"❌ Error al formatear celdas: {str(e)}"

    def _xlsx_insert_row(self, ws, row_index: int, filename: str, sheet_name: str) -> str:
        """Inserta una fila en la posición especificada."""
        try:
            ws.insert_rows(row_index)
            return f"✅ Fila insertada en posición {row_index} en hoja '{sheet_name}'."
        except Exception as e:
            logger.error(f"Error insertando fila: {e}")
            return f"❌ Error al insertar fila: {str(e)}"

    def _xlsx_insert_column(self, ws, column_index: int, filename: str, sheet_name: str) -> str:
        """Inserta una columna en la posición especificada."""
        try:
            ws.insert_cols(column_index)
            return f"✅ Columna insertada en posición {column_index} en hoja '{sheet_name}'."
        except Exception as e:
            logger.error(f"Error insertando columna: {e}")
            return f"❌ Error al insertar columna: {str(e)}"

    def _xlsx_delete_row(self, ws, row_index: int, filename: str, sheet_name: str) -> str:
        """Elimina una fila en la posición especificada."""
        try:
            ws.delete_rows(row_index)
            return f"✅ Fila eliminada en posición {row_index} en hoja '{sheet_name}'."
        except Exception as e:
            logger.error(f"Error eliminando fila: {e}")
            return f"❌ Error al eliminar fila: {str(e)}"

    def _xlsx_delete_column(self, ws, column_index: int, filename: str, sheet_name: str) -> str:
        """Elimina una columna en la posición especificada."""
        try:
            ws.delete_cols(column_index)
            return f"✅ Columna eliminada en posición {column_index} en hoja '{sheet_name}'."
        except Exception as e:
            logger.error(f"Error eliminando columna: {e}")
            return f"❌ Error al eliminar columna: {str(e)}"

    def _xlsx_merge_cells(self, ws, range_address: str, filename: str, sheet_name: str) -> str:
        """Fusiona celdas en un rango."""
        try:
            ws.merge_cells(range_address)
            return f"✅ Celdas fusionadas en rango '{range_address}' de hoja '{sheet_name}'."
        except Exception as e:
            logger.error(f"Error fusionando celdas: {e}")
            return f"❌ Error al fusionar celdas: {str(e)}"

    def _xlsx_set_column_width(self, ws, column_index: int, width: float, filename: str, sheet_name: str) -> str:
        """Establece el ancho de una columna (en caracteres, no cm)."""
        try:
            col_letter = get_column_letter(column_index + 1)
            ws.column_dimensions[col_letter].width = width
            return f"✅ Ancho de columna {col_letter} establecido en {width} en hoja '{sheet_name}'."
        except Exception as e:
            logger.error(f"Error estableciendo ancho de columna: {e}")
            return f"❌ Error al establecer ancho de columna: {str(e)}"

    def _xlsx_set_row_height(self, ws, row_index: int, height: float, filename: str, sheet_name: str) -> str:
        """Establece la altura de una fila."""
        try:
            ws.row_dimensions[row_index].height = height
            return f"✅ Altura de fila {row_index} establecida en {height} en hoja '{sheet_name}'."
        except Exception as e:
            logger.error(f"Error estableciendo altura de fila: {e}")
            return f"❌ Error al establecer altura de fila: {str(e)}"

    def _xlsx_copy_sheet(self, wb, source_name: str, new_name: str, target_index: Optional[int], filename: str) -> str:
        """Copia una hoja del libro a un nuevo nombre."""
        try:
            if source_name not in wb.sheetnames:
                return f"❌ La hoja origen '{source_name}' no existe en '{filename}'."
            if new_name in wb.sheetnames:
                return f"❌ Ya existe una hoja con el nombre '{new_name}' en '{filename}'."
            ws_copy = wb.copy_worksheet(wb[source_name])
            ws_copy.title = new_name
            # Mover a posición si se especificó
            if target_index is not None:
                total = len(wb.sheetnames)
                idx = max(0, min(target_index, total - 1))
                wb.move_sheet(new_name, offset=idx - wb.sheetnames.index(new_name))
            return f"✅ Hoja '{source_name}' copiada como '{new_name}' en '{filename}'."
        except Exception as e:
            logger.error(f"Error copiando hoja: {e}")
            return f"❌ Error al copiar la hoja: {str(e)}"

    def _xlsx_move_sheet(self, wb, sheet_name: str, target_index: int, filename: str) -> str:
        """Mueve una hoja a la posición indicada (0-based)."""
        try:
            if sheet_name not in wb.sheetnames:
                return f"❌ La hoja '{sheet_name}' no existe en '{filename}'."
            total = len(wb.sheetnames)
            idx = max(0, min(target_index, total - 1))
            current_idx = wb.sheetnames.index(sheet_name)
            offset = idx - current_idx
            if offset != 0:
                wb.move_sheet(sheet_name, offset=offset)
            return f"✅ Hoja '{sheet_name}' movida a posición {idx} en '{filename}'."
        except Exception as e:
            logger.error(f"Error moviendo hoja: {e}")
            return f"❌ Error al mover la hoja: {str(e)}"

    def _xlsx_clear_sheet(self, ws, sheet_name: str, filename: str) -> str:
        """Borra todo el contenido de una hoja (celdas, imágenes, tablas) sin eliminarla."""
        try:
            # Eliminar contenido de todas las celdas
            for row in ws.iter_rows():
                for cell in row:
                    cell.value = None
                    cell.font = None
                    cell.fill = None
                    cell.border = None
                    cell.alignment = None
                    cell.number_format = 'General'
            # Eliminar merge de celdas
            for merge_range in list(ws.merged_cells.ranges):
                ws.unmerge_cells(str(merge_range))
            # Reiniciar dimensiones de columnas/filas
            ws.column_dimensions.clear()
            ws.row_dimensions.clear()
            # Eliminar tablas definidas en la hoja
            for tbl in list(ws.tables.values()):
                del ws.tables[tbl.name]
            return f"✅ Contenido de la hoja '{sheet_name}' borrado completamente en '{filename}'."
        except Exception as e:
            logger.error(f"Error limpiando hoja: {e}")
            return f"❌ Error al limpiar la hoja: {str(e)}"

    def _xlsx_get_sheet_info(self, ws, sheet_name: str, filename: str) -> str:
        """Devuelve información detallada sobre una hoja: dimensiones, rango usado, muestra de datos."""
        try:
            max_row = ws.max_row or 0
            max_col = ws.max_column or 0
            min_row = ws.min_row or 1
            min_col = ws.min_column or 1

            # Contar celdas con datos
            cells_with_data = sum(
                1 for row in ws.iter_rows()
                for cell in row
                if cell.value is not None
            )

            # Rango usado
            used_range = f"{get_column_letter(min_col)}{min_row}:{get_column_letter(max_col)}{max_row}" if max_row and max_col else "(vacía)"

            # Muestra de datos (primeras 3 filas, máx 5 columnas)
            preview_lines = []
            for r in ws.iter_rows(min_row=1, max_row=min(3, max_row), max_col=min(5, max_col), values_only=True):
                row_str = " | ".join(str(v) if v is not None else "" for v in r)
                preview_lines.append(f"  {row_str}")
            preview = "\n".join(preview_lines) if preview_lines else "  (sin datos)"

            # Pestañas del libro
            book = ws.parent
            sheets_list = ", ".join(f"'{s}'" for s in book.sheetnames)

            result = (
                f"📊 **Información de la hoja '{sheet_name}'** en '{filename}':\n"
                f"  - Filas usadas: {max_row}\n"
                f"  - Columnas usadas: {max_col}\n"
                f"  - Rango de datos: {used_range}\n"
                f"  - Celdas con datos: {cells_with_data}\n"
                f"  - Celdas fusionadas: {len(list(ws.merged_cells.ranges))}\n"
                f"  - Paneles congelados: {'Sí' if ws.freeze_panes else 'No'}\n"
                f"  - Color de pestaña: #{ws.sheet_properties.tabColor.rgb if ws.sheet_properties.tabColor else 'N/A'}\n"
                f"  - Todas las hojas del libro: {sheets_list}\n"
                f"  Muestra de datos (primeras 3 filas):\n{preview}"
            )
            return result
        except Exception as e:
            logger.error(f"Error obteniendo info de hoja: {e}")
            return f"❌ Error al obtener información de la hoja: {str(e)}"

    def _xlsx_freeze_panes(self, ws, cell: str, filename: str, sheet_name: str) -> str:
        """Congela filas/columnas. Ej: cell='B2' congela la fila 1 y la columna A."""
        try:
            cell_upper = cell.upper()
            ws.freeze_panes = cell_upper
            # Calcular qué se congeló para el mensaje
            from openpyxl.utils import coordinate_to_tuple
            row_num, col_num = coordinate_to_tuple(cell_upper)
            frozen_rows = row_num - 1
            frozen_cols = col_num - 1
            parts = []
            if frozen_rows > 0:
                parts.append(f"{frozen_rows} fila(s) superior(es)")
            if frozen_cols > 0:
                parts.append(f"{frozen_cols} columna(s) izquierda(s)")
            desc = " y ".join(parts) if parts else "ninguna fila/columna"
            return f"✅ Paneles congelados en hoja '{sheet_name}': {desc} (desde celda {cell_upper})."
        except Exception as e:
            logger.error(f"Error congelando paneles: {e}")
            return f"❌ Error al congelar paneles: {str(e)}"

    def _xlsx_protect_sheet(self, ws, password: Optional[str], filename: str, sheet_name: str) -> str:
        """Protege la hoja con contraseña para evitar edición manual en OnlyOffice/Excel."""
        try:
            ws.protection.sheet = True
            ws.protection.enable()
            if password:
                ws.protection.set_password(password)
                return f"✅ Hoja '{sheet_name}' protegida con contraseña en '{filename}'."
            else:
                return f"✅ Hoja '{sheet_name}' protegida (sin contraseña) en '{filename}'."
        except Exception as e:
            logger.error(f"Error protegiendo hoja: {e}")
            return f"❌ Error al proteger la hoja: {str(e)}"

    def _xlsx_set_sheet_tab_color(self, ws, tab_color: str, filename: str, sheet_name: str) -> str:
        """Cambia el color de la pestaña de una hoja existente."""
        try:
            clean_color = tab_color.lstrip('#').upper()
            # Validar longitud hex
            if len(clean_color) not in (6, 8):
                return f"❌ Color inválido '{tab_color}'. Usa formato hex de 6 dígitos (ej: 'FF0000' para rojo)."
            ws.sheet_properties.tabColor = clean_color
            return f"✅ Color de pestaña de '{sheet_name}' cambiado a #{clean_color} en '{filename}'."
        except Exception as e:
            logger.error(f"Error cambiando color de pestaña: {e}")
            return f"❌ Error al cambiar el color de la pestaña: {str(e)}"

    # ------------------------------------------------------------------ #
    #  Utilidades
    # ------------------------------------------------------------------ #
    def _backup_file(self, file_path: str) -> str:
        """Crea un respaldo del archivo antes de modificarlo."""
        backups_dir = os.path.join(os.path.dirname(file_path), ".backups")
        os.makedirs(backups_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = os.path.join(
            backups_dir,
            f"{os.path.basename(file_path)}.{timestamp}.ai_edit.bak"
        )
        shutil.copy2(file_path, backup_path)
        return backup_path

    def _run(self, *args: Any, **kwargs: Any) -> Any:
        raise NotImplementedError("Esta herramienta solo soporta ejecución asíncrona.")
