"""
cli/modules/doc_editor.py
Editor de documentos Word/PDF moderno para el CLI de KognitoAI.
Permite crear, editar y exportar documentos con estilos corporativos.
"""
from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List, Optional, Tuple

# ── Word (python-docx) ────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── PDF (reportlab) ───────────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4, LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# ── WeasyPrint (Markdown → PDF) ───────────────────────────────────────────────
try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


# ── Corporate Palette ─────────────────────────────────────────────────────────
CORP_PRIMARY = RGBColor(0x0A, 0x0E, 0x27) if DOCX_AVAILABLE else None  # Deep Navy
CORP_ACCENT = RGBColor(0x6C, 0x63, 0xFF) if DOCX_AVAILABLE else None   # Purple
CORP_TEXT = RGBColor(0x1A, 0x1A, 0x2E) if DOCX_AVAILABLE else None     # Near-black
CORP_LIGHT = RGBColor(0xF5, 0xF5, 0xF7) if DOCX_AVAILABLE else None    # Off-white


# ── Markdown → Sections parser ────────────────────────────────────────────────

def parse_markdown_to_sections(text: str) -> List[Tuple[str, str]]:
    """
    Convierte Markdown básico en una lista de (tipo, contenido):
    tipos: 'h1', 'h2', 'h3', 'bold', 'bullet', 'paragraph', 'hr', 'code'
    """
    sections: List[Tuple[str, str]] = []
    for line in text.split("\n"):
        stripped = line.rstrip()
        if stripped.startswith("# "):
            sections.append(("h1", stripped[2:]))
        elif stripped.startswith("## "):
            sections.append(("h2", stripped[3:]))
        elif stripped.startswith("### "):
            sections.append(("h3", stripped[4:]))
        elif stripped.startswith("---") or stripped.startswith("___"):
            sections.append(("hr", ""))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            sections.append(("bullet", stripped[2:]))
        elif stripped.startswith("    ") or stripped.startswith("\t"):
            sections.append(("code", stripped.strip()))
        elif stripped.startswith("**") and stripped.endswith("**"):
            sections.append(("bold", stripped[2:-2]))
        elif stripped:
            sections.append(("paragraph", stripped))
        else:
            sections.append(("blank", ""))
    return sections


# ── Word Document ─────────────────────────────────────────────────────────────

def _apply_corporate_word_styles(doc: "Document") -> None:
    """Aplica fuentes, colores y espaciado corporativo al documento Word."""
    if not DOCX_AVAILABLE:
        return

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = CORP_TEXT

    for heading_name, size, bold in [
        ("Heading 1", 28, True),
        ("Heading 2", 20, True),
        ("Heading 3", 15, True),
    ]:
        try:
            h = doc.styles[heading_name]
            h.font.name = "Calibri"
            h.font.size = Pt(size)
            h.font.bold = bold
            h.font.color.rgb = CORP_ACCENT
        except KeyError:
            pass


def _add_header_footer(doc: "Document", title: str, author: str = "KognitoAI") -> None:
    """Agrega encabezado y pie de página corporativos."""
    if not DOCX_AVAILABLE:
        return
    section = doc.sections[0]
    header = section.header
    if not header.paragraphs:
        header.add_paragraph()
    hp = header.paragraphs[0]
    hp.clear()
    run = hp.add_run(f"  {title}  •  {author}")
    run.font.size = Pt(9)
    run.font.color.rgb = CORP_ACCENT
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    fp = footer.paragraphs[0]
    fp.clear()
    fp.add_run("KognitoAI  |  Documento Confidencial").font.size = Pt(8)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_to_word(
    content: str,
    output_path: str,
    title: str = "Documento",
    author: str = "KognitoAI CLI",
    add_cover: bool = True,
) -> str:
    """
    Genera un archivo .docx corporativo desde texto Markdown.
    Retorna la ruta del archivo generado.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx no está instalado. Ejecuta: pip install python-docx")

    doc = Document()
    _apply_corporate_word_styles(doc)
    _add_header_footer(doc, title, author)

    # Márgenes
    section = doc.sections[0]
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Portada
    if add_cover:
        cover_para = doc.add_paragraph()
        cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n\n\n")
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title_para.add_run(title)
        tr.font.size = Pt(32)
        tr.font.bold = True
        tr.font.color.rgb = CORP_ACCENT
        doc.add_paragraph("\n")
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = author_para.add_run(author)
        ar.font.size = Pt(12)
        ar.font.color.rgb = CORP_TEXT
        doc.add_page_break()

    # Contenido
    sections = parse_markdown_to_sections(content)
    for kind, text in sections:
        if kind == "h1":
            p = doc.add_heading(text, level=1)
        elif kind == "h2":
            p = doc.add_heading(text, level=2)
        elif kind == "h3":
            p = doc.add_heading(text, level=3)
        elif kind == "bullet":
            doc.add_paragraph(f"• {text}", style="List Bullet")
        elif kind == "hr":
            doc.add_paragraph("─" * 60)
        elif kind == "code":
            p = doc.add_paragraph(text)
            p.style.font.name = "Courier New"
            p.style.font.size = Pt(10)
        elif kind == "bold":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        elif kind == "paragraph":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif kind == "blank":
            doc.add_paragraph("")

    doc.save(output_path)
    return output_path


# ── PDF Document ──────────────────────────────────────────────────────────────

def export_to_pdf(
    content: str,
    output_path: str,
    title: str = "Documento",
    author: str = "KognitoAI CLI",
) -> str:
    """
    Genera un archivo PDF corporativo desde texto Markdown.
    Retorna la ruta del archivo generado.
    """
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab no está instalado. Ejecuta: pip install reportlab")

    # Estilos
    styles = getSampleStyleSheet()
    corp_primary_color = colors.HexColor("#0A0E27")
    corp_accent_color = colors.HexColor("#6C63FF")
    corp_gray = colors.HexColor("#6B7280")

    title_style = ParagraphStyle(
        "CorporateTitle",
        parent=styles["Title"],
        fontSize=28,
        fontName="Helvetica-Bold",
        textColor=corp_accent_color,
        spaceAfter=12,
        spaceBefore=6,
    )
    h1_style = ParagraphStyle(
        "CorporateH1",
        parent=styles["Heading1"],
        fontSize=20,
        fontName="Helvetica-Bold",
        textColor=corp_accent_color,
        spaceBefore=16,
        spaceAfter=8,
        borderPad=4,
    )
    h2_style = ParagraphStyle(
        "CorporateH2",
        parent=styles["Heading2"],
        fontSize=15,
        fontName="Helvetica-Bold",
        textColor=corp_primary_color,
        spaceBefore=12,
        spaceAfter=6,
    )
    h3_style = ParagraphStyle(
        "CorporateH3",
        parent=styles["Heading3"],
        fontSize=12,
        fontName="Helvetica-Bold",
        textColor=corp_primary_color,
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "CorporateBody",
        parent=styles["BodyText"],
        fontSize=11,
        fontName="Helvetica",
        leading=16,
        textColor=corp_primary_color,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    bullet_style = ParagraphStyle(
        "CorporateBullet",
        parent=body_style,
        bulletText="•",
        leftIndent=18,
        spaceAfter=4,
    )
    code_style = ParagraphStyle(
        "CorporateCode",
        parent=styles["Code"],
        fontSize=9,
        fontName="Courier",
        backColor=colors.HexColor("#F3F4F6"),
        leftIndent=12,
        rightIndent=12,
        spaceBefore=4,
        spaceAfter=4,
    )

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=title,
        author=author,
    )

    story = []

    # Cover section
    story.append(Spacer(1, 3 * cm))
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"<font color='#{corp_gray.hexval()[2:]}'>por {author}</font>", body_style))
    story.append(HRFlowable(width="100%", thickness=2, color=corp_accent_color, spaceAfter=0.5*cm))
    story.append(Spacer(1, 1 * cm))

    # Content
    sections = parse_markdown_to_sections(content)
    for kind, text in sections:
        # Sanitize text for reportlab
        safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if kind == "h1":
            story.append(Paragraph(safe_text, h1_style))
            story.append(HRFlowable(width="100%", thickness=1, color=corp_accent_color))
        elif kind == "h2":
            story.append(Paragraph(safe_text, h2_style))
        elif kind == "h3":
            story.append(Paragraph(safe_text, h3_style))
        elif kind == "bullet":
            story.append(Paragraph(f"• {safe_text}", bullet_style))
        elif kind == "code":
            story.append(Paragraph(safe_text, code_style))
        elif kind == "bold":
            story.append(Paragraph(f"<b>{safe_text}</b>", body_style))
        elif kind == "paragraph":
            story.append(Paragraph(safe_text, body_style))
        elif kind == "blank":
            story.append(Spacer(1, 0.3 * cm))
        elif kind == "hr":
            story.append(HRFlowable(width="100%", thickness=1, color=corp_gray))

    doc.build(story)
    return output_path


# ── Code export ───────────────────────────────────────────────────────────────

def export_code_to_pdf(
    code: str,
    language: str,
    output_path: str,
    title: str = "Código Fuente",
) -> str:
    """Exporta un bloque de código a PDF con resaltado básico de sintaxis."""
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab no está instalado.")

    content = f"# {title}\n\n    {chr(10) + '    '.join(code.split(chr(10)))}"
    return export_to_pdf(content, output_path, title=title, author="KognitoAI CLI")
