# utils/deep_research_word_saver.py

import os
import re
import uuid
import logging
from datetime import datetime
from typing import Optional
from sqlalchemy import select

from core.database import SessionLocal, Document, DocumentFolder
from core.onlyoffice_storage import (
    build_onlyoffice_relative_path,
    ensure_onlyoffice_account_dir,
)
from utils.db_session import DBSession

logger = logging.getLogger(__name__)

# Try to import python-docx and its formatting tools
try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
except ImportError as e:
    DOCX_AVAILABLE = False
    logger.error(f"python-docx is not available: {e}")

# Regex to find inline formatting markers in markdown
inline_regex = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|__.*?__|__.*?__|\*.*?\*|_.*?_|`.*?`)')

def _add_inline_formatted_text(paragraph, text: str, default_font: str = "Segoe UI", default_size = Pt(10.5), default_color = RGBColor(51, 65, 85)):
    """Splits text by markdown inline formatting markers and adds runs to paragraph."""
    if not text:
        return
    
    tokens = inline_regex.split(text)
    for token in tokens:
        if not token:
            continue
        
        is_bold = False
        is_italic = False
        is_code = False
        clean_text = token
        
        if token.startswith("***") and token.endswith("***") and len(token) > 6:
            is_bold = True
            is_italic = True
            clean_text = token[3:-3]
        elif token.startswith("**") and token.endswith("**") and len(token) > 4:
            is_bold = True
            clean_text = token[2:-2]
        elif token.startswith("__") and token.endswith("__") and len(token) > 4:
            is_bold = True
            clean_text = token[2:-2]
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            is_italic = True
            clean_text = token[1:-1]
        elif token.startswith("_") and token.endswith("_") and len(token) > 2:
            is_italic = True
            clean_text = token[1:-1]
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            is_code = True
            clean_text = token[1:-1]
            
        run = paragraph.add_run(clean_text)
        if is_code:
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(239, 68, 68)  # soft red
            try:
                rPr = run._r.get_or_add_rPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                rPr.append(shd)
            except Exception:
                pass
        else:
            run.font.name = default_font
            run.font.size = default_size
            run.font.color.rgb = default_color
            run.bold = is_bold
            run.italic = is_italic

def set_cell_background(cell, hex_color: str):
    """Sets background shading of a cell in Word."""
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)
    except Exception:
        pass

def set_table_borders(table, hex_color: str = "CBD5E1"):
    """Applies minimalist horizontal borders to a table."""
    try:
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'  <w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)
    except Exception:
        pass

def _parse_and_render_markdown(doc_obj, md_text: str):
    """Parses structured Markdown text and appends formatted elements to docx."""
    if not md_text:
        return
        
    lines = md_text.split('\n')
    i = 0
    n = len(lines)
    
    while i < n:
        line = lines[i]
        stripped = line.strip()
        
        # 1. Empty lines
        if not stripped:
            i += 1
            continue
            
        # 2. Code blocks
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < n:
                i += 1
                
            code_text = "\n".join(code_lines)
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.left_indent = Pt(12)
            
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(15, 23, 42)
            
            try:
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
                pPr.append(shd)
                pbdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="CBD5E1"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pbdr)
            except Exception:
                pass
            continue
            
        # 3. Headings (H1, H2, H3)
        h_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if h_match:
            level = len(h_match.group(1))
            content = h_match.group(2).strip()
            
            p = doc_obj.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            if level == 1:
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
                _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(18), default_color=RGBColor(30, 58, 138))
                for r in p.runs:
                    r.bold = True
            elif level == 2:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(14), default_color=RGBColor(37, 99, 235))
                for r in p.runs:
                    r.bold = True
            else:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(12), default_color=RGBColor(30, 41, 59))
                for r in p.runs:
                    r.bold = True
            i += 1
            continue
            
        # 4. Horizontal Rule
        if stripped in ("---", "***", "___") or re.match(r'^[-*_]{3,}$', stripped):
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            try:
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E2E8F0"/></w:pBdr>')
                pPr.append(pBdr)
            except Exception:
                pass
            i += 1
            continue
            
        # 5. Tables
        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
                
            rows = []
            for tl in table_lines:
                cols = [col.strip() for col in tl.split("|")[1:-1]]
                rows.append(cols)
                
            if len(rows) > 1 and all(re.match(r'^:?-+:?$', c) for c in rows[1]):
                rows.pop(1)
                
            if rows:
                num_rows = len(rows)
                num_cols = max(len(r) for r in rows)
                
                table = doc_obj.add_table(rows=num_rows, cols=num_cols)
                table.autofit = True
                set_table_borders(table)
                
                for row_idx, r_data in enumerate(rows):
                    for col_idx, cell_val in enumerate(r_data):
                        if col_idx < num_cols:
                            cell = table.cell(row_idx, col_idx)
                            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                            p.paragraph_format.space_before = Pt(4)
                            p.paragraph_format.space_after = Pt(4)
                            p.paragraph_format.line_spacing = 1.0
                            
                            if row_idx == 0:
                                _add_inline_formatted_text(p, cell_val, default_font="Segoe UI", default_size=Pt(10), default_color=RGBColor(255, 255, 255))
                                for run in p.runs:
                                    run.bold = True
                                set_cell_background(cell, "1E3A8A")
                            else:
                                _add_inline_formatted_text(p, cell_val, default_font="Segoe UI", default_size=Pt(9.5), default_color=RGBColor(51, 65, 85))
                                if row_idx % 2 == 0:
                                    set_cell_background(cell, "F8FAFC")
                                else:
                                    set_cell_background(cell, "FFFFFF")
                
                # Empty spacer paragraph post-table
                p_after = doc_obj.add_paragraph()
                p_after.paragraph_format.space_before = Pt(0)
                p_after.paragraph_format.space_after = Pt(6)
            continue
            
        # 6. Lists
        bullet_match = re.match(r'^([\s]*)([-*+])\s+(.*)$', line)
        num_match = re.match(r'^([\s]*)(\d+)\.\s+(.*)$', line)
        if bullet_match or num_match:
            list_lines = []
            while i < n:
                curr_line = lines[i]
                if not curr_line.strip():
                    break
                b_m = re.match(r'^([\s]*)([-*+])\s+(.*)$', curr_line)
                n_m = re.match(r'^([\s]*)(\d+)\.\s+(.*)$', curr_line)
                if b_m or n_m:
                    list_lines.append((curr_line, b_m, n_m))
                    i += 1
                else:
                    if list_lines and curr_line.startswith("   "):
                        prev_text, prev_bm, prev_nm = list_lines[-1]
                        updated_text = prev_text + " " + curr_line.strip()
                        b_m_up = re.match(r'^([\s]*)([-*+])\s+(.*)$', updated_text)
                        n_m_up = re.match(r'^([\s]*)(\d+)\.\s+(.*)$', updated_text)
                        list_lines[-1] = (updated_text, b_m_up, n_m_up)
                        i += 1
                    else:
                        break
                        
            for item_line, bm, nm in list_lines:
                p = doc_obj.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                
                if bm:
                    indent_level = len(bm.group(1)) // 2
                    content = bm.group(3).strip()
                    p.style = 'List Bullet'
                    p.paragraph_format.left_indent = Pt(18 * (indent_level + 1))
                    _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(10.5), default_color=RGBColor(51, 65, 85))
                else:
                    indent_level = len(nm.group(1)) // 2
                    content = nm.group(3).strip()
                    p.style = 'List Number'
                    p.paragraph_format.left_indent = Pt(18 * (indent_level + 1))
                    _add_inline_formatted_text(p, content, default_font="Segoe UI", default_size=Pt(10.5), default_color=RGBColor(51, 65, 85))
            continue
            
        # 7. Standard Paragraph
        para_lines = []
        while i < n:
            curr_line = lines[i]
            curr_stripped = curr_line.strip()
            if not curr_stripped:
                break
            if (curr_stripped.startswith("```") or 
                re.match(r'^(#{1,6})\s+(.*)$', curr_stripped) or 
                curr_stripped.startswith("|") or 
                re.match(r'^([\s]*)([-*+])\s+(.*)$', curr_line) or 
                re.match(r'^([\s]*)(\d+)\.\s+(.*)$', curr_line) or
                curr_stripped in ("---", "***", "___") or 
                re.match(r'^[-*_]{3,}$', curr_stripped)):
                break
            para_lines.append(curr_stripped)
            i += 1
            
        if para_lines:
            para_text = " ".join(para_lines)
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            _add_inline_formatted_text(p, para_text, default_font="Segoe UI", default_size=Pt(10.5), default_color=RGBColor(51, 65, 85))

async def save_deep_research_as_word(
    account_id: str,
    query: str,
    report_text: str,
    workspace_id: Optional[str] = None
) -> Optional[uuid.UUID]:
    """
    Saves the deep research result parallelly to the user's OnlyOffice cloud documents
    inside a folder named 'Investigaciones Profundas'.
    Creates the folder automatically if it does not exist.
    """
    if not DOCX_AVAILABLE:
        logger.error("Cannot save Deep Research as Word document because python-docx is not installed.")
        return None

    try:
        account_uuid = uuid.UUID(account_id)
    except ValueError:
        logger.error(f"Invalid account_id format: {account_id}")
        return None

    workspace_uuid = None
    if workspace_id and workspace_id != "null" and workspace_id != "none":
        try:
            workspace_uuid = uuid.UUID(workspace_id)
        except ValueError:
            logger.warning(f"Invalid workspace_id format: {workspace_id}")

    logger.info(f"Initiating parallel Word document save for account {account_id}, workspace {workspace_id}...")

    async with DBSession(SessionLocal) as db_session:
        try:
            # 1. Check if the folder 'Investigaciones Profundas' exists
            stmt = select(DocumentFolder).where(
                DocumentFolder.account_id == account_uuid,
                DocumentFolder.name == "Investigaciones Profundas",
                DocumentFolder.workspace_id == workspace_uuid
            )
            result = await db_session.execute(stmt)
            folder = result.scalars().first()

            if not folder:
                logger.info("Folder 'Investigaciones Profundas' not found. Creating a new one...")
                folder = DocumentFolder(
                    account_id=account_uuid,
                    workspace_id=workspace_uuid,
                    name="Investigaciones Profundas"
                )
                db_session.add(folder)
                await db_session.commit()
                await db_session.refresh(folder)
                logger.info(f"Folder 'Investigaciones Profundas' created with ID: {folder.id}")

            # 2. Sanitize filename based on query slug
            query_slug = "".join(c if c.isalnum() or c in " _-" else "_" for c in query)
            query_slug = query_slug.strip()[:60]
            if not query_slug:
                query_slug = "Reporte"
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Investigacion_{query_slug}_{timestamp}.docx"
            unique_filename = f"{uuid.uuid4()}.docx"

            # 3. Physically create Word document in user's cloud folder
            user_dir = ensure_onlyoffice_account_dir(str(account_uuid))
            file_path = user_dir / unique_filename

            logger.info(f"Creating Word document at {file_path}...")
            doc = docx.Document()
            
            # Title page/Heading
            p_title = doc.add_paragraph()
            p_title.alignment = 1  # CENTER
            p_title.paragraph_format.space_before = Pt(12)
            p_title.paragraph_format.space_after = Pt(12)
            _add_inline_formatted_text(p_title, f"Investigación Profunda: {query}", default_size=Pt(20), default_color=RGBColor(30, 58, 138))
            for run in p_title.runs:
                run.bold = True

            # Spacer
            doc.add_paragraph()

            # Render markdown content
            _parse_and_render_markdown(doc, report_text)

            # Save the file
            doc.save(str(file_path))
            logger.info(f"Word document saved to filesystem successfully.")

            # 4. Register in database
            new_doc = Document(
                account_id=account_uuid,
                workspace_id=workspace_uuid,
                folder_id=folder.id,
                filename=filename,
                extension="docx",
                file_path=build_onlyoffice_relative_path(str(account_uuid), unique_filename)
            )
            db_session.add(new_doc)
            await db_session.commit()
            await db_session.refresh(new_doc)

            logger.info(f"Document registered in OnlyOffice database with ID: {new_doc.id}")
            return new_doc.id

        except Exception as e:
            logger.error(f"Error inside save_deep_research_as_word database/file operations: {e}", exc_info=True)
            await db_session.rollback()
            return None
